import sys
import uuid
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from client_delivery.extract import build_delivery_package  # noqa: E402
from client_delivery.quality import (  # noqa: E402
    MISSING_ACTION_OWNER_WARNING,
    NON_NUMERIC_THRESHOLD_WARNING,
    delivery_quality_warnings,
)
from client_delivery.render_docx import render_board_memo_docx  # noqa: E402
from client_delivery.render_xlsx import render_execution_tracker_xlsx  # noqa: E402
from tests.fixtures import fake_state  # noqa: E402


class OddObject:
    def __str__(self):
        return "odd-object"


def _docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _xlsx_values(path: Path) -> list[object]:
    workbook = load_workbook(path, data_only=False)
    values: list[object] = []
    for worksheet in workbook.worksheets:
        for row in worksheet.iter_rows(values_only=True):
            values.extend(value for value in row if value is not None)
    return values


def test_docx_is_readable(tmp_path):
    package = build_delivery_package(fake_state())
    path = render_board_memo_docx(package, tmp_path / "memo.docx")

    text = _docx_text(path)

    assert "Board Memo" in text
    assert "Recommendation" in text
    assert "Human Review" in text


def test_xlsx_is_readable(tmp_path):
    package = build_delivery_package(fake_state())
    path = render_execution_tracker_xlsx(package, tmp_path / "tracker.xlsx")

    workbook = load_workbook(path)

    assert workbook.sheetnames == [
        "Decision Summary",
        "30-60-90 Actions",
        "KPI Tracker",
        "Assumptions",
        "Review Triggers",
    ]


def test_renderers_tolerate_object_evidence(tmp_path):
    state = fake_state()
    state["execution_plan"][0]["evidence"] = [
        uuid.uuid4(),
        {"source": "dict"},
        7,
        None,
        ["nested", uuid.uuid4()],
        OddObject(),
    ]
    package = build_delivery_package(state)

    render_board_memo_docx(package, tmp_path / "memo.docx")
    render_execution_tracker_xlsx(package, tmp_path / "tracker.xlsx")


def test_renderers_use_deterministic_client_visible_serialization(tmp_path):
    state = fake_state()
    state["execution_plan"][0]["action"] = "  =cmd('calc')"
    state["execution_plan"][0]["owner"] = " +SUM(1,1)"
    state["execution_plan"][0]["evidence"] = [
        {"b": "second", "a": "first"},
        ["nested", {"source": "placeholder workflow review"}],
        "knowledge_alpha",
        "source_ref=upload:file-1:secret.csv#row=1",
        r"C:\Users\nicoc\secret.xlsx",
        "api_key=sk-test",
    ]
    state["execution_plan"][0]["notes"] = "@operator note"
    state["execution_plan"][0]["success_criteria"] = " -1+2"
    state["critical_assumptions"][0]["evidence"] = [{"source": "placeholder cadence agreement"}]
    state["kpis"][0]["notes"] = "=HYPERLINK(\"http://bad\")"
    package = build_delivery_package(state)
    package.execution_plan[0].action = "  =cmd('calc')"
    package.execution_plan[0].owner = " +SUM(1,1)"
    package.execution_plan[0].notes = "@operator note"
    package.execution_plan[0].success_criteria = " -1+2"

    docx_path = render_board_memo_docx(package, tmp_path / "memo.docx")
    xlsx_path = render_execution_tracker_xlsx(package, tmp_path / "tracker.xlsx")

    combined = _docx_text(docx_path) + "\n" + "\n".join(str(value) for value in _xlsx_values(xlsx_path))
    for forbidden in (
        '{"source"',
        "'source'",
        "knowledge_alpha",
        "source_ref=",
        "upload:file",
        r"C:\Users",
        "secret.xlsx",
        "api_key=sk-test",
    ):
        assert forbidden not in combined
    assert "first, second" in combined
    assert "nested, placeholder workflow review" in combined
    assert "placeholder cadence agreement" in combined
    assert "project evidence" in combined
    assert "Evidence source unavailable" in combined
    assert "redacted local path" in combined
    assert "credential=[redacted]" in combined

    values = _xlsx_values(xlsx_path)
    assert "'  =cmd('calc')" in values
    assert "' +SUM(1,1)" in values
    assert "'@operator note" in values
    assert "' -1+2" in values
    assert "'=HYPERLINK(\"http://bad\")" in values


def test_xlsx_preserves_numeric_kpi_threshold_cells(tmp_path):
    package = build_delivery_package(fake_state())
    path = render_execution_tracker_xlsx(package, tmp_path / "tracker.xlsx")
    worksheet = load_workbook(path, data_only=False)["KPI Tracker"]

    assert worksheet["C2"].value == 40
    assert worksheet["D2"].value == 60
    assert isinstance(worksheet["C2"].value, (int, float))
    assert isinstance(worksheet["D2"].value, (int, float))


def test_non_numeric_thresholds_do_not_crash(tmp_path):
    state = fake_state()
    state["kpis"][0]["threshold_red"] = "red if adoption is blocked"
    state["kpis"][0]["threshold_amber"] = "amber if adoption is uncertain"
    package = build_delivery_package(state)

    warnings = delivery_quality_warnings(package)
    path = render_execution_tracker_xlsx(package, tmp_path / "tracker.xlsx")

    assert path.exists()
    assert NON_NUMERIC_THRESHOLD_WARNING in warnings


def test_missing_action_owner_warning_is_emitted():
    state = fake_state()
    state["execution_plan"][1]["owner"] = ""
    package = build_delivery_package(state)

    warnings = delivery_quality_warnings(package)

    assert MISSING_ACTION_OWNER_WARNING in warnings
