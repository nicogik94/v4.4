"""Focused tests for the upload foundation and overview surface."""
import io
import os
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import AsyncMock, patch

from starlette.datastructures import Headers, UploadFile


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import api  # noqa: E402
import config  # noqa: E402
from config import UPLOAD_LAYER  # noqa: E402
from extensions.connectors import CSVColumnMapping  # noqa: E402
from knowledge.evidence_snapshot.capture import DeletionBlockedError  # noqa: E402
from knowledge.file_parsers import UploadParseError, parse_upload_bytes  # noqa: E402
from knowledge.files import UploadStorageError, delete_uploaded_file, ingest_uploaded_file  # noqa: E402
from knowledge.retrieval import evaluate_phase_retrieval  # noqa: E402
from overview import build_operator_overview  # noqa: E402
from tests.test_decision_objects import make_state  # noqa: E402


def _make_upload(filename: str, content: bytes, media_type: str) -> UploadFile:
    return UploadFile(file=io.BytesIO(content), filename=filename, headers=Headers({"content-type": media_type}))


def _make_docx_bytes(text: str) -> bytes:
    from docx import Document

    buffer = io.BytesIO()
    document = Document()
    document.add_heading("Context memo", level=1)
    document.add_paragraph(text)
    document.save(buffer)
    return buffer.getvalue()


def _make_pdf_bytes(text: str) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    pdf.drawString(72, 720, text[:100])
    pdf.save()
    return buffer.getvalue()


def _make_xlsx_bytes(rows: list[list[object]]) -> bytes:
    from openpyxl import Workbook

    buffer = io.BytesIO()
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Sheet1"
    for row in rows:
        worksheet.append(row)
    workbook.save(buffer)
    return buffer.getvalue()


class TestUploadParsers(unittest.TestCase):
    def test_supported_parsers_accept_all_supported_types(self):
        samples = [
            ("note.txt", "text/plain", b"Short context note for the operator.", "txt", "document"),
            ("note.md", "text/markdown", b"# Title\n\nMarkdown context.", "md", "document"),
            ("table.csv", "text/csv", b"name,value\nctr,0.42\n", "csv", "table"),
            (
                "sheet.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                _make_xlsx_bytes([["metric", "value"], ["ctr", 0.42]]),
                "xlsx",
                "table",
            ),
            (
                "memo.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                _make_docx_bytes("This file provides additional operating context."),
                "docx",
                "document",
            ),
            ("memo.pdf", "application/pdf", _make_pdf_bytes("PDF context for the decision engine"), "pdf", "document"),
        ]

        for filename, media_type, content, parser_kind, file_kind in samples:
            parsed = parse_upload_bytes(filename, media_type, content)
            self.assertEqual(parsed.parser_kind, parser_kind)
            self.assertEqual(parsed.file_kind, file_kind)

    def test_malformed_binary_files_raise_clear_errors(self):
        with self.assertRaises(UploadParseError):
            parse_upload_bytes("bad.pdf", "application/pdf", b"not a real pdf")
        with self.assertRaises(UploadParseError):
            parse_upload_bytes(
                "bad.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                b"not a docx",
            )
        with self.assertRaises(UploadParseError):
            parse_upload_bytes(
                "bad.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                b"not an xlsx",
            )

    def test_unsupported_type_is_rejected(self):
        with self.assertRaises(UploadParseError):
            parse_upload_bytes("photo.png", "image/png", b"fake")

    def test_json_upload_is_rejected_until_ingestion_is_enabled(self):
        with self.assertRaises(UploadParseError):
            parse_upload_bytes("cohort_snapshot.json", "application/json", b'{"retention": 0.82}')

    def test_oversize_upload_is_rejected(self):
        with patch.object(UPLOAD_LAYER, "max_file_bytes", 10):
            with self.assertRaises(UploadParseError):
                parse_upload_bytes("note.txt", "text/plain", b"This is longer than ten bytes.")


class TestUploadIngestion(unittest.TestCase):
    def test_document_upload_creates_source_manifest_and_knowledge_items_without_phase_changes(self):
        state = make_state("upload-doc")
        before_status = dict(state.phase_status)
        with tempfile.TemporaryDirectory() as tempdir, patch.object(UPLOAD_LAYER, "storage_dir", tempdir):
            result = ingest_uploaded_file(
                state,
                filename="brief.txt",
                media_type="text/plain",
                content=b"Fresh context for the project.\n\nDemand is shifting quickly.",
                actor="operator",
            )

            self.assertEqual(result.manifest.filename, "brief.txt")
            self.assertEqual(result.manifest.parse_summary.status.value, "completed")
            self.assertEqual(result.source.source_kind, "uploaded_file")
            self.assertGreaterEqual(len(result.knowledge_items), 1)
            self.assertTrue(Path(result.manifest.storage_ref).exists())
            self.assertEqual(dict(state.phase_status), before_status)

            dumped = state.model_dump()
            self.assertFalse(_contains_bytes(dumped))
            strategy_retrieval = evaluate_phase_retrieval(state, "strategy")
            self.assertGreaterEqual(len(strategy_retrieval.eligible_items), 1)
            overview = build_operator_overview(state)
            self.assertEqual(len(overview.files), 1)
            self.assertIn("uploaded file", overview.sources_and_files_message.lower())

    def test_table_upload_can_run_structured_import_and_delete_cleanly(self):
        state = make_state("upload-table")
        mapping = [
            CSVColumnMapping(column="title", target_type="evidence", target_field="title", required=True),
            CSVColumnMapping(column="summary", target_type="evidence", target_field="summary"),
            CSVColumnMapping(column="signal_name", target_type="signal", target_field="name"),
            CSVColumnMapping(column="signal_desc", target_type="signal", target_field="description"),
        ]
        rows = _make_xlsx_bytes(
            [
                ["title", "summary", "signal_name", "signal_desc"],
                ["Urgent metric", "CTR is falling", "CTR", "Weekly click-through rate"],
            ]
        )
        with tempfile.TemporaryDirectory() as tempdir, patch.object(UPLOAD_LAYER, "storage_dir", tempdir):
            result = ingest_uploaded_file(
                state,
                filename="metrics.xlsx",
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                content=rows,
                actor="operator",
                role="data",
                import_mode="structured_import",
                mapping=mapping,
            )

            self.assertEqual(result.manifest.import_mode, "structured_import")
            self.assertEqual(len(state.imported_evidence), 1)
            self.assertEqual(len(state.imported_signals), 1)
            self.assertIsNotNone(result.import_summary)
            self.assertEqual(result.import_summary["evidence_count"], 1)
            self.assertEqual(result.import_summary["signal_count"], 1)

            file_id = result.manifest.file_id
            storage_ref = result.manifest.storage_ref
            delete_result = delete_uploaded_file(state, file_id)
            self.assertTrue(delete_result["deleted"])
            self.assertFalse(Path(storage_ref).exists())
            self.assertEqual(len(state.imported_evidence), 0)
            self.assertEqual(len(state.imported_signals), 0)
            self.assertEqual(len(state.knowledge_layer.uploaded_files), 0)

    def test_storage_mkdir_failure_raises_controlled_error_without_manifest(self):
        state = make_state("upload-storage-failure")

        with patch("knowledge.files.Path.mkdir", side_effect=OSError(5, "Input/output error")):
            with self.assertRaises(UploadStorageError) as ctx:
                ingest_uploaded_file(
                    state,
                    filename="brief.txt",
                    media_type="text/plain",
                    content=b"Fresh context for the project.",
                    actor="operator",
                )

        self.assertEqual(str(ctx.exception), UploadStorageError.public_message)
        self.assertNotIn("Input/output error", str(ctx.exception))
        self.assertEqual(len(state.knowledge_layer.uploaded_files), 0)


class TestUploadApiAndOverview(unittest.IsolatedAsyncioTestCase):
    async def asyncSetUp(self):
        api.running.clear()

    async def asyncTearDown(self):
        api.running.clear()

    async def test_upload_route_and_overview_surface_are_additive(self):
        state = make_state("upload-api")
        before_status = dict(state.phase_status)
        upload = _make_upload("context.md", b"# Market note\n\nPricing pressure is rising.", "text/markdown")

        with tempfile.TemporaryDirectory() as tempdir, patch.object(UPLOAD_LAYER, "storage_dir", tempdir):
            with patch("api.store.load", new=AsyncMock(return_value=state)), patch("api.store.save", new=AsyncMock()) as save_mock:
                response = await api.upload_project_file(
                    state.project_id,
                    file=upload,
                    actor="operator",
                    role="context",
                    import_mode="knowledge",
                    sheet_name="",
                    mapping_json="",
                )
                overview = await api.get_overview(state.project_id)
                files = await api.get_uploaded_files(state.project_id)
                detail = await api.get_uploaded_file(state.project_id, response["manifest"]["file_id"])

        self.assertEqual(response["status"], "uploaded")
        self.assertEqual(len(files["files"]), 1)
        self.assertEqual(detail["manifest"]["filename"], "context.md")
        self.assertEqual(overview.files[0].filename, "context.md")
        self.assertEqual(dict(state.phase_status), before_status)
        save_mock.assert_awaited()

    async def test_upload_route_returns_controlled_503_for_storage_failure(self):
        state = make_state("upload-api-storage-failure")
        upload = _make_upload("context.md", b"# Market note\n\nPricing pressure is rising.", "text/markdown")
        storage_error = UploadStorageError(
            path=r"C:\private\upload_store\upload-api-storage-failure",
            operation="mkdir",
            cause=OSError(5, "Input/output error"),
        )

        with patch("api.store.load", new=AsyncMock(return_value=state)):
            with patch("api.ingest_uploaded_file", side_effect=storage_error):
                with self.assertRaises(api.HTTPException) as ctx:
                    await api.upload_project_file(
                        state.project_id,
                        file=upload,
                        actor="operator",
                        role="context",
                        import_mode="knowledge",
                        sheet_name="",
                        mapping_json="",
                    )

        self.assertEqual(ctx.exception.status_code, 503)
        detail = str(ctx.exception.detail)
        self.assertIn("Upload storage is unavailable", detail)
        self.assertNotIn("Traceback", detail)
        self.assertNotIn("Input/output error", detail)
        self.assertNotIn("C:\\", detail)


class TestDeleteApiSnapshotGuard(unittest.IsolatedAsyncioTestCase):
    """Regression tests for the Slice A deletion guard at the DELETE API boundary."""

    async def asyncSetUp(self):
        api.running.clear()

    async def asyncTearDown(self):
        api.running.clear()

    def _state_with_one_upload(self, name: str, tempdir: str):
        state = make_state(name)
        with patch.object(UPLOAD_LAYER, "storage_dir", tempdir):
            result = ingest_uploaded_file(
                state,
                filename="evidence.txt",
                media_type="text/plain",
                content=b"context linked to an evidence snapshot",
                actor="operator",
            )
        return state, result.manifest

    async def test_linked_upload_delete_returns_409_and_leaves_state_unchanged(self):
        with tempfile.TemporaryDirectory() as tempdir, patch.object(UPLOAD_LAYER, "storage_dir", tempdir):
            state, manifest = self._state_with_one_upload("delete-linked", tempdir)
            files_before = len(state.knowledge_layer.uploaded_files)

            with patch("api.store.load", new=AsyncMock(return_value=state)), \
                    patch("api.store.save", new=AsyncMock()) as save_mock, \
                    patch(
                        "knowledge.evidence_snapshot.capture.assert_safe_to_delete_storage_ref",
                        side_effect=DeletionBlockedError("linked to snapshot"),
                    ):
                with self.assertRaises(api.HTTPException) as ctx:
                    await api.delete_project_file(state.project_id, manifest.file_id)

            # Boundary translated to a clear 409 with a safe, concise message.
            self.assertEqual(ctx.exception.status_code, 409)
            detail = str(ctx.exception.detail)
            self.assertIn("evidence snapshot", detail.lower())
            self.assertNotIn("Traceback", detail)
            # Fail-closed: no storage deletion, no state mutation, no persistence.
            self.assertTrue(Path(manifest.storage_ref).exists())
            self.assertEqual(len(state.knowledge_layer.uploaded_files), files_before)
            save_mock.assert_not_awaited()

    async def test_unlinked_upload_delete_succeeds(self):
        with tempfile.TemporaryDirectory() as tempdir, patch.object(UPLOAD_LAYER, "storage_dir", tempdir):
            state, manifest = self._state_with_one_upload("delete-unlinked", tempdir)

            # Flag-on but the storage_ref is not snapshot-linked: guard is a no-op.
            with patch("api.store.load", new=AsyncMock(return_value=state)), \
                    patch("api.store.save", new=AsyncMock()) as save_mock, \
                    patch(
                        "knowledge.evidence_snapshot.capture.assert_safe_to_delete_storage_ref",
                        return_value=None,
                    ):
                result = await api.delete_project_file(state.project_id, manifest.file_id)

            self.assertTrue(result["deleted"])
            self.assertFalse(Path(manifest.storage_ref).exists())
            self.assertEqual(len(state.knowledge_layer.uploaded_files), 0)
            save_mock.assert_awaited()

    async def test_flag_off_delete_behaves_normally(self):
        with tempfile.TemporaryDirectory() as tempdir, patch.object(UPLOAD_LAYER, "storage_dir", tempdir), \
                patch.dict(os.environ, {"MAS_EVIDENCE_SNAPSHOT_ENABLED": "false"}):
            # Genuine flag-off path: the real guard short-circuits to a no-op.
            self.assertFalse(config.evidence_snapshot_enabled())
            state, manifest = self._state_with_one_upload("delete-flag-off", tempdir)

            with patch("api.store.load", new=AsyncMock(return_value=state)), \
                    patch("api.store.save", new=AsyncMock()) as save_mock:
                result = await api.delete_project_file(state.project_id, manifest.file_id)

            self.assertTrue(result["deleted"])
            self.assertFalse(Path(manifest.storage_ref).exists())
            self.assertEqual(len(state.knowledge_layer.uploaded_files), 0)
            save_mock.assert_awaited()


def _contains_bytes(value) -> bool:
    if isinstance(value, (bytes, bytearray)):
        return True
    if isinstance(value, dict):
        return any(_contains_bytes(item) for item in value.values())
    if isinstance(value, list):
        return any(_contains_bytes(item) for item in value)
    return False


if __name__ == "__main__":
    unittest.main(verbosity=2)
