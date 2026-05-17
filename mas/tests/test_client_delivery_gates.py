import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = ROOT.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from client_delivery.service import generate_client_delivery_package  # noqa: E402
from tests.fixtures import fake_state  # noqa: E402


EXPECTED_SHEETS = {
    "Decision Summary",
    "30-60-90 Actions",
    "KPI Tracker",
    "Assumptions",
    "Review Triggers",
}
EXPECTED_ACTION_OWNERS = ["Pilot owner", "Operations lead", "Executive sponsor"]


def _manifest(path: str | Path) -> dict:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def _docx_action_owners(path: str | Path) -> list[str]:
    document = Document(path)
    for table in document.tables:
        if not table.rows:
            continue
        headers = [cell.text for cell in table.rows[0].cells]
        if "Action" not in headers or "Owner" not in headers:
            continue
        owner_idx = headers.index("Owner")
        return [row.cells[owner_idx].text for row in table.rows[1:]]
    return []


def _xlsx_action_owners(path: str | Path) -> list[str]:
    worksheet = load_workbook(path)["30-60-90 Actions"]
    headers = [cell.value for cell in worksheet[1]]
    owner_idx = headers.index("owner")
    return [row[owner_idx] for row in worksheet.iter_rows(min_row=2, values_only=True)]


def test_no_v06_modules_present():
    blocked = [
        ROOT / "client_delivery" / "render_pptx.py",
        ROOT / "client_delivery" / "render_five_pager.py",
        ROOT / "monitoring",
        ROOT / "client_delivery" / "slack.py",
        ROOT / "client_delivery" / "email_digest.py",
        ROOT / "client_delivery" / "bayesian_feedback.py",
        ROOT / "client_delivery" / "sheets_api.py",
        ROOT / "client_delivery" / "connectors",
    ]

    assert all(not path.exists() for path in blocked)


def test_manifest_carries_validation_status(tmp_path):
    result = generate_client_delivery_package(fake_state(), tmp_path)

    manifest = _manifest(result["outputs"]["manifest_json"])

    assert manifest["validation_status"] == "awaiting_side_by_side_defense_test"


def test_worked_example_script_runs():
    completed = subprocess.run(
        [sys.executable, "-m", "scripts.render_example"],
        cwd=REPO_ROOT,
        check=False,
        capture_output=True,
        text=True,
    )

    assert completed.returncode == 0, completed.stderr
    result = json.loads(completed.stdout)
    outputs = result["outputs"]

    assert Document(outputs["board_memo_docx"])
    workbook = load_workbook(outputs["execution_tracker_xlsx"])
    assert set(workbook.sheetnames) == EXPECTED_SHEETS
    manifest = _manifest(outputs["manifest_json"])
    assert manifest["quality_warnings"] == []
    assert _docx_action_owners(outputs["board_memo_docx"]) == EXPECTED_ACTION_OWNERS
    assert _xlsx_action_owners(outputs["execution_tracker_xlsx"]) == EXPECTED_ACTION_OWNERS


def test_example_input_exercises_all_fields():
    payload = json.loads((REPO_ROOT / "examples" / "sunforest_redacted_state.json").read_text(encoding="utf-8"))

    assert len(payload["strategy"]["strategies"]) >= 3
    assert len(payload["critical_assumptions"]) >= 2
    assert len(payload["kpis"]) >= 3
    assert len(payload["review"]["reentry_triggers"]) >= 1


def test_xlsx_sheet_count_ceiling(tmp_path):
    result = generate_client_delivery_package(fake_state(), tmp_path)
    workbook = load_workbook(result["outputs"]["execution_tracker_xlsx"])

    assert len(workbook.sheetnames) <= 5
    assert set(workbook.sheetnames) == EXPECTED_SHEETS


def test_docx_heading_count_ceiling(tmp_path):
    result = generate_client_delivery_package(fake_state(), tmp_path)
    document = Document(result["outputs"]["board_memo_docx"])
    top_level_headings = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style and paragraph.style.name in {"Title", "Heading 1"}
    ]

    assert len(top_level_headings) <= 7
