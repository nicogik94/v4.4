"""Tests for profile-based project exports."""
import json
import os
import re
import sys
import unittest
import zipfile
from io import BytesIO
from pathlib import Path
from unittest.mock import AsyncMock, patch

from docx import Document
from fastapi import HTTPException
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import api  # noqa: E402
import report_freshness  # noqa: E402
from report_quality import (  # noqa: E402
    CLIENT_BF_CONFIDENCE_CAVEAT,
    PROVISIONAL_CLARIFICATION_CAVEAT,
    PROVISIONAL_CLARIFICATION_NEXT_ACTION,
    RISK_CLASSIFICATION_WARNING,
    THRESHOLD_CONFLICT_UNKNOWN_WARNING,
)
from clarifications import (  # noqa: E402
    ClarificationAnswer,
    ClarificationCycle,
    ClarificationPriority,
    ClarificationQuestion,
    ClarificationStatus,
)
from exporters import (  # noqa: E402
    MONITORING_TEMPLATE_OPERATOR_NOTE,
    _safe_report_markdown,
    build_client_dossier_markdown,
    build_export_manifest,
    build_machine_archive_payload,
    build_operator_dossier_markdown,
    export_project_profile_bytes,
    operator_monitoring_summary,
    sanitize_for_export,
)
from monitoring_templates import (  # noqa: E402
    CLIENT_MONITORING_TEMPLATE_HEADERS,
    OPERATOR_MONITORING_TEMPLATE_HEADERS,
    SHEET_NAME,
    monitoring_template_cell_rows,
)
from state import (  # noqa: E402
    AuditOutput,
    ClassifyOutput,
    FileParseStatus,
    FileParseSummary,
    FMEAItem,
    Hypothesis,
    KnowledgeItem,
    KnowledgeLayerState,
    MonitorCanary,
    MonitorCircuitBreaker,
    MonitorOODASchedule,
    MonitorOutput,
    MonitorScheduleItem,
    PhaseStatus,
    PreliminaryVerdict,
    ProjectState,
    StrategyOutput,
    UploadedFileManifest,
    Verdict,
)
from tests.test_decision_objects import make_state  # noqa: E402


def _docx_text(payload: bytes) -> str:
    document = Document(BytesIO(payload))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _docx_table_count(payload: bytes) -> int:
    return len(Document(BytesIO(payload)).tables)


def _xlsx_rows(payload: bytes) -> tuple[tuple[str, ...], ...]:
    workbook = load_workbook(BytesIO(payload), data_only=False)
    worksheet = workbook[SHEET_NAME]
    rows = []
    for row in worksheet.iter_rows(values_only=True):
        rows.append(tuple("" if value is None else str(value) for value in row))
    return tuple(rows)


def _attach_report_generation_metadata(
    state: ProjectState,
    *,
    code_version: str,
    generated_at: str = "2026-05-01T00:00:00Z",
) -> None:
    state.policy_audit_log.append(
        {
            "ts": 1770000000.0,
            "event_type": "report_generated",
            "phase": "report",
            "details": report_freshness.build_report_generation_metadata(
                state.report,
                code_version=code_version,
                generated_at=generated_at,
            ),
        }
    )


CLIENT_HEADINGS = [
    "What decision we reviewed",
    "Recommended path",
    "Why this is recommended",
    "What evidence was used",
    "What should happen next",
    "Timeline / 7-30-60-90 roadmap",
    "Key risks",
    "What to monitor",
    "Open assumptions / questions",
    "Human review note",
]


OPERATOR_HEADINGS = [
    "Cover / project metadata",
    "Executive summary",
    "Current recommendation",
    "Decision snapshot",
    "Phase completion status",
    "Dashboard overview",
    "Original input",
    "Classification summary",
    "Hypotheses table",
    "Gauntlet / stress-test summary",
    "Audit findings",
    "Evidence and source summary",
    "Strategy plan",
    "SQI / quality review",
    "Monitoring plan",
    "Workspace summary",
    "Risks and open questions",
    "Decision trace / explainability",
    "Clarifications / assumptions",
    "Report appendix",
    "Technical appendix",
]


def _assert_in_order(testcase: unittest.TestCase, text: str, headings: list[str]) -> None:
    position = -1
    for heading in headings:
        next_position = text.find(heading)
        testcase.assertGreater(next_position, position, heading)
        position = next_position


def make_export_state(project_id: str = "export-profile"):
    state = make_state(project_id)
    state.report = """# Executive Summary
Proceed with a bounded audit pilot [Evidence: ev-market | chunk=1].

# The Decision
Decide whether to fund the pilot.

# Recommended Path
Run the 30-day pilot before the full program.

# Why This Is Recommended
The pilot limits commitment while validating demand.

# Options Considered
| Option | Verdict |
|---|---|
| Pilot | Recommended |

# Evidence Used
| Evidence | What it suggests |
|---|---|
| Market note | There is a concrete source marker. |

# Key Risks
Scope creep and unclear ownership.

# Assumptions and Open Questions
Who owns the pilot remains open.

# Roadmap
Day 1-30: run the audit.

# Next Steps
- Confirm owner.
- Confirm data access.

# Monitoring and Kill Criteria
Stop if data access is unavailable.

# Appendix: Technical Analysis
Technical notes stay here.
"""
    for phase in ("classify", "hypotheses", "gauntlet", "audit", "strategy", "sqi", "monitor", "report"):
        state.phase_status[phase] = PhaseStatus.COMPLETED
        state.phase_confidence[phase] = 1.0
        state.phase_summaries[phase] = f"{phase} summary"
    state.knowledge_layer = KnowledgeLayerState(
        items=[
            KnowledgeItem(
                item_id="ev-market",
                source_id="src-1",
                title="Market note",
                source_ref="upload:file-1:market.pdf#chunk=1",
                structured_payload={"locator": "chunk=1"},
            )
        ],
        uploaded_files=[
            UploadedFileManifest(
                file_id="file-1",
                source_id="src-1",
                filename="market.pdf",
                media_type="application/pdf",
                size_bytes=321,
                storage_ref=r"C:\Users\nicoc\private\market.pdf",
            )
        ],
    )
    state.clarification_cycles = [
        ClarificationCycle(
            cycle_id="cycle-1",
            project_id=state.project_id,
            questions=[
                ClarificationQuestion(
                    question_id="q1",
                    text="Who owns the pilot?",
                    why_it_matters="Ownership controls follow-through.",
                    priority=ClarificationPriority.CRITICAL,
                    affected_phase="strategy",
                    source_gap="stakeholder_audience",
                    status=ClarificationStatus.OPEN,
                )
            ],
        )
    ]
    state.policy_audit_log = [
        {
            "event_type": "policy_gate_blocked",
            "phase": "strategy",
            "details": {
                "api_key": "sk-test-secret",
                "raw_prompt": "hidden prompt",
                "raw_provider_payload": {"token": "provider-token"},
                "chain_of_thought": "hidden reasoning",
                "scenario_probability": 0.91,
                "local_path": r"C:\Users\nicoc\private\payload.json",
            },
        }
    ]
    state.data = r"Local notes at C:\Users\nicoc\private\data.csv with api_key=sk-test-secret"
    return state


def make_sparse_growth_state(project_id: str = "sparse-growth"):
    state = ProjectState(
        project_id=project_id,
        project_name="Improving growth performance",
        brief="Improve growth performance across revenue operations, retention, churn, acquisition, and pipeline.",
        report="""# Executive Summary
Recommend a diagnostic hold before scaling. This has high risk and circuit breaker implications.
H1 model-generated prior probability 0. 68 and structured risk priority 70. 0 [Evidence: ev-growth | citation unavailable].

# The Decision
Decide whether to scale growth spend before measurement is repaired.

# Recommended Path
1. Repair tracking.
1. Reconcile billing and product metrics.
1. Interview churned customers.
DQ greater than 70 and BF greater than 10 are required before major spend.
model-generated prior probability 0. 68 and structured risk priority 70. 0 should remain provisional.

# Monitoring and Kill Criteria
Top channel greater than 70 and CAC worsening. Stop threshold is 2. 1 issues.

# Evidence Used
No uploaded files, imported evidence, imported signals, or concrete locators are available.

# Roadmap
Run Sprint 0 first.
""",
    )
    return state


class TestCodeVersionFreshness(unittest.TestCase):
    def test_current_code_version_uses_safe_env_var_first(self):
        env = {
            "V4_CODE_VERSION": "abc123",
            "GIT_COMMIT": "",
            "COMMIT_SHA": "",
            "SOURCE_VERSION": "",
        }
        with patch.dict(os.environ, env, clear=False):
            with patch("report_freshness.subprocess.run") as run_mock:
                self.assertEqual(report_freshness.current_code_version(), "abc123")

        run_mock.assert_not_called()

    def test_report_generation_metadata_uses_app_version_by_default(self):
        env = {name: "" for name in report_freshness.CODE_VERSION_ENV_VARS}
        with patch.dict(os.environ, env, clear=False):
            metadata = report_freshness.build_report_generation_metadata("current report")

        self.assertEqual(metadata["code_version"], "4.4.0")
        self.assertNotEqual(metadata["code_version"], report_freshness.UNKNOWN_CODE_VERSION)

    def test_current_code_version_ignores_unsafe_env_and_uses_app_version(self):
        env = {
            "V4_CODE_VERSION": r"abc123 C:\Users\example\secret.txt",
            "GIT_COMMIT": "",
            "COMMIT_SHA": "",
            "SOURCE_VERSION": "",
        }
        with patch.dict(os.environ, env, clear=False):
            with patch("report_freshness.subprocess.run") as run_mock:
                self.assertEqual(report_freshness.current_code_version(), "4.4.0")

        run_mock.assert_not_called()

    def test_current_code_version_returns_unknown_only_when_env_app_and_git_fail(self):
        env = {name: "" for name in report_freshness.CODE_VERSION_ENV_VARS}
        with patch.dict(os.environ, env, clear=False):
            with patch("report_freshness.APP_VERSION", ""):
                with patch("report_freshness._detect_repo_root", return_value=ROOT.parent):
                    with patch("report_freshness.subprocess.run", side_effect=FileNotFoundError):
                        self.assertEqual(report_freshness.current_code_version(), report_freshness.UNKNOWN_CODE_VERSION)

    def test_export_manifest_uses_code_version_helper_for_manifest_and_freshness(self):
        state = make_export_state("manifest-version")
        _attach_report_generation_metadata(state, code_version="same999")

        with patch("report_freshness.current_code_version", return_value="same999"):
            manifest = build_export_manifest(state, "machine_archive", "zip")

        self.assertEqual(manifest["code_version"], "same999")
        self.assertEqual(manifest["report_freshness"]["current_code_version"], "same999")
        self.assertEqual(manifest["report_freshness"]["generated_code_version"], "same999")
        self.assertEqual(manifest["report_freshness"]["status"], "fresh")


class TestProfileExporterHelpers(unittest.TestCase):
    def test_report_profile_exports_report_only(self):
        state = make_export_state("report-only")
        state.report = "Standalone report only."

        payload, media_type, filename = export_project_profile_bytes(state, "report", "docx")
        text = _docx_text(payload)

        self.assertEqual(media_type, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertIn("report-only-report-", filename)
        self.assertIn("Standalone report only.", text)
        self.assertNotIn("Executive Summary", text)
        self.assertNotIn("classify summary", text)

    def test_all_profiles_export_valid_formats(self):
        state = make_export_state("all-profiles")
        expected = {
            ("report", "docx"): "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ("report", "pdf"): "application/pdf",
            ("client_dossier", "docx"): "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ("client_dossier", "pdf"): "application/pdf",
            ("client_monitoring_template", "xlsx"): "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ("operator_dossier", "docx"): "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ("operator_dossier", "pdf"): "application/pdf",
            ("operator_monitoring_template", "xlsx"): "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ("machine_archive", "zip"): "application/zip",
        }

        for (profile, fmt), media_type in expected.items():
            with self.subTest(profile=profile, format=fmt):
                payload, actual_media_type, filename = export_project_profile_bytes(state, profile, fmt)
                self.assertGreater(len(payload), 100)
                self.assertEqual(actual_media_type, media_type)
                self.assertIn(f"{profile}-", filename)
                if fmt == "pdf":
                    self.assertTrue(payload.startswith(b"%PDF"))
                if fmt == "zip":
                    with zipfile.ZipFile(BytesIO(payload)) as archive:
                        self.assertNotIn("raw_project_state.json", archive.namelist())
                if fmt == "xlsx":
                    self.assertEqual(_xlsx_rows(payload)[0], (
                        tuple(OPERATOR_MONITORING_TEMPLATE_HEADERS)
                        if profile.startswith("operator")
                        else tuple(CLIENT_MONITORING_TEMPLATE_HEADERS)
                    ))

    def test_monitoring_template_profiles_export_xlsx_with_stable_headers(self):
        state = make_export_state("monitor-template")
        state.report = """# Executive Summary
Proceed with the pilot.

# Decision Gates
| Signal to watch | Good sign | Warning sign | Stop/change-course threshold | Owner / role | Review cadence | Action if triggered | Evidence source |
|---|---|---|---|---|---|---|---|
| Activation quality | at least 12 qualified pilots | fewer than 6 qualified pilots | stop if data access is unavailable | Product Lead | Weekly | pause rollout | Market note |

# Monitoring and Kill Criteria
This section is narrative only and should not replace Decision Gates.
"""

        payload, media_type, filename = export_project_profile_bytes(state, "client_monitoring_template", "xlsx")
        rows = _xlsx_rows(payload)

        self.assertEqual(media_type, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        self.assertIn("monitor-template-client_monitoring_template-", filename)
        self.assertTrue(filename.endswith(".xlsx"))
        self.assertEqual(rows[0], tuple(CLIENT_MONITORING_TEMPLATE_HEADERS))
        self.assertEqual(rows[1][0], "Activation quality")
        self.assertEqual(rows[1][7], "stop if data access is unavailable")
        self.assertEqual(rows[1][8], "pause rollout")
        self.assertGreaterEqual(len(rows), 2)

    def test_monitoring_template_uses_safe_placeholders_for_missing_values(self):
        state = ProjectState(project_id="monitor-placeholders", project_name="Monitor placeholders", brief="Brief")

        payload, _, _ = export_project_profile_bytes(state, "client_monitoring_template", "xlsx")
        rows = _xlsx_rows(payload)
        combined = "\n".join("\t".join(row) for row in rows)

        self.assertIn("Operator to define", combined)
        self.assertIn("Evidence source unavailable", combined)
        self.assertIn("Threshold not yet confirmed", combined)
        self.assertIn("Validation required", combined)

    def test_client_monitoring_template_uses_concrete_values_and_direction_fidelity(self):
        state = ProjectState(
            project_id="monitor-fidelity",
            project_name="Monitor fidelity",
            brief="Decide whether activation and lead-quality remediation are ready to scale.",
        )
        state.report = """# Decision Gates
| Signal to watch | Good sign | Warning sign | Stop/change-course threshold | Owner / role | Review cadence | Action if triggered | Evidence source |
|---|---|---|---|---|---|---|---|
| Activation rate | >= 40% by Day 7 | below 30% by Day 10 | stop if under 30% within 14 days | Growth Lead | 7-day rolling | pause rollout | Product telemetry |

# Monitoring and Kill Criteria
Decision Gates remain the threshold source of truth.
"""
        state.hypotheses = [
            Hypothesis(
                id="H1",
                text="Lead quality controls activation-to-conversion lag.",
                signal="Lead quality conversion",
                confirm="lead quality improves 1.5pp within 14 days",
                reject="time-to-value stays above 72 hours",
                evidence_ids=["ev-market"],
            )
        ]
        state.strategy = StrategyOutput(
            preliminary_verdicts=[
                PreliminaryVerdict(
                    id="H1",
                    verdict=Verdict.NEEDS_MONITORING,
                    evidence="ev-market",
                    monitoring_plan="Track activation-to-conversion lag down 0.5pp by Day 10.",
                )
            ],
            success_metrics=[
                "Time-to-value down 24 hours",
                "Activation-to-conversion lag down 1.5pp within 14 days",
            ],
            review_date="Day 10",
        )
        state.monitor = MonitorOutput(
            ooda_schedule=MonitorOODASchedule(
                weekly=[
                    MonitorScheduleItem(
                        metric="Activation rate",
                        owner="Growth Lead",
                        source="Product telemetry 40% by Day 7",
                    )
                ]
            ),
            canaries=[
                MonitorCanary(
                    signal="activation-to-conversion lag",
                    direction="up",
                    window="14-day rolling",
                    meaning="lag should not increase",
                )
            ],
            circuit_breakers=[
                MonitorCircuitBreaker(
                    strategy_ref="H1 lead-quality follow-up",
                    trip="time-to-value above 72 hours",
                    reset="below 48 hours for 7-day rolling",
                )
            ],
        )

        client_payload, _, _ = export_project_profile_bytes(state, "client_monitoring_template", "xlsx")
        client_rows = _xlsx_rows(client_payload)
        client_combined = "\n".join("\t".join(row) for row in client_rows)

        for weak_placeholder in ("Validation required", "Threshold not yet confirmed", "Operator to define"):
            self.assertNotIn(weak_placeholder, client_combined)
        for concrete in (
            "72 hours",
            "24 hours",
            "48 hours",
            "7-day rolling",
            "14-day rolling",
            "Day 7",
            "Day 10",
            "within 14 days",
            "40%",
            "30%",
            "1.5pp",
            "0.5pp",
        ):
            self.assertIn(concrete, client_combined)
        self.assertNotRegex(client_combined, r"\bH(10|[1-9])\b")
        self.assertIn("hypothesis 1", client_combined)
        self.assertIn("Expected trend under remediation: down", client_combined)
        self.assertIn("Expected trend under remediation: up", client_combined)
        self.assertNotIn("Direction: up", client_combined)
        self.assertNotIn("Signal direction: up", client_combined)

        operator_payload, _, _ = export_project_profile_bytes(state, "operator_monitoring_template", "xlsx")
        operator_combined = "\n".join("\t".join(row) for row in _xlsx_rows(operator_payload))
        self.assertRegex(operator_combined, r"\bH1\b")
        self.assertIn("72 hours", operator_combined)
        self.assertIn("1.5pp", operator_combined)

    def test_monitoring_template_ambiguous_decision_gate_uses_placeholder(self):
        state = make_export_state("ambiguous-gates")
        state.report = """# Decision Gates
| Topic | Comment |
|---|---|
| Pilot | Needs discussion |

# Evidence Used
| Evidence | What it suggests |
|---|---|
| ev-market | Should remain unrelated to gate parsing |
"""

        payload, _, _ = export_project_profile_bytes(state, "operator_monitoring_template", "xlsx")
        rows = _xlsx_rows(payload)

        self.assertEqual(rows[1][0], "Decision Gate")
        self.assertEqual(rows[1][7], "Threshold not yet confirmed")
        self.assertIn("not parsed into a clear gate table", "\n".join(rows[1]))
        self.assertNotIn("ev-market", rows[1])

    def test_monitoring_template_client_redacts_internal_ids_refs_and_formula_cells(self):
        state = make_export_state("client-monitor-safe")
        state.report = r"""# Decision Gates
| Signal to watch | Good sign | Warning sign | Stop/change-course threshold | Owner / role | Review cadence | Action if triggered | Evidence source | Notes |
|---|---|---|---|---|---|---|---|---|
| =Activation | +good | @warning | -stop if BF 12 or RPN 90 worsens | @Owner | Weekly | =HYPERLINK("http://bad") | upload:file-1:metrics.csv#row=2 | ev-market knowledge_alpha storage_ref=C:\Users\nicoc\secret.xlsx operator trace |
"""
        state.monitor = MonitorOutput(
            ooda_schedule=MonitorOODASchedule(
                daily=[MonitorScheduleItem(metric="=CTR", owner="+Owner", source="source_ref=upload:file-2:private.csv#row=1")]
            ),
            circuit_breakers=[MonitorCircuitBreaker(strategy_ref="S1", trip="-stop now", reset="@reset")],
            canaries=[MonitorCanary(signal="@canary", direction="up", window="7d", meaning="+lift")],
        )

        payload, _, _ = export_project_profile_bytes(state, "client_monitoring_template", "xlsx")
        rows = _xlsx_rows(payload)
        combined = "\n".join("\t".join(row) for row in rows)

        self.assertIn("'=Activation", combined)
        self.assertIn("'+good", combined)
        self.assertIn("'@warning", combined)
        self.assertIn("'-stop if internal diagnostic redacted", combined)
        self.assertIn("Uploaded project document", combined)
        for forbidden in (
            "ev-market",
            "knowledge_alpha",
            "upload:",
            "storage_ref",
            "source_ref",
            r"C:\Users\nicoc",
            "BF 12",
            "RPN 90",
            "operator trace",
        ):
            self.assertNotIn(forbidden, combined)

    def test_monitoring_template_operator_retains_allowed_trace_after_redaction(self):
        state = make_export_state("operator-monitor-trace")
        state.report = r"""# Decision Gates
| Signal to watch | Stop/change-course threshold | Action if triggered | Evidence source |
|---|---|---|---|
| Pipeline quality | stop if below target [Evidence: ev-market] | Escalate | upload:file-1:metrics.csv#row=2 C:\Users\nicoc\secret.xlsx |
"""

        payload, _, _ = export_project_profile_bytes(state, "operator_monitoring_template", "xlsx")
        rows = _xlsx_rows(payload)
        header = rows[0]
        data = rows[1]

        self.assertEqual(header, tuple(OPERATOR_MONITORING_TEMPLATE_HEADERS))
        self.assertIn("ev-market", data[header.index("Evidence IDs")])
        self.assertIn("upload:file-1:metrics.csv#row=2", data[header.index("Internal source refs")])
        self.assertNotIn(r"C:\Users\nicoc", "\n".join(data))

    def test_monitoring_template_cell_content_is_deterministic(self):
        state = make_export_state("deterministic-monitor")

        first = _xlsx_rows(export_project_profile_bytes(state, "client_monitoring_template", "xlsx")[0])
        second = _xlsx_rows(export_project_profile_bytes(state, "client_monitoring_template", "xlsx")[0])

        self.assertEqual(first, second)
        self.assertEqual(first, monitoring_template_cell_rows(state, audience="client"))

    def test_client_quality_gate_cleans_raw_ids_sources_and_unsupported_certainty(self):
        state = ProjectState(
            project_id="client-quality-gate",
            project_name="Client quality gate",
            brief="Evaluate a growth pilot with partial evidence.",
            report=r"""# Executive Summary
Direct project evidence: Moderate — three supplied documents (GTM plan, market research, web proposal, social calendar)
BF=12.0 says the confirmed causal hypothesis is retention. RPN 90 remains high.

# The Decision
Decide whether to scale.

# Recommended Path
Proceed only after validation.

# Why This Is Recommended
knowledge_alpha suggests the pilot is promising.
evidence_alpha and source_ref=upload:file-9:source.md#chunk=7 mechanically explained the slowdown.
source_ref=upload:file-1:metrics.md#chunk=2 provides evidence interpretation context.
target threshold <provisional threshold.
secret=supersecret C:\Users\operator\project.txt

# Evidence Used
| Evidence | What it suggests | Citation |
|---|---|---|
| knowledge_table | knowledge_cell confirms traction | No citation available |
| evidence_table | mechanically explained demand | citation unavailable |

# Monitoring and Kill Criteria
Stop if knowledge_monitor deteriorates.
""",
        )

        markdown = build_client_dossier_markdown(state)

        for forbidden in (
            "knowledge_alpha",
            "knowledge_table",
            "knowledge_cell",
            "knowledge_monitor",
            "evidence_alpha",
            "evidence_table",
            "source_ref=",
            "upload:file",
            r"C:\Users",
            "supersecret",
            "BF=12.0",
            "RPN 90",
            "owner_decision_authority",
            "variable_coverage",
            "provisional threshold",
            "operator-defined",
            "No citation available",
            "Evidence source unavailable",
            "Citation |",
            "three supplied documents",
            "confirmed causal hypothesis",
            "confirms traction",
            "mechanically explained",
            "provides evidence interpretation context",
        ):
            self.assertNotIn(forbidden, markdown)
        self.assertEqual(markdown.count("No concrete source locators were available for this project"), 1)
        self.assertIn("project evidence", markdown)
        self.assertIn("working diagnosis", markdown)
        self.assertIn("Evidence maturity: Hypothesis-only", markdown)

    def test_client_metadata_and_report_omit_internal_ids_and_operator_notes(self):
        state = make_sparse_growth_state("d07ade65-f9d3-462d-b7c0-23bf6c505a6e")
        state.project_name = "Mini Test — Activation Bottleneck Decision"
        state.risk_classification = "minimal_risk"
        state.report = (
            "# Executive Summary\n"
            "Recommend dashboard telemetry, product analytics, legal review, circuit breaker, and regeneration-event logging.\n"
        )

        client_markdown = build_client_dossier_markdown(state)
        report_markdown = _safe_report_markdown(state)

        for markdown in (client_markdown, report_markdown):
            self.assertNotIn("Telemetry privacy note", markdown)
            self.assertNotIn("Risk classification may understate", markdown)
            self.assertNotIn("minimal_risk", markdown)
            self.assertNotRegex(
                markdown,
                r"[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}",
            )
        self.assertIn("Project name: Mini Test", client_markdown)
        self.assertIn("Generated:", client_markdown)
        self.assertNotIn("Project ID:", client_markdown)
        self.assertNotIn("Risk:", client_markdown)

        report_payload, _, _ = export_project_profile_bytes(state, "report", "docx")
        report_text = _docx_text(report_payload)
        self.assertNotIn("Telemetry privacy note", report_text)
        self.assertNotIn("Risk classification may understate", report_text)
        self.assertNotIn("minimal_risk", report_text)
        self.assertNotRegex(
            report_text,
            r"[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}",
        )

    def test_client_exports_remove_citation_noise_but_preserve_concrete_timing_values(self):
        state = ProjectState(
            project_id="client-citation-timing",
            project_name="Client citation timing",
            brief="Decide whether to continue an activation pilot with sparse evidence.",
            report="""# Executive Summary
Use a legal-review SLA of 24 hours or less, a 72 hours escalation window, and a 48 hours support follow-up.

# Evidence Used
| Evidence | What it suggests | Citation |
|---|---|---|
| Activation note | Keep response inside 72 hours. | No citation available |
| Cadence note | Compare 7-day rolling and 14-day rolling activation by Day 7 and Day 10. | Evidence source unavailable |

# Monitoring and Kill Criteria
Stop if activation falls below provisional threshold of 7-day rolling baseline.
Resolve blockers within 14 days.
""",
        )

        client_markdown = build_client_dossier_markdown(state)
        report_markdown = _safe_report_markdown(state)

        for markdown in (client_markdown, report_markdown):
            for concrete in (
                "24 hours or less",
                "72 hours",
                "48 hours",
                "7-day rolling",
                "14-day rolling",
                "Day 7",
                "Day 10",
                "within 14 days",
            ):
                self.assertIn(concrete, markdown)
            self.assertNotIn("No citation available", markdown)
            self.assertNotIn("Evidence source unavailable", markdown)
            self.assertNotIn("Citation |", markdown)
            self.assertNotRegex(markdown, r"(?mi)^\s*Citation\s*:?\s*$")
            self.assertNotIn("planning estimate to validate in Sprint 0", markdown)
            self.assertNotIn("threshold to validate in Sprint 0", markdown)
            self.assertNotIn("operator-defined", markdown)
            self.assertNotIn("provisional threshold", markdown)
        self.assertEqual(client_markdown.count("No concrete source locators were available for this project"), 1)

    def test_operator_monitoring_template_note_is_single_and_nonduplicating(self):
        state = make_export_state("operator-monitor-note")

        first = build_operator_dossier_markdown(state)
        second = build_operator_dossier_markdown(state)
        self.assertEqual(first.count(MONITORING_TEMPLATE_OPERATOR_NOTE), 1)
        self.assertEqual(second.count(MONITORING_TEMPLATE_OPERATOR_NOTE), 1)

        state.monitor.commitment_rationale = MONITORING_TEMPLATE_OPERATOR_NOTE
        summary = operator_monitoring_summary(state)
        self.assertEqual(summary.count(MONITORING_TEMPLATE_OPERATOR_NOTE), 1)

    def test_report_profile_includes_versioned_freshness_warning_when_stale(self):
        state = make_export_state("stale-report-profile")
        _attach_report_generation_metadata(state, code_version="old123")

        with patch("report_freshness.current_code_version", return_value="new456"):
            payload, _, _ = export_project_profile_bytes(state, "report", "docx")

        text = _docx_text(payload)
        self.assertIn(
            "Freshness check: report was generated with code version old123; "
            "current code version is new456. Regenerate from the current branch before client delivery.",
            text,
        )

    def test_client_dossier_includes_concise_freshness_warning_when_stale(self):
        state = make_export_state("stale-client-profile")
        _attach_report_generation_metadata(state, code_version="old123")

        with patch("report_freshness.current_code_version", return_value="new456"):
            markdown = build_client_dossier_markdown(state)

        self.assertIn("Freshness check:", markdown)
        self.assertIn("old123", markdown)
        self.assertIn("new456", markdown)
        self.assertNotIn("Freshness metadata", markdown)

    def test_operator_dossier_includes_detailed_freshness_warning_when_stale(self):
        state = make_export_state("stale-operator-profile")
        _attach_report_generation_metadata(state, code_version="old123")

        with patch("report_freshness.current_code_version", return_value="new456"):
            markdown = build_operator_dossier_markdown(state)

        self.assertIn("Freshness check:", markdown)
        self.assertIn("Freshness metadata", markdown)
        self.assertIn("Generated code version", markdown)
        self.assertIn("old123", markdown)
        self.assertIn("Current code version", markdown)
        self.assertIn("new456", markdown)

    def test_matching_report_generation_version_suppresses_freshness_warning(self):
        state = make_export_state("fresh-report-profile")
        _attach_report_generation_metadata(state, code_version="same123")

        with patch("report_freshness.current_code_version", return_value="same123"):
            report_payload, _, _ = export_project_profile_bytes(state, "report", "docx")
            client_markdown = build_client_dossier_markdown(state)
            operator_markdown = build_operator_dossier_markdown(state)

        self.assertNotIn("Freshness check:", _docx_text(report_payload))
        self.assertNotIn("Freshness check:", client_markdown)
        self.assertNotIn("Freshness check:", operator_markdown)
        self.assertIn("Freshness metadata", operator_markdown)
        self.assertIn("fresh", operator_markdown)
        self.assertIn("Matching report hash", operator_markdown)

    def test_legacy_missing_generation_version_triggers_unproven_freshness_warning(self):
        state = make_export_state("legacy-report-profile")

        with patch("report_freshness.current_code_version", return_value="new456"):
            payload, _, _ = export_project_profile_bytes(state, "report", "docx")

        self.assertIn(
            "Freshness check: report generation metadata is missing or incomplete. "
            "Regenerate from the current branch before client delivery.",
            _docx_text(payload),
        )

    def test_matching_hash_unknown_version_uses_exact_unverified_warning(self):
        state = make_export_state("unknown-version-report")
        _attach_report_generation_metadata(state, code_version="unknown")

        with patch("report_freshness.current_code_version", return_value="4.4.0"):
            freshness = report_freshness.assess_report_freshness(state)
            client_markdown = build_client_dossier_markdown(state)

        expected = (
            "Freshness check: report hash matches the stored report, but code-version metadata is unverified. "
            "Regenerate from the current branch before client delivery."
        )
        self.assertEqual(freshness.status, "unproven")
        self.assertTrue(freshness.matching_report_hash)
        self.assertEqual(freshness.warning, expected)
        self.assertIn(expected, client_markdown)

    def test_content_mismatch_uses_content_mismatch_status_and_warning(self):
        state = make_export_state("content-mismatch-report")
        _attach_report_generation_metadata(state, code_version="4.4.0")
        state.report += "\nChanged after report metadata was recorded."

        with patch("report_freshness.current_code_version", return_value="4.4.0"):
            freshness = report_freshness.assess_report_freshness(state)
            operator_markdown = build_operator_dossier_markdown(state)

        self.assertEqual(freshness.status, "content_mismatch")
        self.assertFalse(freshness.matching_report_hash)
        self.assertIn("Freshness check: stored report-generation metadata does not match", freshness.warning)
        self.assertIn("content_mismatch", operator_markdown)
        self.assertIn("Matching report hash", operator_markdown)
        self.assertIn("False", operator_markdown)

    def test_client_dossier_includes_expected_sections_and_safe_cdp_wording(self):
        markdown = build_client_dossier_markdown(make_export_state("client-profile"))

        _assert_in_order(self, markdown, CLIENT_HEADINGS)
        for expected in (
            "Run the 30-day pilot before the full program.",
            "Day 1-30: run the audit.",
            "Stop if data access is unavailable.",
            "Who owns the pilot?",
            "This export is intended to support human review and decision-making.",
        ):
            self.assertIn(expected, markdown)
        for forbidden in (
            "policy_audit_log",
            "raw_prompt",
            "raw_provider_payload",
            "sk-test-secret",
            "chain_of_thought",
            "scenario_probability",
            "semantic proof",
            "defensibility score",
            "evidence validated",
            "claim proven",
            "owner_decision_authority",
            "variable_coverage",
            "Variable coverage summary",
            "Missing decision-critical categories",
        ):
            self.assertNotIn(forbidden, markdown)

    def test_operator_dossier_includes_summaries_and_excludes_unsafe_detail(self):
        markdown = build_operator_dossier_markdown(make_export_state("operator-profile"))

        _assert_in_order(self, markdown, OPERATOR_HEADINGS)
        for expected in (
            "Phase completion status",
            "Dashboard overview",
            "Evidence and source summary",
            "Decision trace / explainability",
            "Technical appendix",
            "policy_gate_blocked=1",
            "Market note",
            "Variable coverage summary",
            "Missing decision-critical categories",
            "Evidence needs:",
            "Variable coverage limitation",
        ):
            self.assertIn(expected, markdown)
        for forbidden in (
            "sk-test-secret",
            "raw_provider_payload",
            "raw_prompt",
            "chain_of_thought",
            "ProjectState",
            r"C:\Users\nicoc",
            "owner_decision_authority",
            "coverage-debug",
        ):
            self.assertNotIn(forbidden, markdown)

    def test_sparse_dossiers_include_explicit_empty_states(self):
        sparse = ProjectState(
            project_id="sparse-export",
            project_name="Sparse export",
            brief="Sparse decision brief",
        )
        client_markdown = build_client_dossier_markdown(sparse)
        operator_markdown = build_operator_dossier_markdown(sparse)
        combined = client_markdown + "\n\n" + operator_markdown

        _assert_in_order(self, client_markdown, CLIENT_HEADINGS)
        _assert_in_order(self, operator_markdown, OPERATOR_HEADINGS)
        for expected in (
            "No hypotheses have been generated yet.",
            "No uploaded files were attached.",
            "No monitoring plan is available yet.",
            "This section will populate after the strategy phase runs.",
            "No imported evidence is available yet.",
            "No decision trace is available yet.",
            "No clarification answers have been submitted yet.",
            "No audit findings are available yet.",
        ):
            self.assertIn(expected, combined)
        self.assertIn("This is a structured hypothesis map, not a measured audit.", combined)
        self.assertNotIn(PROVISIONAL_CLARIFICATION_CAVEAT, combined)
        self.assertNotIn(PROVISIONAL_CLARIFICATION_NEXT_ACTION, combined)
        self.assertIn("Sprint 0 evidence collection should validate", client_markdown)

    def test_client_dossier_simplifies_internal_jargon_but_operator_keeps_detail(self):
        state = make_export_state("client-jargon")
        state.report = """# Executive Summary
H1 has BF 12, DQ 65, H_norm 0.12, rho 0.45, correlation=0.44, FMEA RPN 336 [#10], citation unavailable.
H2 has Jaccard index 0.42, Brier score 0.20, ECE 0.12, probability 70%, scenario_probability: 0.91, structural probability=0.73, and failure probability 0.70.

# The Decision
Choose whether to productize the dashboard.

# Recommended Path
Run Sprint 0 before productization.

# Evidence Used
H1 BF 12, DQ 65, H_norm 0.12, rho 0.45, correlation=0.44, FMEA row citation unavailable. H2 Jaccard index 0.42, Brier score 0.20, ECE 0.12, scenario_probability: 0.91, citation unavailable.

# Key Risks
RPN 336 and FMEA gaps remain.

# Appendix: Technical Analysis
FMEA RPN 336 BF 12 DQ 65 H_norm 0.12 rho 0.45 [#24]
"""
        state.audit = AuditOutput(
            fmea=[FMEAItem(component="dashboard", failure_mode="unclear owner", rpn=336, action="assign owner")],
            top_findings=["FMEA RPN 336 needs review"],
        )

        client_markdown = build_client_dossier_markdown(state)
        operator_markdown = build_operator_dossier_markdown(state)

        for forbidden in (
            "RPN",
            "FMEA",
            "BF",
            "DQ",
            "H_norm",
            "rho",
            "correlation=0.44",
            "Jaccard",
            "Brier score",
            "ECE",
            "probability 70%",
            "scenario_probability: 0.91",
            "structural probability=0.73",
            "[#10]",
            "[#24]",
        ):
            self.assertNotIn(forbidden, client_markdown)
        for expected in (
            "risk priority",
            "structured risk review",
            "structural confidence signal",
            "evidence quality signal",
            "uncertainty signal",
            "related-hypothesis risk",
            "user-value hypothesis",
            "schema overlap score",
            "forecast accuracy check",
            "calibration check",
            "No concrete source locators were available for this project",
        ):
            self.assertIn(expected, client_markdown)
        self.assertIn("FMEA", operator_markdown)
        self.assertIn("RPN", operator_markdown)
        self.assertIn("structural BF estimate=12", operator_markdown)
        self.assertIn("operator trace, not measured posterior", operator_markdown)
        self.assertIn("DQ=65, diagnostic score", operator_markdown)
        self.assertNotIn("model-generated prior", client_markdown)
        self.assertNotIn("internal confidence diagnostic", client_markdown)
        self.assertNotIn("evidence quality diagnostic", client_markdown)
        self.assertNotIn("structured risk priority", client_markdown)

    def test_monitor_signals_fill_success_metrics_when_strategy_metrics_empty(self):
        state = make_export_state("monitor-success")
        state.strategy.success_metrics = []
        state.monitor = MonitorOutput(
            ooda_schedule=MonitorOODASchedule(
                daily=[MonitorScheduleItem(metric="Pilot activation", owner="Product Owner", source="product telemetry")]
            ),
            canaries=[MonitorCanary(signal="Regeneration failure rate", direction="down", window="7d", meaning="workflow quality")],
            circuit_breakers=[MonitorCircuitBreaker(strategy_ref="S1", trip="failure rate more than 15%", reset="two stable weeks")],
            commitment_score=70,
            commitment_rationale="Owner and cadence confirmed.",
        )

        operator_markdown = build_operator_dossier_markdown(state)
        client_markdown = build_client_dossier_markdown(state)

        self.assertIn("Success metrics are captured in the monitoring plan below.", operator_markdown)
        self.assertIn("Regeneration failure rate", operator_markdown)
        self.assertIn("Pilot activation", client_markdown)
        self.assertNotIn("No success metrics saved.", operator_markdown)

    def test_unavailable_commitment_score_renders_not_scored(self):
        state = make_export_state("commitment-unscored")
        state.monitor.commitment_score = 0
        state.monitor.commitment_rationale = ""

        combined = build_client_dossier_markdown(state) + "\n\n" + build_operator_dossier_markdown(state)

        self.assertIn("Commitment score: Not scored — requires operator confirmation.", combined)
        self.assertNotIn("Commitment score: 0", combined)

    def test_numeric_commitment_with_placeholder_rationale_renders_not_scored(self):
        state = make_export_state("commitment-placeholder")
        state.monitor.commitment_score = 71
        state.monitor.commitment_rationale = "TBD / operator confirmation"

        combined = build_client_dossier_markdown(state) + "\n\n" + build_operator_dossier_markdown(state)

        self.assertIn("Commitment score: Not scored — requires operator confirmation.", combined)
        self.assertNotIn("Commitment score: 71", combined)

    def test_substantive_commitment_rationale_preserves_numeric_score(self):
        state = make_export_state("commitment-substantive")
        state.monitor.commitment_score = 61
        state.monitor.commitment_rationale = "Executive sponsor, owner, budget review, and weekly cadence are confirmed."

        combined = build_client_dossier_markdown(state) + "\n\n" + build_operator_dossier_markdown(state)

        self.assertIn("Commitment score: 61", combined)
        self.assertIn("Executive sponsor, owner, budget review, and weekly cadence are confirmed.", combined)

    def test_threshold_conflicts_warn_and_consistent_thresholds_do_not(self):
        conflict = make_export_state("threshold-conflict")
        conflict.report = (
            "# Executive Summary\n"
            "Use rho target 0.45 for portfolio risk.\n"
            "# Monitoring and Kill Criteria\n"
            "Canary success more than 5/20 in one table and more than 15% elsewhere. rho less than 0.50.\n"
        )
        consistent = make_export_state("threshold-consistent")
        consistent.report = (
            "# Executive Summary\n"
            "Use rho target 0.45 for portfolio risk.\n"
            "# Monitoring and Kill Criteria\n"
            "Canary success more than 5/20 in each table. rho target 0.45.\n"
        )

        conflict_operator = build_operator_dossier_markdown(conflict)
        conflict_client = build_client_dossier_markdown(conflict)
        self.assertNotIn(THRESHOLD_CONFLICT_UNKNOWN_WARNING, conflict_operator)
        self.assertNotIn(THRESHOLD_CONFLICT_UNKNOWN_WARNING, conflict_client)
        self.assertNotIn("Threshold consistency warning", conflict_operator)
        self.assertNotIn("Confirm one decision matrix in Sprint 0", conflict_client)
        self.assertNotIn("Threshold consistency warning", build_operator_dossier_markdown(consistent))

    def test_exact_dollar_estimate_without_budget_evidence_is_caveated(self):
        state = ProjectState(
            project_id="dollar-without-budget",
            project_name="Dollar without budget",
            brief="Decide productization direction for a product pilot.",
            report="# Executive Summary\nExpected value is $125,000 with high confidence.",
        )

        operator_markdown = build_operator_dossier_markdown(state)

        self.assertIn("Exact dollar estimates appear without budget or spend evidence", operator_markdown)
        self.assertIn("Moderate confidence in the need for Sprint 0 evidence collection", operator_markdown)
        self.assertNotIn("High confidence only that evidence collection is required", operator_markdown)

    def test_sparse_client_dossier_replaces_exact_probabilities_but_keeps_provisional_gates(self):
        state = ProjectState(
            project_id="sparse-probabilities",
            project_name="Sparse probabilities",
            brief="Improve growth performance across sales, retention, and pipeline.",
            report="""# Executive Summary
H1 probability 70%, scenario_probability: 0.91, structural probability=0.73, and predicted failure probability 0.70.

# Monitoring and Kill Criteria
Stop threshold is more than 15% churn.
The proposed planning gate is more than 20% activation.
""",
        )

        markdown = build_client_dossier_markdown(state)

        self.assertNotIn("probability 70%", markdown)
        self.assertNotIn("scenario_probability: 0.91", markdown)
        self.assertNotIn("structural probability=0.73", markdown)
        self.assertNotIn("failure probability 0.70", markdown)
        self.assertIn("structural prior", markdown)
        self.assertNotIn("model-generated prior", markdown)
        self.assertIn("high provisional failure risk", markdown)
        self.assertIn("above the threshold to validate in Sprint 0", markdown)
        self.assertNotIn("operator-defined", markdown)
        self.assertNotIn("operator-confirmed threshold required", markdown)
        self.assertIn("proposed planning gate is more than 20% activation", markdown)

    def test_sparse_report_preserves_exact_values_in_technical_appendix(self):
        state = ProjectState(
            project_id="sparse-technical-appendix",
            project_name="Sparse technical appendix",
            brief="Decide productization direction for a product pilot.",
            report="""# Executive Summary
BF=42 DQ=70 RPN=336 correlation=0.44 probability 70%.

# Appendix: Technical Analysis
BF=42 DQ=70 RPN=336 correlation=0.44 probability 70%.
""",
        )

        markdown = _safe_report_markdown(state)
        main, appendix = markdown.split("# Appendix: Technical Analysis", 1)

        for exact in ("BF=42", "DQ=70", "RPN=336", "correlation=0.44", "probability 70%"):
            self.assertNotIn(exact, main)
        self.assertIn("structural BF estimate=42 (operator trace, not measured posterior)", appendix)
        self.assertNotIn("BF=42", appendix)
        for exact in ("DQ=70", "RPN=336", "correlation=0.44", "probability 70%"):
            self.assertIn(exact, appendix)

    def test_sparse_growth_client_dossier_has_evidence_badge_and_decision_package(self):
        state = make_sparse_growth_state("sparse-growth-client")

        markdown = build_client_dossier_markdown(state)

        self.assertIn("Evidence maturity: Hypothesis-only", markdown)
        self.assertIn("Client-use status: Internal planning only", markdown)
        self.assertIn("Validation required: Sprint 0 evidence pack", markdown)
        self.assertEqual(len(re.findall(r"(?m)^## Decision Gates$", markdown)), 1)
        for row in (
            "Data quality",
            "Measurement artifact",
            "Retention",
            "PMF",
            "Channel concentration",
            "Strategic action",
            "Governance",
        ):
            self.assertIn(row, markdown)
        self.assertNotIn("Threshold consistency warning", markdown)
        self.assertNotIn("Confirm one decision matrix", markdown)

    def test_sparse_growth_report_profile_has_single_decision_gates_section(self):
        state = make_sparse_growth_state("sparse-growth-report")

        markdown = _safe_report_markdown(state)

        self.assertEqual(len(re.findall(r"(?m)^## Decision Gates$", markdown)), 1)
        self.assertNotIn(PROVISIONAL_CLARIFICATION_CAVEAT, markdown)
        self.assertNotIn(PROVISIONAL_CLARIFICATION_NEXT_ACTION, markdown)
        self.assertIn("Governance fallback if leadership overrides the diagnostic hold", markdown)
        self.assertIn("What the team may do during Sprint 0", markdown)
        self.assertIn("Minimum staffing assumption", markdown)

    def test_sparse_growth_no_generated_questions_omits_provisional_warning_everywhere(self):
        state = make_sparse_growth_state("sparse-growth-no-clarification-warning")
        state.clarification_cycles = []
        state.clarification_answers = []
        _attach_report_generation_metadata(state, code_version="4.4.0")

        client_markdown = build_client_dossier_markdown(state, current_code_version="4.4.0")
        operator_markdown = build_operator_dossier_markdown(state, current_code_version="4.4.0")
        report_markdown = _safe_report_markdown(state)

        for markdown in (client_markdown, operator_markdown, report_markdown):
            self.assertNotIn(PROVISIONAL_CLARIFICATION_CAVEAT, markdown)
            self.assertNotIn(PROVISIONAL_CLARIFICATION_NEXT_ACTION, markdown)
        self.assertNotIn("Freshness check:", client_markdown)

    def test_open_required_clarifications_show_provisional_operator_action(self):
        state = make_sparse_growth_state("sparse-growth-open-clarification")
        state.clarification_cycles = [
            ClarificationCycle(
                cycle_id="cycle-critical",
                project_id=state.project_id,
                questions=[
                    ClarificationQuestion(
                        question_id="critical",
                        text="Who owns the spend gate?",
                        why_it_matters="Ownership makes the gate enforceable.",
                        priority=ClarificationPriority.CRITICAL,
                        affected_phase="strategy",
                        source_gap="owner",
                        status=ClarificationStatus.OPEN,
                    )
                ],
            )
        ]
        _attach_report_generation_metadata(state, code_version="4.4.0")

        markdown = build_client_dossier_markdown(state)
        operator_markdown = build_operator_dossier_markdown(state)
        report_markdown = _safe_report_markdown(state)

        self.assertIn(PROVISIONAL_CLARIFICATION_CAVEAT, markdown)
        self.assertIn(PROVISIONAL_CLARIFICATION_NEXT_ACTION, markdown)
        self.assertIn(PROVISIONAL_CLARIFICATION_CAVEAT, operator_markdown)
        self.assertIn(PROVISIONAL_CLARIFICATION_NEXT_ACTION, operator_markdown)
        self.assertIn(PROVISIONAL_CLARIFICATION_CAVEAT, report_markdown)
        self.assertIn(PROVISIONAL_CLARIFICATION_NEXT_ACTION, report_markdown)
        self.assertNotIn("Freshness check:", markdown)

    def test_resolved_required_clarifications_hide_provisional_warning(self):
        state = make_sparse_growth_state("sparse-growth-resolved-clarification")
        state.clarification_cycles = [
            ClarificationCycle(
                cycle_id="cycle-critical",
                project_id=state.project_id,
                questions=[
                    ClarificationQuestion(
                        question_id="critical",
                        text="Who owns the spend gate?",
                        why_it_matters="Ownership makes the gate enforceable.",
                        priority=ClarificationPriority.CRITICAL,
                        affected_phase="strategy",
                        source_gap="owner",
                        status=ClarificationStatus.ANSWERED,
                    )
                ],
            )
        ]
        state.clarification_answers = [
            ClarificationAnswer(
                answer_id="answer-critical",
                question_id="critical",
                answer_text="Growth Lead",
                status=ClarificationStatus.ANSWERED,
            )
        ]
        _attach_report_generation_metadata(state, code_version="4.4.0")

        markdown = build_client_dossier_markdown(state)
        report_markdown = _safe_report_markdown(state)

        self.assertNotIn(PROVISIONAL_CLARIFICATION_CAVEAT, markdown)
        self.assertNotIn(PROVISIONAL_CLARIFICATION_NEXT_ACTION, markdown)
        self.assertNotIn(PROVISIONAL_CLARIFICATION_CAVEAT, report_markdown)
        self.assertNotIn(PROVISIONAL_CLARIFICATION_NEXT_ACTION, report_markdown)

    def test_sparse_growth_governance_sprint0_and_capacity_sections_render(self):
        markdown = build_client_dossier_markdown(make_sparse_growth_state("sparse-growth-package"))

        self.assertIn("Governance fallback if leadership overrides the diagnostic hold", markdown)
        self.assertIn("capped canary budget", markdown)
        self.assertIn("one explicit hypothesis", markdown)
        self.assertIn("one success metric, one stop metric, and one review date", markdown)
        self.assertIn("Block permanent headcount, major acquisition spend", markdown)
        self.assertIn("What the team may do during Sprint 0", markdown)
        self.assertIn("repair tracking", markdown)
        self.assertIn("full strategy pivot", markdown)
        self.assertIn("Minimum staffing assumption", markdown)
        self.assertIn("billing reconciliation", markdown)
        self.assertIn("cohort retention", markdown)
        self.assertIn("funnel conversion", markdown)
        self.assertIn("10 churn/user interviews", markdown)

    def test_sparse_growth_spend_gate_owner_and_enforcement_section_renders(self):
        state = make_sparse_growth_state("sparse-growth-spend-gate")

        client_markdown = build_client_dossier_markdown(state)
        operator_markdown = build_operator_dossier_markdown(state)

        for markdown in (client_markdown, operator_markdown):
            self.assertIn("Spend Gate Owner and Enforcement", markdown)
            self.assertIn(
                "Before Sprint 0 begins, assign one named owner for the Spend Authorization Gate.",
                markdown,
            )
            self.assertIn("advisory rather than enforceable", markdown)
            self.assertIn("Gate owner name / role", markdown)
            self.assertIn("Spend categories covered", markdown)
            self.assertIn("Diagnostic-spend exemptions", markdown)
            self.assertIn("DQ and BF thresholds", markdown)
            self.assertIn("Override process", markdown)
            self.assertIn("Review date", markdown)
            self.assertIn("Default owner: Executive Sponsor or Growth Lead.", markdown)

    def test_sparse_growth_client_bf_guard_and_limitation_section_render(self):
        state = make_sparse_growth_state("sparse-growth-bf-guard")
        state.classify = ClassifyOutput(bf=12.0)
        state.report += "\n# Additional Analysis\nBF=12.0 says a confirmed causal hypothesis is retention."

        markdown = build_client_dossier_markdown(state)
        operator_markdown = build_operator_dossier_markdown(state)

        self.assertIn(CLIENT_BF_CONFIDENCE_CAVEAT, markdown)
        self.assertNotIn("confirmed causal hypothesis", markdown.lower())
        self.assertNotIn("BF=12.0", markdown)
        self.assertIn("BF >10", markdown)
        self.assertIn("structural BF estimate=12.0", operator_markdown)
        self.assertIn("operator trace, not measured posterior", operator_markdown)
        self.assertIn("Main limitation of this recommendation", markdown)
        self.assertIn("This recommendation is strongest as a diagnostic plan, not as a growth strategy.", markdown)
        self.assertIn("If reliable existing data already shows one bottleneck", markdown)
        self.assertIn("If leadership will not honor the spend gate", markdown)
        self.assertIn("If measurement cannot be repaired within 30 days", markdown)

    def test_sparse_growth_main_limitation_precedes_appendix_in_report(self):
        state = make_sparse_growth_state("sparse-growth-limitation-before-appendix")
        state.report += "\n# Appendix: Technical Analysis\nTrace details."

        markdown = _safe_report_markdown(state)

        self.assertLess(
            markdown.find("Main limitation of this recommendation"),
            markdown.find("Appendix: Technical Analysis"),
        )

    def test_sparse_growth_threshold_warning_suppressed_when_decision_gates_are_canonical(self):
        state = make_sparse_growth_state("sparse-growth-canonical-thresholds")

        client_markdown = build_client_dossier_markdown(state)
        operator_markdown = build_operator_dossier_markdown(state)

        for markdown in (client_markdown, operator_markdown):
            self.assertNotIn("Threshold consistency warning", markdown)
            self.assertNotIn("Confirm one decision matrix", markdown)
            self.assertNotIn(THRESHOLD_CONFLICT_UNKNOWN_WARNING, markdown)

    def test_operator_threshold_debug_is_safe_and_client_hidden(self):
        state = make_export_state("threshold-debug-safe")
        state.report = """# Decision Gates
Proceed if DQ >70.

# Technical Appendix
Stop/change-course threshold is churn >15%.

# Convergence Gate Status
Status remains pending until DQ >70.

# Framework References
Canary example: CAC >20%.
"""

        client_markdown = build_client_dossier_markdown(state)
        operator_markdown = build_operator_dossier_markdown(state)

        self.assertNotIn("Threshold section classification", client_markdown)
        self.assertIn("Threshold section classification", operator_markdown)
        self.assertIn("Decision Gates", operator_markdown)
        self.assertIn("primary", operator_markdown)
        self.assertIn("Technical Appendix", operator_markdown)
        self.assertIn("subordinate", operator_markdown)
        self.assertIn("Convergence Gate Status", operator_markdown)
        self.assertIn("Framework References", operator_markdown)
        self.assertNotIn(THRESHOLD_CONFLICT_UNKNOWN_WARNING, operator_markdown)
        self.assertNotIn(r"C:\\Users", operator_markdown)
        self.assertNotIn("upload:file-1", operator_markdown)

    def test_threshold_conflict_warnings_are_specific_or_source_unknown(self):
        specific = make_export_state("specific-threshold-conflict")
        specific.report = """# Decision Gates
Proceed if DQ >70.

# Alternative Thresholds
Proceed if DQ >50.
"""
        ambiguous = make_export_state("ambiguous-threshold-conflict")
        ambiguous.report = "Canary success more than 5/20 in one place and more than 15% elsewhere."

        self.assertIn(
            "Threshold conflict detected between: Decision Gates and Alternative Thresholds.",
            build_operator_dossier_markdown(specific),
        )
        self.assertIn(THRESHOLD_CONFLICT_UNKNOWN_WARNING, build_operator_dossier_markdown(ambiguous))

    def test_sparse_growth_client_has_single_sprint0_evidence_section(self):
        markdown = build_client_dossier_markdown(make_sparse_growth_state("sparse-growth-evidence-dedupe"))

        self.assertEqual(len(re.findall(r"(?m)^## Sprint 0 Evidence Pack Required$", markdown)), 1)
        self.assertNotIn("Validates assumptions before implementation.", markdown)
        self.assertIn("billing/product metric reconciliation", markdown)
        self.assertIn("cohort retention curves", markdown)

    def test_sparse_growth_client_artifact_regression(self):
        markdown = build_client_dossier_markdown(make_sparse_growth_state("sparse-growth-artifacts"))

        for forbidden in (
            "operator-confirmed threshold required",
            "model-generated prior",
            "internal confidence diagnostic",
            "evidence quality diagnostic",
            "structured risk priority",
            "greater than greater than",
            "2. 1",
            "0. 68",
            "70. 0",
            "provisional threshold of the expected signal",
            "provisional threshold of the planned run time",
            "less than provisional threshold",
            "more than provisional threshold",
            "provisional effort estimate",
            "structural prior s",
            "Threshold consistency warning",
            "Confirm one decision matrix",
            THRESHOLD_CONFLICT_UNKNOWN_WARNING,
            "BF=12.0",
            "BF = 12.0",
            "BF=12",
            "ev-growth",
        ):
            self.assertNotIn(forbidden, markdown)
        for expected in (
            "Evidence maturity: Hypothesis-only",
            "Client-use status: Internal planning only",
            "Validation required: Sprint 0 evidence pack",
            "Decision Gates",
            "Monitoring Details",
            "Governance fallback if leadership overrides the diagnostic hold",
            "Spend Gate Owner and Enforcement",
            "What the team may do during Sprint 0",
            "Minimum staffing assumption",
            "Main limitation of this recommendation",
            CLIENT_BF_CONFIDENCE_CAVEAT,
        ):
            self.assertIn(expected, markdown)
        self.assertIn("structural prior", markdown)
        self.assertIn("risk priority score", markdown)
        self.assertIn("1. Repair tracking.", markdown)
        self.assertIn("2. Reconcile billing and product metrics.", markdown)
        self.assertIn("3. Interview churned customers.", markdown)

    def test_evidence_backed_operator_accounting_separates_markers_from_locators(self):
        state = ProjectState(
            project_id="evidence-accounting",
            project_name="Evidence accounting",
            brief="Improve growth performance with cohort retention evidence.",
            report="# Evidence Used\nUploaded notes support the direction [Evidence: ev1 | chunk=1].",
            knowledge_layer=KnowledgeLayerState(
                items=[
                    KnowledgeItem(
                        item_id="ev1",
                        evidence_id="ev1",
                        title="Uploaded note",
                        source_ref="upload:file-1:metrics.md#chunk=1",
                        structured_payload={"locator": "chunk=1"},
                    )
                ],
                uploaded_files=[
                    UploadedFileManifest(
                        file_id="file-1",
                        filename="metrics.md",
                        parse_summary=FileParseSummary(
                            status=FileParseStatus.COMPLETED,
                            knowledge_item_count=1,
                            chunk_count=1,
                        ),
                    )
                ],
            ),
        )

        markdown = build_operator_dossier_markdown(state)

        self.assertIn("Evidence maturity: Partial evidence", markdown)
        self.assertIn("citation_marker_count", markdown)
        self.assertIn("citation_markers_resolved_count", markdown)
        self.assertIn("citation_markers_resolved", markdown)
        self.assertIn("concrete_source_locator_count", markdown)
        self.assertIn("concrete_source_locators_available", markdown)
        self.assertIn("uploaded_file_count", markdown)
        self.assertIn("parsed_file_count", markdown)
        self.assertIn("rejected_or_unsupported_file_count", markdown)
        self.assertIn("imported_evidence_count", markdown)
        self.assertIn("imported_signal_count", markdown)
        self.assertIn("| concrete_source_locator_count | 0 |", markdown)
        self.assertIn("| concrete_source_locators_available | False |", markdown)
        self.assertIn("Uploaded knowledge chunks available; imported evidence records unavailable.", markdown)
        self.assertIn("metrics.md", markdown)
        self.assertNotIn("upload:file-1:metrics.md", markdown)

    def test_evidence_backed_missing_warning_requires_explicit_filename(self):
        explicit = ProjectState(
            project_id="explicit-missing-evidence",
            project_name="Explicit missing evidence",
            brief="Expected evidence files: 02_cohort_retention_snapshot.json and metrics.md.",
            knowledge_layer=KnowledgeLayerState(
                uploaded_files=[
                    UploadedFileManifest(
                        file_id="file-1",
                        filename="metrics.md",
                        parse_summary=FileParseSummary(status=FileParseStatus.COMPLETED),
                    )
                ]
            ),
        )
        generic = ProjectState(
            project_id="generic-missing-evidence",
            project_name="Generic missing evidence",
            brief="Improve growth with cohort retention and channel performance evidence.",
            knowledge_layer=KnowledgeLayerState(
                uploaded_files=[
                    UploadedFileManifest(
                        file_id="file-1",
                        filename="metrics.md",
                        parse_summary=FileParseSummary(status=FileParseStatus.COMPLETED),
                    )
                ]
            ),
        )

        explicit_markdown = build_operator_dossier_markdown(explicit)
        generic_markdown = build_operator_dossier_markdown(generic)

        self.assertIn(
            "Some uploaded evidence files were not ingested. Convert unsupported files to .md/.txt or enable JSON ingestion.",
            explicit_markdown,
        )
        self.assertIn("02_cohort_retention_snapshot.json", explicit_markdown)
        self.assertNotIn("Some uploaded evidence files were not ingested", generic_markdown)
        self.assertIn(
            "Some expected evidence categories may be incomplete; verify uploaded source coverage.",
            generic_markdown,
        )

    def test_final_client_markdown_cleanup_handles_bullets_pricing_and_partial_confirmed(self):
        state = ProjectState(
            project_id="final-client-cleanup",
            project_name="Final client cleanup",
            brief="Improve growth performance with cohort retention evidence.",
            data="Pricing notes: Starter tier is $499/month.",
            report="""# Executive Summary
Direct project evidence: Partial — 7 evidence files supplied.
Onboarding friction is confirmed by the supplied files.
Support ticket themes confirm that setup and integration friction is preventing value delivery.
Pricing and packaging notes confirm that the starter tier is constrained.
Support ticket volume confirms an onboarding problem.
BF = 12.0 — domain complexity confirmed.
target threshold <provisional threshold.
exceeds crux threshold by provisional threshold.
If the top channel's share rises above provisional threshold of new ARR, pause acquisition spend.
The source note says "confirmed by customer interview" and should remain quoted.
The source note says "confirm that setup is difficult" and should remain quoted.

# Why This Is Recommended
Direct project evidence: Partial — seven source files supplied; four evidence channels remain incomplete or absent.
Direct project evidence: Partial — seven evidence files supplied.
Direct project evidence: Partial — 7 source files supplied.
Onboarding friction is confirmed by the supplied files.
Support ticket themes confirm that setup and integration friction is preventing value delivery.
Pricing and packaging notes confirm that the starter tier is constrained.
Support ticket volume confirms an onboarding problem.
[Evidence: knowledge_x | chunk=1] suggests the growth metrics snapshot shows meaningful deceleration.
[Evidence: knowledge_y | chunk=2] provides evidence interpretation context indicating the severity of the trend.
The analytics audit suggests instrumentation is degraded [Evidence: knowledge_z | chunk=3].
BF = 12.0 — domain complexity confirmed.
target threshold <provisional threshold.
exceeds crux threshold by provisional threshold.
If the top channel's share rises above provisional threshold of new ARR, pause acquisition spend.
target threshold: >provisional threshold.
target threshold: <provisional threshold.
crosses provisional threshold threshold.
provisional planning estimateK/mo estimate.
If Sprint 0 data confirms that activation exceeds the gate, proceed.
The source note says "confirm that setup is difficult" and should remain quoted.

# Recommended Path
-
Repair measurement
*
Review onboarding
1 billing reconciliation
2 cohort retention
3 funnel conversion
4 10 churn/user interviews
Starter tier at provisional planning estimate.

# Evidence Used
| Evidence | What it says |
|---|---|
| Interview note | "confirmed by customer interview" |
""",
            knowledge_layer=KnowledgeLayerState(
                uploaded_files=[
                    UploadedFileManifest(
                        file_id="file-1",
                        filename="pricing.md",
                        parse_summary=FileParseSummary(status=FileParseStatus.COMPLETED),
                    ),
                    UploadedFileManifest(
                        file_id="file-2",
                        filename="interviews.md",
                        parse_summary=FileParseSummary(status=FileParseStatus.COMPLETED),
                    ),
                ]
            ),
        )

        markdown = build_client_dossier_markdown(state)

        self.assertNotRegex(markdown, r"(?m)^\s*[-*•]\s*$")
        self.assertIn("- Repair measurement", markdown)
        self.assertIn("* Review onboarding", markdown)
        self.assertIn("1. billing reconciliation", markdown)
        self.assertIn("2. cohort retention", markdown)
        self.assertIn("3. funnel conversion", markdown)
        self.assertIn("4. 10 churn/user interviews", markdown)
        self.assertIn("Starter tier at $499/month.", markdown)
        self.assertNotIn("Starter tier at provisional planning estimate", markdown)
        self.assertIn("Onboarding friction is supported by multiple supplied evidence files", markdown)
        self.assertIn(
            "Direct project evidence: Partial — supplied evidence exists, but several decision-critical evidence channels remain incomplete or unavailable.",
            markdown,
        )
        self.assertNotIn("7 evidence files supplied", markdown)
        self.assertNotIn("seven source files supplied", markdown)
        self.assertNotIn("seven evidence files supplied", markdown)
        self.assertNotIn("7 source files supplied", markdown)
        self.assertIn(
            "Support ticket themes indicate that setup and integration friction is preventing value delivery.",
            markdown,
        )
        self.assertIn("Pricing and packaging notes indicate that the starter tier is constrained.", markdown)
        self.assertIn("The analytics audit suggests instrumentation is degraded.", markdown)
        self.assertNotIn("suggests the growth metrics snapshot", markdown)
        self.assertNotIn("provides evidence interpretation context", markdown)
        self.assertNotIn("[Evidence: knowledge_x", markdown)
        self.assertNotIn("[Evidence: knowledge_y", markdown)
        self.assertNotIn("[Evidence: knowledge_z", markdown)
        self.assertIn("Support ticket volume supports an onboarding problem.", markdown)
        self.assertIn("structural BF estimate=12.0 (operator trace, not measured posterior)", markdown)
        self.assertIn("target threshold below the threshold to validate in Sprint 0.", markdown)
        self.assertIn("exceeds the crux threshold by the margin to validate in Sprint 0.", markdown)
        self.assertIn("rises above the new ARR share threshold to validate in Sprint 0", markdown)
        self.assertIn("target threshold: above the threshold to validate in Sprint 0.", markdown)
        self.assertIn("target threshold: below the threshold to validate in Sprint 0.", markdown)
        self.assertIn("crosses the threshold to validate in Sprint 0.", markdown)
        self.assertIn("If Sprint 0 data confirms that activation exceeds the gate, proceed.", markdown)
        self.assertNotIn("domain complexity confirmed", markdown)
        self.assertNotIn("provisional threshold", markdown)
        self.assertNotIn("provisional threshold threshold", markdown)
        self.assertNotIn("provisional planning estimateK", markdown)
        self.assertNotIn("operator-defined", markdown)
        self.assertIn('"confirmed by customer interview"', markdown)
        self.assertIn('"confirm that setup is difficult"', markdown)

    def test_client_markdown_cleanup_handles_moderate_document_count_and_narrative_tables(self):
        state = ProjectState(
            project_id="moderate-client-count-cleanup",
            project_name="Moderate client count cleanup",
            brief="Improve growth performance with supplied planning documents.",
            report="""# Executive Summary
Direct project evidence: Moderate — three supplied documents (GTM plan, market research, web proposal, social calendar)

# Why This Is Recommended
| Area | Client-facing summary |
|---|---|
| Evidence | Direct project evidence: Moderate — four supplied files (GTM plan, market research, web proposal, social calendar) |
| Legal | Require a legal-review SLA of operator-defined effort estimate or less before campaign claims ship. |
""",
            knowledge_layer=KnowledgeLayerState(
                uploaded_files=[
                    UploadedFileManifest(
                        file_id="file-1",
                        filename="gtm-plan.md",
                        parse_summary=FileParseSummary(status=FileParseStatus.COMPLETED),
                    ),
                    UploadedFileManifest(
                        file_id="file-2",
                        filename="market-research.md",
                        parse_summary=FileParseSummary(status=FileParseStatus.COMPLETED),
                    ),
                    UploadedFileManifest(
                        file_id="file-3",
                        filename="web-proposal.md",
                        parse_summary=FileParseSummary(status=FileParseStatus.COMPLETED),
                    ),
                    UploadedFileManifest(
                        file_id="file-4",
                        filename="social-calendar.md",
                        parse_summary=FileParseSummary(status=FileParseStatus.COMPLETED),
                    ),
                ]
            ),
        )

        markdown = build_client_dossier_markdown(state)

        expected = (
            "Direct project evidence: Moderate — supplied project documents provide "
            "planning-level evidence, but validation gaps remain."
        )
        self.assertIn(expected, markdown)
        self.assertNotIn("three supplied documents", markdown)
        self.assertNotIn("four supplied files", markdown)
        self.assertIn("| Evidence | " + expected + " |", markdown)
        self.assertIn("| Legal | Require a legal-review SLA of 24 hours or less before campaign claims ship. |", markdown)
        self.assertNotIn("operator-defined effort estimate or less", markdown)

    def test_client_count_cleanup_cleans_client_visible_source_excerpt_tables(self):
        state = ProjectState(
            project_id="protected-source-table",
            project_name="Protected source table",
            brief="Improve growth performance with supplied planning documents.",
            report="""# Executive Summary
Use the supplied evidence for planning.

# Why This Is Recommended
| Evidence ID | Source excerpt | Locator |
|---|---|---|
| ev-growth | Direct project evidence: Moderate — three supplied documents (GTM plan, market research, web proposal, social calendar) | upload:file-1:gtm-plan.pdf#page=2 |
""",
            knowledge_layer=KnowledgeLayerState(
                uploaded_files=[
                    UploadedFileManifest(
                        file_id="file-1",
                        filename="gtm-plan.md",
                        parse_summary=FileParseSummary(status=FileParseStatus.COMPLETED),
                    ),
                ]
            ),
        )

        markdown = build_client_dossier_markdown(state)

        self.assertIn(
            "Direct project evidence: Moderate — supplied project documents provide planning-level evidence, but validation gaps remain.",
            markdown,
        )
        self.assertNotIn("three supplied documents", markdown)
        self.assertNotIn("upload:file-1:gtm-plan.pdf#page=2", markdown)

    def test_risk_classification_warning_only_for_minimal_risk_with_structured_high_risk(self):
        minimal = make_sparse_growth_state("risk-minimal")
        limited = make_sparse_growth_state("risk-limited")
        minimal.audit = AuditOutput(
            fmea=[
                FMEAItem(
                    component="Growth spend",
                    failure_mode="Scaling spend before measurement repair",
                    rpn=240,
                    action="Hold scale-up until Sprint 0 evidence exists.",
                )
            ]
        )
        limited.risk_classification = "limited_risk"
        limited.audit = AuditOutput(
            fmea=[
                FMEAItem(
                    component="Growth spend",
                    failure_mode="Scaling spend before measurement repair",
                    rpn=240,
                    action="Hold scale-up until Sprint 0 evidence exists.",
                )
            ]
        )

        self.assertNotIn(RISK_CLASSIFICATION_WARNING, build_client_dossier_markdown(minimal))
        operator_markdown = build_operator_dossier_markdown(minimal)
        self.assertIn(RISK_CLASSIFICATION_WARNING, operator_markdown)
        self.assertIn("Risk classification note", operator_markdown)
        self.assertIn("High/critical structured risk diagnostics", operator_markdown)
        self.assertIn("audit.fmea", operator_markdown)
        self.assertIn("Growth spend", operator_markdown)
        self.assertIn("critical", operator_markdown)
        self.assertNotIn(RISK_CLASSIFICATION_WARNING, build_client_dossier_markdown(limited))
        self.assertNotIn(RISK_CLASSIFICATION_WARNING, build_operator_dossier_markdown(limited))

    def test_client_dossier_omits_empty_citation_marker_column_without_locators(self):
        state = ProjectState(
            project_id="empty-citation-column",
            project_name="Empty citation column",
            brief="Decide productization direction for a product pilot.",
            report="""# Executive Summary
Run Sprint 0 first.

# Evidence Used
| Evidence | What it suggests | Citation Marker |
|---|---|---|
| Supplied context | Placeholder only | citation unavailable |
| Open questions | Needs Sprint 0 | - |
""",
        )

        client_markdown = build_client_dossier_markdown(state)
        operator_markdown = build_operator_dossier_markdown(state)

        self.assertEqual(client_markdown.count("No concrete source locators were available for this project"), 1)
        self.assertIn("| Evidence | What it suggests |", client_markdown)
        self.assertNotIn("Citation Marker", client_markdown)
        self.assertNotIn("No citation available", client_markdown)
        self.assertNotIn("citation unavailable", client_markdown)
        self.assertIn("Citation Marker", operator_markdown)
        self.assertIn("citation unavailable", operator_markdown)

    def test_growth_client_dossier_avoids_generated_web_language_without_explicit_context(self):
        state = ProjectState(
            project_id="growth-client-no-web",
            project_name="Growth performance",
            brief="Improve growth performance across revenue operations, retention, and pipeline.",
            report="""# Executive Summary
Generated text says Search Console, GA4, crawl, CMS/schema capability, SEO Lead, and Web/CMS Owner.

# Evidence Used
Generated text says editorial evidence and CMS/schema capability.
""",
        )

        markdown = build_client_dossier_markdown(state)

        self.assertIn("cohort retention", markdown)
        for forbidden in ("Search Console", "GA4", "CMS/schema capability", "SEO Lead", "Web/CMS Owner"):
            self.assertNotIn(forbidden, markdown)

    def test_productization_client_and_operator_include_wave2_matrix_without_cms_leakage(self):
        state = ProjectState(
            project_id="productization-wave2",
            project_name="Productization direction",
            brief="Choose the v4 productization direction, Wave 2 roadmap, template abstraction, ROI engine, and pilot sessions.",
            report="# Executive Summary\nProceed with Wave 0 and Wave 1 before Wave 2.",
        )

        client_markdown = build_client_dossier_markdown(state)
        operator_markdown = build_operator_dossier_markdown(state)

        for markdown in (client_markdown, operator_markdown):
            self.assertIn("Wave 2 Graduation Matrix", markdown)
            self.assertIn("Proceed to Wave 2 if", markdown)
            self.assertIn("operator-set threshold", markdown)
            self.assertNotIn("70%", markdown)
            self.assertNotIn("80%", markdown)
        self.assertNotIn("CMS/schema capability", client_markdown)

    def test_telemetry_privacy_caveat_appears_for_logging_recommendations(self):
        state = make_export_state("telemetry-export")
        state.report = "# Executive Summary\nRecommend dashboard telemetry, product analytics, session replay, and regeneration-event logging."

        client_markdown = build_client_dossier_markdown(state)
        operator_markdown = build_operator_dossier_markdown(state)

        self.assertNotIn("Telemetry privacy note", client_markdown)
        self.assertNotIn("Log event metadata by default.", client_markdown)
        self.assertNotIn("Do not log raw briefs, uploaded content, report text, provider payloads, secrets, local paths, API keys, or sensitive user text", client_markdown)
        self.assertIn("Telemetry privacy note", operator_markdown)
        self.assertIn("Log event metadata by default.", operator_markdown)
        self.assertIn("Do not log raw briefs, uploaded content, report text, provider payloads, secrets, local paths, API keys, or sensitive user text", operator_markdown)
        self.assertNotIn("raw_provider_payload", client_markdown)
        self.assertNotIn("raw_prompt", client_markdown)
        self.assertNotIn("sk-test-secret", client_markdown)
        self.assertIn("This export is intended to support human review and decision-making.", client_markdown)

    def test_profile_docx_outputs_include_headings_and_tables(self):
        state = make_export_state("docx-profile")

        client_payload, _, _ = export_project_profile_bytes(state, "client_dossier", "docx")
        operator_payload, _, _ = export_project_profile_bytes(state, "operator_dossier", "docx")
        client_text = _docx_text(client_payload)
        operator_text = _docx_text(operator_payload)

        self.assertIn("What decision we reviewed", client_text)
        self.assertIn("Human review note", client_text)
        self.assertIn("Cover / project metadata", operator_text)
        self.assertIn("Technical appendix", operator_text)
        self.assertGreaterEqual(_docx_table_count(client_payload), 1)
        self.assertGreaterEqual(_docx_table_count(operator_payload), 1)

    def test_profile_pdfs_build_for_report_and_dossiers(self):
        state = make_export_state("pdf-profile")

        for profile in ("report", "client_dossier", "operator_dossier"):
            with self.subTest(profile=profile):
                payload, media_type, filename = export_project_profile_bytes(state, profile, "pdf")
                self.assertEqual(media_type, "application/pdf")
                self.assertIn(f"{profile}-", filename)
                self.assertTrue(payload.startswith(b"%PDF"))

    def test_report_docx_normalizes_blockquoted_at_a_glance_table(self):
        state = make_export_state("blockquote-report")
        state.report = """# Executive Summary
> **At a Glance**
>
> | Field | Detail |
> |---|---|
> | Decision | Run Sprint 0 first |
> | Recommendation | Validate before implementation |

---

# Evidence Used
| Evidence | What It Suggests |
|---|---|
| Supplied context | Structural analysis only |

Stop if >2 critical assumptions remain unknown.
"""

        payload, _, _ = export_project_profile_bytes(state, "report", "docx")
        text = _docx_text(payload)

        self.assertIn("At a Glance", text)
        self.assertIn("Decision", text)
        self.assertIn("Run Sprint 0 first", text)
        self.assertNotIn("| Field | Detail |", text)
        self.assertNotIn("|---|---|", text)
        self.assertNotIn("---", text)
        self.assertIn(">2 critical assumptions", text)
        for line in text.splitlines():
            self.assertFalse(line.strip().startswith(">"), line)

    def test_report_pdf_builds_with_blockquoted_table_and_wide_table_cards(self):
        state = make_export_state("wide-table-report")
        state.report = """# Executive Summary
> **At a Glance**
>
> | Field | Detail |
> |---|---|
> | Decision | Run Sprint 0 first |
> | Recommendation | Validate before implementation |

# Options Considered
| Option | Upside | Downside | Best Use Case | Verdict |
|---|---|---|---|---|
| Sprint 0 | Validates assumptions | Adds discovery time | Evidence-light SEO work | Recommended |
"""

        payload, media_type, filename = export_project_profile_bytes(state, "report", "pdf")

        self.assertEqual(media_type, "application/pdf")
        self.assertIn("wide-table-report-report-", filename)
        self.assertTrue(payload.startswith(b"%PDF"))

    def test_machine_archive_contains_expected_sanitized_files(self):
        state = make_export_state("archive-profile")
        payload, media_type, filename = export_project_profile_bytes(state, "machine_archive", "zip")

        self.assertEqual(media_type, "application/zip")
        self.assertIn("archive-profile-machine_archive-", filename)
        with zipfile.ZipFile(BytesIO(payload)) as archive:
            names = set(archive.namelist())
            self.assertNotIn("raw_project_state.json", names)
            self.assertTrue(
                {
                    "export_manifest.json",
                    "project_state.json",
                    "report.md",
                    "phase_outputs.json",
                    "clarifications.json",
                    "evidence_locator_register.json",
                    "uploaded_file_manifest.json",
                    "policy_summary.json",
                }.issubset(names)
            )
            combined = "\n".join(archive.read(name).decode("utf-8") for name in names)
            self.assertNotIn("sk-test-secret", combined)
            self.assertNotIn("hidden reasoning", combined)
            self.assertNotIn("hidden prompt", combined)
            self.assertNotIn("provider-token", combined)
            self.assertNotIn(r"C:\Users\nicoc", combined)

            manifest = json.loads(archive.read("export_manifest.json").decode("utf-8"))
            self.assertEqual(manifest["export_schema_version"], "1.0")
            self.assertEqual(manifest["export_profile"], "machine_archive")
            self.assertEqual(manifest["export_format"], "zip")
            self.assertIn("code_version", manifest)
            self.assertIn("report_freshness", manifest)
            self.assertIn("status", manifest["report_freshness"])
            self.assertIn("is_fresh", manifest["report_freshness"])
            self.assertIn("project_state.json", manifest["included_files"])
            self.assertNotIn("client_monitoring_template", manifest["included_files"])
            self.assertNotIn("operator_monitoring_template", manifest["included_files"])

            uploads = json.loads(archive.read("uploaded_file_manifest.json").decode("utf-8"))
            self.assertEqual(uploads[0]["file_id"], "file-1")
            self.assertEqual(uploads[0]["original_filename"], "market.pdf")
            self.assertEqual(uploads[0]["content_type"], "application/pdf")
            self.assertEqual(uploads[0]["size_bytes"], 321)
            self.assertEqual(uploads[0]["storage_ref"], "[REDACTED]")

    def test_monitoring_template_profiles_do_not_mutate_machine_archive_report(self):
        state = make_export_state("archive-monitor-invariance")
        original_report = state.report
        before_payload = build_machine_archive_payload(state)
        before_report = before_payload["report.md"]
        before_zip_payload, _, _ = export_project_profile_bytes(state, "machine_archive", "zip")
        with zipfile.ZipFile(BytesIO(before_zip_payload)) as archive:
            before_zip_report = archive.read("report.md").decode("utf-8")

        build_client_dossier_markdown(state)
        build_operator_dossier_markdown(state)
        export_project_profile_bytes(state, "report", "docx")
        export_project_profile_bytes(state, "client_dossier", "docx")
        export_project_profile_bytes(state, "operator_dossier", "docx")
        export_project_profile_bytes(state, "client_monitoring_template", "xlsx")
        export_project_profile_bytes(state, "operator_monitoring_template", "xlsx")
        after_payload = build_machine_archive_payload(state)
        payload, _, _ = export_project_profile_bytes(state, "machine_archive", "zip")

        self.assertEqual(state.report, original_report)
        self.assertEqual(after_payload["report.md"], before_report)
        with zipfile.ZipFile(BytesIO(payload)) as archive:
            report_md = archive.read("report.md").decode("utf-8")
            manifest = json.loads(archive.read("export_manifest.json").decode("utf-8"))
            self.assertEqual(report_md, before_zip_report)
            self.assertNotIn("client_monitoring_template", report_md)
            self.assertNotIn("operator_monitoring_template", report_md)
            self.assertNotIn("client_monitoring_template", manifest["included_files"])
            self.assertNotIn("operator_monitoring_template", manifest["included_files"])

    def test_invalid_profile_and_format_combinations_raise(self):
        state = make_export_state("invalid-profile")

        with self.assertRaises(ValueError):
            export_project_profile_bytes(state, "unknown", "pdf")
        with self.assertRaises(ValueError):
            export_project_profile_bytes(state, "client_dossier", "zip")
        with self.assertRaises(ValueError):
            export_project_profile_bytes(state, "machine_archive", "pdf")
        with self.assertRaises(ValueError):
            export_project_profile_bytes(state, "client_monitoring_template", "pdf")
        with self.assertRaises(ValueError):
            export_project_profile_bytes(state, "report", "exe")

    def test_sanitizer_redacts_sensitive_keys_and_paths(self):
        payload = {
            "api_key": "sk-test-secret",
            "nested": {
                "raw_provider_payload": {"token": "provider-token"},
                "note": r"See C:\Users\nicoc\private\file.txt",
            },
        }

        sanitized = sanitize_for_export(payload, "machine_archive", mode="redact")

        self.assertEqual(sanitized["api_key"], "[REDACTED]")
        self.assertEqual(sanitized["nested"]["raw_provider_payload"], "[REDACTED]")
        self.assertNotIn(r"C:\Users\nicoc", sanitized["nested"]["note"])

    def test_build_machine_archive_payload_has_no_raw_project_state_file(self):
        payload = build_machine_archive_payload(make_export_state("payload-profile"))

        self.assertIn("project_state.json", payload)
        self.assertNotIn("raw_project_state.json", payload)


class TestProfileExportApi(unittest.IsolatedAsyncioTestCase):
    async def asyncSetUp(self):
        api.running.clear()

    async def asyncTearDown(self):
        api.running.clear()

    async def test_legacy_export_route_still_returns_attachments(self):
        state = make_export_state("legacy-route")
        with patch("api.store.load", new=AsyncMock(return_value=state)):
            docx_response = await api.export_project(state.project_id, "docx")
            pdf_response = await api.export_project(state.project_id, "pdf")

        self.assertEqual(docx_response.media_type, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertEqual(pdf_response.media_type, "application/pdf")
        self.assertIn("attachment;", docx_response.headers["content-disposition"])
        self.assertIn("attachment;", pdf_response.headers["content-disposition"])
        self.assertIn(".docx", docx_response.headers["content-disposition"])
        self.assertIn(".pdf", pdf_response.headers["content-disposition"])

    async def test_query_route_defaults_profile_to_report(self):
        state = make_export_state("query-default")
        state.report = "Only this report content."
        with patch("api.store.load", new=AsyncMock(return_value=state)):
            response = await api.export_project_profile(state.project_id, format="docx")

        self.assertEqual(response.media_type, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertIn("query-default-report-", response.headers["content-disposition"])
        text = _docx_text(response.body)
        self.assertIn("Only this report content.", text)
        self.assertNotIn("classify summary", text)

    async def test_query_route_rejects_missing_or_invalid_format(self):
        state = make_export_state("query-invalid")
        with patch("api.store.load", new=AsyncMock(return_value=state)):
            with self.assertRaises(HTTPException) as missing:
                await api.export_project_profile(state.project_id)
            with self.assertRaises(HTTPException) as unknown:
                await api.export_project_profile(state.project_id, profile="bogus", format="pdf")
            with self.assertRaises(HTTPException) as invalid_combo:
                await api.export_project_profile(state.project_id, profile="client_dossier", format="zip")

        self.assertEqual(missing.exception.status_code, 400)
        self.assertEqual(unknown.exception.status_code, 400)
        self.assertEqual(invalid_combo.exception.status_code, 400)


if __name__ == "__main__":
    unittest.main(verbosity=2)
