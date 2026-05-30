"""Project dossier exporters for DOCX, PDF, ZIP, and XLSX downloads."""
from __future__ import annotations

import json
import zipfile
from io import BytesIO
from datetime import datetime, timezone
from xml.sax.saxutils import escape
import re
from typing import Any

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    HRFlowable,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from report_quality import (
    EVIDENCE_CATEGORY_COVERAGE_WARNING,
    NO_CONCRETE_LOCATORS_CLIENT_NOTE,
    PARTIAL_EVIDENCE_CAVEAT,
    RISK_CLASSIFICATION_WARNING,
    SPARSE_CONFIDENCE_RULE,
    TELEMETRY_PRIVACY_CAVEAT,
    UPLOADED_KNOWLEDGE_NO_IMPORTED_EVIDENCE_NOTE,
    UNSUPPORTED_EVIDENCE_FILES_WARNING,
    WAVE2_GRADUATION_MATRIX,
    assess_risk_classification_gate,
    assess_report_quality_context,
    client_simplify_text,
    commitment_score_text,
    constraint_adherence_projection,
    evidence_accounting_projection,
    evidence_maturity_projection,
    guard_client_bf_confidence,
    monitor_has_signals,
    monitor_success_metric_lines,
    normalize_export_text,
    requires_productization_wave_matrix,
    requires_telemetry_privacy_caveat,
    suppress_client_raw_evidence_ids,
    threshold_consistency_warnings,
    threshold_section_classification,
)
import report_freshness
from hypothesis_coverage import assess_hypothesis_variable_coverage
from monitoring_templates import monitoring_template_xlsx_bytes
from state import ProjectState


def build_export_filename(state: ProjectState, ext: str) -> str:
    name = (state.project_name or "decision-dossier").strip().lower()
    safe = re.sub(r"[^a-z0-9]+", "-", name).strip("-") or "decision-dossier"
    return f"{safe}-dossier.{ext}"


def export_project_docx_bytes(state: ProjectState) -> bytes:
    document = Document()
    document.core_properties.title = f"{state.project_name or 'Decision Dossier'}"
    document.core_properties.subject = "Decision Engine dossier export"

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    for block in _build_dossier_blocks(state):
        if block["type"] == "heading":
            document.add_heading(block["text"], level=min(block["level"], 4))
        elif block["type"] == "paragraph":
            document.add_paragraph(block["text"])
        elif block["type"] == "bullets":
            for item in block["items"]:
                document.add_paragraph(item, style="List Bullet")
        elif block["type"] == "numbered":
            for item in block["items"]:
                document.add_paragraph(item, style="List Number")
        elif block["type"] == "table":
            _add_docx_table(document, block["rows"])
        elif block["type"] == "divider":
            _add_docx_divider(document)

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def export_project_pdf_bytes(state: ProjectState) -> bytes:
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=state.project_name or "Decision Dossier",
    )

    styles = getSampleStyleSheet()
    heading1 = ParagraphStyle(
        "DossierHeading1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=17,
        leading=21,
        textColor=colors.HexColor("#0A1628"),
        spaceAfter=8,
    )
    heading2 = ParagraphStyle(
        "DossierHeading2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=17,
        textColor=colors.HexColor("#1E293B"),
        spaceBefore=8,
        spaceAfter=6,
    )
    heading3 = ParagraphStyle(
        "DossierHeading3",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#1E293B"),
        spaceBefore=6,
        spaceAfter=4,
    )
    body = ParagraphStyle(
        "DossierBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#0A1628"),
        spaceAfter=4,
    )

    story = []
    for block in _build_dossier_blocks(state):
        if block["type"] == "heading":
            style = heading1 if block["level"] == 1 else heading2 if block["level"] == 2 else heading3
            story.append(Paragraph(_as_pdf_text(block["text"]), style))
        elif block["type"] == "paragraph":
            story.append(Paragraph(_as_pdf_text(block["text"]), body))
        elif block["type"] == "bullets":
            items = [
                ListItem(Paragraph(_as_pdf_text(item), body), leftIndent=10)
                for item in block["items"]
            ]
            story.append(ListFlowable(items, bulletType="bullet", leftIndent=12))
        elif block["type"] == "numbered":
            items = [
                ListItem(Paragraph(_as_pdf_text(item), body), leftIndent=10)
                for item in block["items"]
            ]
            story.append(ListFlowable(items, bulletType="1", leftIndent=12))
        elif block["type"] == "table":
            story.extend(_pdf_table_flowables(block["rows"], body, doc.width))
        elif block["type"] == "divider":
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CBD5E1"), spaceBefore=4, spaceAfter=4))
        story.append(Spacer(1, 2))

    doc.build(story)
    return buf.getvalue()


REPORT_CLARITY_HEADINGS = [
    "Executive Summary",
    "The Decision",
    "Recommended Path",
    "Why This Is Recommended",
    "Options Considered",
    "Evidence Used",
    "Key Risks",
    "Assumptions and Open Questions",
    "Roadmap",
    "Next Steps",
    "Monitoring and Kill Criteria",
    "Appendix: Technical Analysis",
]

CLIENT_DOSSIER_HEADINGS = [
    "What decision we reviewed",
    "Recommended path",
    "Why this is recommended",
    "What evidence was used",
    "What should happen next",
    "Timeline / 7-30-60-90 roadmap",
    "Key risks",
    "What to monitor",
    "Open assumptions / questions",
    "Human review note",
]

OPERATOR_DOSSIER_HEADINGS = [
    "Cover / project metadata",
    "Executive summary",
    "Current recommendation",
    "Decision snapshot",
    "Phase completion status",
    "Dashboard overview",
    "Original input",
    "Classification summary",
    "Hypotheses table",
    "Gauntlet / stress-test summary",
    "Audit findings",
    "Evidence and source summary",
    "Strategy plan",
    "SQI / quality review",
    "Monitoring plan",
    "Workspace summary",
    "Risks and open questions",
    "Decision trace / explainability",
    "Clarifications / assumptions",
    "Report appendix",
    "Technical appendix",
]

HUMAN_REVIEW_NOTE = (
    "This export is intended to support human review and decision-making. "
    "It should not replace expert judgment where legal, financial, medical, "
    "safety, or compliance stakes are involved."
)
CLIENT_DELIVERY_VALIDATION_BANNER = (
    "Validate before client delivery. This is a hypothesis-driven diagnostic memo, not a measured audit."
)

SPARSE_GROWTH_DECISION_GATES = """| Gate | Proceed | Extend | Stop / Escalate |
|---|---|---|---|
| Data quality | DQ ≥70 and billing delta <5% | DQ 50-69 with clear repair path | DQ <50 by Day 30 |
| Measurement artifact | H10 rejected with reconciled data | H10 unresolved but improving data | H10 confirmed or data remains unreliable |
| Retention | D90 >40% and NRR healthy | D90 30-40% | D90 <30% or NRR <85% |
| PMF | Ellis score ≥40% with usable sample | 25-39%, segment-specific signal | <25% or activation <30% |
| Channel concentration | No channel >60% or CAC stable | Concentration high but stable | Top channel >70% and CAC worsening |
| Strategic action | One causal hypothesis BF >10 | Multiple hypotheses BF 3-8 | No hypothesis BF >10 by Week 8 |
| Governance | Leadership accepts diagnostic hold | Limited canary only | Major spend/headcount proposed before gates |"""

SPARSE_GROWTH_GOVERNANCE_FALLBACK = """If leadership chooses to act before diagnostic gates are met:

- Limit spend to a capped canary budget.
- Require one explicit hypothesis.
- Require one success metric, one stop metric, and one review date.
- Block permanent headcount, major acquisition spend, pricing overhaul, or full strategy pivot until DQ ≥70 and at least one causal hypothesis has BF >10.
- Log the decision as an override with owner, reason, risk accepted, and revisit date."""

SPARSE_GROWTH_SPRINT0_ACTIONS = """Allowed:

- repair tracking
- reconcile billing/product metrics
- interview churned customers
- pull cohort retention curves
- map funnel and channel mix
- prepare experiment briefs
- run small no-regret lifecycle/onboarding fixes that do not require major spend

Not allowed:

- full strategy pivot
- large acquisition spend increase
- new channel scale-up
- pricing overhaul
- headcount expansion justified by unvalidated growth assumptions
- permanent roadmap shift before diagnostic gates are met"""

SPARSE_GROWTH_CAPACITY_NOTE = """Minimum staffing assumption: this plan assumes access to analytics, finance/billing, product usage data, sales/win-loss data, and customer success/churn data.

If those roles or data owners are unavailable, use a reduced Sprint 0:

1. billing reconciliation
2. cohort retention
3. funnel conversion
4. 10 churn/user interviews"""

SPARSE_GROWTH_SPEND_GATE = """Before Sprint 0 begins, assign one named owner for the Spend Authorization Gate. Until this is documented, the gate is advisory rather than enforceable.

Required fields:

- Gate owner name / role
- Spend categories covered
- Diagnostic-spend exemptions
- DQ and BF thresholds
- Override process
- Review date

Default owner: Executive Sponsor or Growth Lead. If no owner is assigned, the recommendation should be treated as incomplete."""

SPARSE_GROWTH_MAIN_LIMITATION = """This recommendation is strongest as a diagnostic plan, not as a growth strategy. Its weakest point is governance: the spend gate only works if leadership assigns a named owner, defines covered spend categories, and agrees to an override process before Sprint 0 begins.

What would change this recommendation?

- If reliable existing data already shows one bottleneck, skip broad diagnosis and act on that bottleneck.
- If leadership will not honor the spend gate, run only capped canary experiments.
- If measurement cannot be repaired within 30 days, reduce scope to billing reconciliation, cohort retention, funnel conversion, and 10 user/churn interviews."""

SPARSE_GROWTH_SPRINT0_EVIDENCE_PACK = """Collect one evidence pack before treating this as a measured growth recommendation:

- billing/product metric reconciliation
- cohort retention curves
- funnel and channel-mix review
- churned-customer interviews
- win/loss or pipeline conversion review
- CAC, LTV, NRR, and expansion baseline"""

EMPTY_HYPOTHESES = "No hypotheses have been generated yet."
EMPTY_UPLOADED_FILES = "No uploaded files were attached."
EMPTY_MONITORING = "No monitoring plan is available yet."
EMPTY_STRATEGY = "This section will populate after the strategy phase runs."
EMPTY_EVIDENCE = "No imported evidence is available yet."
EMPTY_TRACE = "No decision trace is available yet."
EMPTY_CLARIFICATIONS = "No clarification answers have been submitted yet."
EMPTY_AUDIT = "No audit findings are available yet."
MONITORING_TEMPLATE_OPERATOR_NOTE = (
    "Monitoring template exports are available via "
    "profile=client_monitoring_template&format=xlsx and "
    "profile=operator_monitoring_template&format=xlsx."
)

EXPORT_PROFILE_FORMATS = {
    "report": {"pdf", "docx"},
    "client_dossier": {"pdf", "docx"},
    "client_monitoring_template": {"xlsx"},
    "operator_dossier": {"pdf", "docx"},
    "operator_monitoring_template": {"xlsx"},
    "machine_archive": {"zip"},
}

PROFILE_MEDIA_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "zip": "application/zip",
}

SENSITIVE_KEY_SUBSTRINGS = (
    "secret",
    "api_key",
    "apikey",
    "token",
    "password",
    "credential",
    "authorization",
    "cookie",
    "session",
    "provider",
    "raw_payload",
    "payload_raw",
    "raw_response",
    "raw_request",
    "raw_llm",
    "llm_response",
    "chain_of_thought",
    "chain-of-thought",
    "scratchpad",
    "debug",
    "prompt",
    "storage_ref",
    "file_path",
    "local_path",
)

_DROP = object()
REDACTED = "[REDACTED]"


def export_project_profile_bytes(state: ProjectState, profile: str, format: str) -> tuple[bytes, str, str]:
    profile_name = (profile or "report").strip().lower()
    fmt = (format or "").strip().lower()
    if not fmt:
        raise ValueError("format is required")
    if profile_name not in EXPORT_PROFILE_FORMATS:
        raise ValueError(f"profile must be one of {sorted(EXPORT_PROFILE_FORMATS)}")
    if fmt not in PROFILE_MEDIA_TYPES:
        raise ValueError("format must be pdf, docx, or zip")
    if fmt not in EXPORT_PROFILE_FORMATS[profile_name]:
        raise ValueError(f"profile={profile_name} does not support format={fmt}")

    filename = _profile_filename(state, profile_name, fmt)
    if profile_name == "machine_archive":
        archive_payload = build_machine_archive_payload(state)
        return _zip_archive_bytes(archive_payload), PROFILE_MEDIA_TYPES[fmt], filename
    if profile_name in {"client_monitoring_template", "operator_monitoring_template"}:
        audience = "operator" if profile_name.startswith("operator") else "client"
        return monitoring_template_xlsx_bytes(state, audience=audience), PROFILE_MEDIA_TYPES[fmt], filename

    current_version = report_freshness.current_code_version()
    if profile_name == "report":
        markdown = _finalize_export_markdown(
            _prepend_report_freshness_warning(
                _safe_report_markdown(state),
                state,
                current_code_version=current_version,
            ),
            state,
            audience="client",
        )
    elif profile_name == "client_dossier":
        markdown = build_client_dossier_markdown(state, current_code_version=current_version)
    else:
        markdown = build_operator_dossier_markdown(state, current_code_version=current_version)

    client_visible_profile = profile_name in {"report", "client_dossier"}
    if client_visible_profile:
        markdown = _finalize_client_runtime_render_markdown(markdown, state)

    if fmt == "docx":
        return _export_markdown_docx_bytes(markdown, title=filename), PROFILE_MEDIA_TYPES[fmt], filename
    return _export_markdown_pdf_bytes(
        markdown,
        title=filename,
        client_visible=client_visible_profile,
    ), PROFILE_MEDIA_TYPES[fmt], filename


def build_client_dossier_markdown(
    state: ProjectState,
    *,
    current_code_version: str | None = None,
) -> str:
    quality = assess_report_quality_context(state)
    sections = _extract_report_sections(state.report or "")
    lines = [
        "# Client Dossier",
        _client_project_metadata_line(state),
    ]
    warning = _report_freshness_warning(state, current_code_version=current_code_version)
    if warning:
        lines.append(warning)
    lines.extend(_evidence_maturity_badge_markdown(state, quality, client=True))
    lines.extend(_quality_warning_blocks(state, quality, client=True))
    open_questions = client_section_from_report_or_fallback(
        state,
        sections,
        ("Assumptions and Open Questions",),
        _client_open_questions(state),
    )
    clarification_questions = _client_open_questions(state)
    if (
        clarification_questions
        and clarification_questions != "No open assumptions or questions are recorded yet."
        and clarification_questions not in open_questions
    ):
        open_questions = "\n\n".join([open_questions, clarification_questions])

    section_bodies = {
        "What decision we reviewed": client_section_from_report_or_fallback(
            state,
            sections,
            ("The Decision",),
            _client_decision_reviewed(state),
        ),
        "Recommended path": client_section_from_report_or_fallback(
            state,
            sections,
            ("Recommended Path", "Executive Summary"),
            _client_recommended_path(state),
        ),
        "Why this is recommended": client_section_from_report_or_fallback(
            state,
            sections,
            ("Why This Is Recommended",),
            _client_why_recommended(state),
        ),
        "What evidence was used": client_section_from_report_or_fallback(
            state,
            sections,
            ("Evidence Used",),
            _client_evidence_used(state),
        ),
        "What should happen next": client_section_from_report_or_fallback(
            state,
            sections,
            ("Next Steps",),
            _client_next_steps(state),
        ),
        "Timeline / 7-30-60-90 roadmap": client_section_from_report_or_fallback(
            state,
            sections,
            ("Roadmap",),
            _client_roadmap(state),
        ),
        "Key risks": client_section_from_report_or_fallback(
            state,
            sections,
            ("Key Risks",),
            _client_key_risks(state),
        ),
        "What to monitor": client_section_from_report_or_fallback(
            state,
            sections,
            ("Monitoring and Kill Criteria",),
            _client_monitoring(state),
        ),
        "Open assumptions / questions": client_section_from_report_or_fallback(
            state,
            sections,
            (),
            open_questions,
        ),
        "Human review note": HUMAN_REVIEW_NOTE,
    }
    monitor_fallback = _client_monitoring(state)
    if state.monitor and monitor_fallback and monitor_fallback not in section_bodies["What to monitor"]:
        section_bodies["What to monitor"] = "\n\n".join([section_bodies["What to monitor"], monitor_fallback])
    sprint0_pack = _client_sprint0_pack(quality)
    if (
        sprint0_pack
        and not _requires_sparse_growth_decision_package(quality)
        and sprint0_pack not in section_bodies["What evidence was used"]
    ):
        section_bodies["What evidence was used"] = "\n\n".join([section_bodies["What evidence was used"], sprint0_pack])

    for heading in CLIENT_DOSSIER_HEADINGS:
        body = section_bodies.get(heading) or "Not available in current project output."
        if heading != "Human review note":
            body = _client_safe_text(body, quality)
            if not quality.has_concrete_locators:
                body = _remove_empty_client_citation_marker_columns(body)
                body = suppress_client_raw_evidence_ids(body)
            if heading == "What to monitor" and _requires_sparse_growth_decision_package(quality):
                body = "\n\n".join(["### Monitoring Details", body])
        lines.extend([
            f"## {heading}",
            body,
        ])
        if heading == "Recommended path" and _requires_sparse_growth_decision_package(quality):
            lines.extend(_sparse_growth_decision_package_sections(client=True))
        if heading == "Timeline / 7-30-60-90 roadmap" and requires_productization_wave_matrix(state, quality):
            lines.extend([
                "## Wave 2 Graduation Matrix",
                _client_safe_text(WAVE2_GRADUATION_MATRIX.replace("## Wave 2 Graduation Matrix\n\n", ""), quality),
            ])

    markdown = normalize_export_text(
        "\n\n".join(part for part in lines if str(part).strip()),
        audience="client",
    )
    if _requires_sparse_growth_decision_package(quality):
        markdown = _dedupe_client_sparse_growth_sprint0(markdown)
        markdown = guard_client_bf_confidence(markdown, state)
    return _finalize_export_markdown(markdown, state, audience="client", quality=quality)


def build_operator_dossier_markdown(
    state: ProjectState,
    *,
    current_code_version: str | None = None,
    include_freshness: bool = True,
) -> str:
    quality = assess_report_quality_context(state)
    lines = [
        "# Operator Dossier",
    ]
    if include_freshness:
        warning = _operator_report_freshness_warning(state, current_code_version=current_code_version)
        if warning:
            lines.append(warning)
    lines.extend(_evidence_maturity_badge_markdown(state, quality, client=False))
    lines.extend(_quality_warning_blocks(state, quality, client=False))
    section_bodies = {
        "Cover / project metadata": operator_project_metadata(state),
        "Executive summary": operator_executive_summary(state),
        "Current recommendation": operator_current_recommendation(state),
        "Decision snapshot": operator_decision_snapshot(state),
        "Phase completion status": summarize_phase_outputs(state),
        "Dashboard overview": operator_dashboard_overview(state),
        "Original input": operator_original_input(state),
        "Classification summary": operator_classification_summary(state),
        "Hypotheses table": operator_hypotheses_table(state),
        "Gauntlet / stress-test summary": operator_gauntlet_summary(state),
        "Audit findings": operator_audit_summary(state),
        "Evidence and source summary": operator_evidence_summary(state),
        "Strategy plan": operator_strategy_summary(state),
        "SQI / quality review": operator_sqi_summary(state),
        "Monitoring plan": operator_monitoring_summary(state),
        "Workspace summary": operator_workspace_summary(state),
        "Risks and open questions": operator_risks_and_questions(state),
        "Decision trace / explainability": operator_trace_summary(state),
        "Clarifications / assumptions": operator_clarifications_summary(state),
        "Report appendix": operator_report_appendix(state),
        "Technical appendix": operator_technical_appendix(state),
    }
    for heading in OPERATOR_DOSSIER_HEADINGS:
        lines.extend([
            f"## {heading}",
            section_bodies.get(heading) or "No dashboard-facing summary is available yet.",
        ])
        if heading == "Strategy plan" and requires_productization_wave_matrix(state, quality):
            lines.append(WAVE2_GRADUATION_MATRIX)
        if heading == "Strategy plan" and _requires_sparse_growth_decision_package(quality):
            lines.extend(_sparse_growth_decision_package_sections(client=False))
    return _finalize_export_markdown(
        "\n\n".join(part for part in lines if str(part).strip()),
        state,
        audience="operator",
        quality=quality,
    )


def _finalize_export_markdown(
    markdown: str,
    state: ProjectState,
    *,
    audience: str,
    quality=None,
) -> str:
    mode = "operator" if str(audience or "").lower() == "operator" else "client"
    quality = quality or assess_report_quality_context(state)
    value = _apply_pricing_placeholder_cleanup(str(markdown or ""), state)
    metric_fragments: list[str] = []
    if mode == "client":
        value, metric_fragments = _protect_client_concrete_metric_values(value)
    value = normalize_export_text(value, audience=mode)
    if mode == "client":
        value = _redact_unsafe_string(value)
        value = _remove_empty_client_citation_marker_columns(value)
        value = suppress_client_raw_evidence_ids(value)
        value = _finalize_client_visible_artifacts(value, state)
        value = _normalize_client_evidence_count_language(value, state, quality)
        value = _label_client_bf_trace_language(value)
        value = _soften_unvalidated_confirmed_language(value, state, quality)
        value = normalize_export_text(value, audience=mode)
        value = _remove_empty_client_citation_marker_columns(value)
        value = suppress_client_raw_evidence_ids(value)
        value = _finalize_client_visible_artifacts(value, state)
        value = _ensure_single_client_source_locator_note(value, quality)
        value = _ensure_client_delivery_validation_banner(value)
        value = _restore_client_concrete_metric_values(value, metric_fragments)
        value = _polish_client_report_citation_rendering(value)
    return value


def _apply_pricing_placeholder_cleanup(markdown: str, state: ProjectState) -> str:
    replacement = (
        "Starter tier at $499/month"
        if _state_contains_starter_price(state)
        else "Starter tier based on the supplied pricing notes."
    )
    return re.sub(
        r"\bStarter tier at (?:provisional planning estimate|planning estimate to validate in Sprint 0)\b",
        replacement,
        str(markdown or ""),
        flags=re.I,
    )


def _finalize_client_visible_artifacts(markdown: str, state: ProjectState) -> str:
    value = _drop_client_operator_only_locator_lines(str(markdown or ""))
    value = _fix_client_generated_duplicate_words(value)
    value = _remove_empty_client_citation_marker_columns(value)
    value = _strip_client_citation_placeholder_noise(value)
    value = _remove_standalone_client_citation_rows(value)
    value = _rename_client_citation_table_headers(value)
    value = _suppress_client_internal_locator_tokens(value)
    value = _fix_client_generated_duplicate_words(value)
    value = _remove_empty_client_citation_marker_columns(value)
    value = _strip_client_citation_placeholder_noise(value)
    value = _remove_standalone_client_citation_rows(value)
    value = _rename_client_citation_table_headers(value)
    value = _replace_client_threshold_placeholders(value, state)
    value = _polish_client_report_citation_rendering(value)
    return _collapse_markdown_blank_lines(value)


def _finalize_client_runtime_render_markdown(markdown: str, state: ProjectState) -> str:
    quality = assess_report_quality_context(state)
    value = _finalize_export_markdown(markdown, state, audience="client", quality=quality)
    return _polish_client_report_citation_rendering(value)


def _polish_client_report_citation_rendering(markdown: str) -> str:
    value = str(markdown or "")
    value = re.sub(
        r"(?im)\b(Sprint 0\.)[ \t]*(#{1,6}\s+Evidence maturity\b)",
        r"\1\n\n\2",
        value,
    )
    value = re.sub(
        r"(?im)\b(Sprint 0\.)[ \t]*(#{1,6}\s+Client Dossier\b)",
        r"\1\n\n\2",
        value,
    )
    value = re.sub(r"\bCitation\s+(?:citation unavailable|no citation available)\b\.?", "", value, flags=re.I)
    value = re.sub(r"(?mi)^\s*Citation\s*$", "", value)
    value = re.sub(r"\bThreshold not yet confirmed\b", "validation threshold to confirm", value, flags=re.I)
    value = re.sub(r"\bOperator to define\b", "decision owner to confirm", value, flags=re.I)
    value = re.sub(
        r"\bthreshold probability that above\s+(\d+(?:\.\d+)?\s*pp)\s+divergence\b",
        r"validation signal that divergence above \1",
        value,
        flags=re.I,
    )
    value = re.sub(r"\bthreshold probability that above\b", "validation signal that", value, flags=re.I)
    value = re.sub(r"\bthreshold probability\b", "validation signal", value, flags=re.I)
    value = _polish_client_metric_comparator_phrasing(value)
    value = re.sub(
        r"\b(?:What It Suggests|Why It Is Needed|Why it is needed|Upside)\s+supports\s+whether\b",
        "This helps determine whether",
        value,
        flags=re.I,
    )
    value = re.sub(
        r"\b(?:What It Suggests|Why It Is Needed|Why it is needed|Upside)\s+supports\s+which\b",
        "This helps identify which",
        value,
        flags=re.I,
    )
    value = re.sub(r"\bWhat It Suggests supports\b", "This supports", value, flags=re.I)
    value = re.sub(r"\b(?:Why It Is Needed|Why it is needed|Upside)\s+supports\b", "This supports", value, flags=re.I)
    value = re.sub(r"\bsupports\s+whether\b", "helps determine whether", value, flags=re.I)
    value = re.sub(r"\bsupports\s+which\b", "helps identify which", value, flags=re.I)
    return value


def _polish_client_metric_comparator_phrasing(markdown: str) -> str:
    lines: list[str] = []
    for line in str(markdown or "").splitlines():
        if re.search(r"\bproposed planning gate\b", line, re.I):
            lines.append(line)
            continue
        value = re.sub(r"\bmore than\s+(\d+(?:\.\d+)?\s*%)", r"above \1", line, flags=re.I)
        value = re.sub(r"\bless than\s+(\d+(?:\.\d+)?\s*%)", r"below \1", value, flags=re.I)
        lines.append(value)
    return "\n".join(lines)


def _protect_client_concrete_metric_values(markdown: str) -> tuple[str, list[str]]:
    fragments: list[str] = []
    patterns = [
        r"(?<![A-Za-z0-9_])(?:[<>]=?|≥|≤)\s*\d+(?:\.\d+)?\s*(?:%|pp|h|hrs?|hours?|days?|weeks?|months?)\b",
        r"\b\d+(?:\.\d+)?\s*%",
        r"\b\d+(?:\.\d+)?\s*percentage[- ]points?\b",
        r"\b\d+(?:\.\d+)?\s*(?:h|hrs?|hours?|days?|weeks?|months?)\b",
    ]
    value = str(markdown or "")

    def repl(match: re.Match[str]) -> str:
        line_start = value.rfind("\n", 0, match.start()) + 1
        line_end = value.find("\n", match.end())
        if line_end == -1:
            line_end = len(value)
        line_context = value[line_start:line_end]
        if re.search(
            r"\b(?:probability|scenario[_ -]?probability|structural probability|prior|likelihood|chance|failure probability|predicted failure probability)\b",
            line_context,
            re.I,
        ):
            return match.group(0)
        fragments.append(match.group(0))
        return f"CLIENTMETRICVALUE{len(fragments) - 1}MARKER"

    for pattern in patterns:
        value = re.sub(pattern, repl, value, flags=re.I)
    return value, fragments


def _restore_client_concrete_metric_values(markdown: str, fragments: list[str]) -> str:
    value = str(markdown or "")
    for index, fragment in enumerate(fragments):
        value = value.replace(f"CLIENTMETRICVALUE{index}MARKER", fragment)
        value = value.replace(f"CLIENTMETRICVALUE{index}TOKEN", fragment)
    return value


def _ensure_client_delivery_validation_banner(markdown: str) -> str:
    source = str(markdown or "")
    banner_pattern = re.compile(
        r"(?:\*\*)?\s*Validate before client delivery\.\s+"
        r"This is a hypothesis-driven diagnostic memo, not a measured audit\.\s*(?:\*\*)?",
        re.I,
    )
    cleaned = banner_pattern.sub("", source)
    cleaned = _collapse_markdown_blank_lines(cleaned)
    if not cleaned:
        return CLIENT_DELIVERY_VALIDATION_BANNER
    return _collapse_markdown_blank_lines("\n\n".join([CLIENT_DELIVERY_VALIDATION_BANNER, cleaned]))


def _fix_client_generated_duplicate_words(markdown: str) -> str:
    value = str(markdown or "")
    value = re.sub(r"\bThe\s+The\s+problem\b", "The problem", value, flags=re.I)
    return value


def _drop_client_operator_only_locator_lines(markdown: str) -> str:
    lines: list[str] = []
    raw_reference = re.compile(
        r"\[Evidence:\s*|\bknowledge[_-][A-Za-z0-9_.:-]+\b|"
        r"\b(?:ev|src)[-_][A-Za-z0-9_.:-]+\b|"
        r"\bevidence[_-](?!based\b|backed\b|driven\b|quality\b|maturity\b|source\b|used\b)[A-Za-z0-9_.:-]+\b|"
        r"\bupload:[^\s|,)>\]]+|\bstorage_ref\s*[:=]|\bsource_ref\s*[:=]",
        re.I,
    )
    for line in str(markdown or "").splitlines():
        if re.search(r"\boperator[-\s]?only\b", line, re.I) and raw_reference.search(line):
            continue
        lines.append(line)
    return "\n".join(lines)


def _strip_client_citation_placeholder_noise(markdown: str) -> str:
    lines: list[str] = []
    in_code_block = False
    for line in str(markdown or "").splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            lines.append(line)
            continue
        if not stripped:
            lines.append(line)
            continue
        if _looks_like_json_line(stripped):
            lines.append(_strip_client_citation_placeholder_text(line))
            continue
        cleaned = _strip_client_citation_placeholder_text(line)
        if _is_standalone_client_citation_placeholder_line(cleaned):
            continue
        if in_code_block and cleaned.strip():
            lines.append(cleaned)
            continue
        if cleaned.strip():
            lines.append(cleaned)
    return "\n".join(lines)


def _strip_client_citation_placeholder_text(value: str) -> str:
    text = str(value or "")
    if not re.search(
        r"\b(?:Citation|No citation available|citation unavailable|Evidence source unavailable)\b",
        text,
        re.I,
    ):
        return text.rstrip()
    text = re.sub(r"\b(?:No citation available|citation unavailable|Evidence source unavailable)\b\.?", "", text, flags=re.I)
    text = re.sub(r"\bCitation\s*:\s*(?=$|[|,.;])", "", text, flags=re.I)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\s+([|,.;:])", r"\1", text)
    text = re.sub(r"\|\s+\|", "| |", text)
    return text.rstrip()


def _remove_standalone_client_citation_rows(markdown: str) -> str:
    source_lines = str(markdown or "").splitlines()
    output: list[str] = []
    index = 0
    while index < len(source_lines):
        line = source_lines[index]
        if _looks_like_markdown_table_line(line):
            block: list[str] = []
            while index < len(source_lines) and _looks_like_markdown_table_line(source_lines[index]):
                block.append(source_lines[index])
                index += 1
            cleaned = [row for row in block if not _is_client_citation_noise_table_line(row)]
            if _client_table_block_has_meaningful_rows(cleaned):
                output.extend(cleaned)
            continue
        if not _is_standalone_client_citation_placeholder_line(line):
            output.append(line)
        index += 1
    return "\n".join(output)


def _client_table_block_has_meaningful_rows(block: list[str]) -> bool:
    for row in block:
        cells = _split_markdown_table_row(row)
        if not cells:
            continue
        if _is_client_markdown_separator_cells(cells):
            continue
        if any(cell.strip() for cell in cells):
            return True
    return False


def _rename_client_citation_table_headers(markdown: str) -> str:
    source_lines = str(markdown or "").splitlines()
    output: list[str] = []
    index = 0
    while index < len(source_lines):
        line = source_lines[index]
        if not _looks_like_markdown_table_line(line):
            output.append(line)
            index += 1
            continue
        block: list[str] = []
        while index < len(source_lines) and _looks_like_markdown_table_line(source_lines[index]):
            block.append(source_lines[index])
            index += 1
        if len(block) < 2:
            output.extend(block)
            continue
        rows = [_split_markdown_table_row(row) for row in block]
        if not rows or not rows[0] or not _is_client_markdown_separator_cells(rows[1]):
            output.extend(block)
            continue
        header = [
            "Citation status" if _is_client_citation_column_header(cell) else cell
            for cell in rows[0]
        ]
        output.append(_format_markdown_table_row(header))
        output.append(_format_markdown_table_row(["---"] * len(header)))
        for row in rows[2:]:
            output.append(_format_markdown_table_row(row))
    return "\n".join(output)


def _is_client_citation_noise_table_line(value: str) -> bool:
    cells = _split_markdown_table_row(value)
    if not cells:
        return False
    nonempty = [cell for cell in cells if cell.strip()]
    if not nonempty:
        return True
    if not any(
        _is_client_citation_column_header(cell)
        or _normalize_heading(cell) in {"citation unavailable", "no citation available"}
        for cell in nonempty
    ):
        return False
    return all(
        _is_client_citation_column_header(cell)
        or _is_empty_client_citation_cell(cell)
        for cell in nonempty
    )


def _is_standalone_client_citation_placeholder_line(value: str) -> bool:
    stripped = str(value or "").strip()
    if not stripped:
        return False
    stripped = re.sub(r"^\s*(?:[-*•]|\d+\.)\s+", "", stripped)
    stripped = stripped.strip("|*_` ")
    if not stripped:
        return False
    if _is_client_citation_column_header(stripped.rstrip(":")):
        return True
    return bool(re.fullmatch(r"citation\s*:?", stripped, re.I))


def _suppress_client_internal_locator_tokens(markdown: str) -> str:
    lines: list[str] = []
    for line in str(markdown or "").splitlines():
        value = str(line or "")
        value = re.sub(
            r"\s*\[Evidence:\s*[A-Za-z0-9_.:-]+\s*(?:\|[^\]]*)?\]",
            "",
            value,
            flags=re.I,
        )
        value = re.sub(r"\s*\[#\d+\]", "", value)
        value = re.sub(r"\bknowledge[_-][A-Za-z0-9_.:-]+\b", "project evidence", value, flags=re.I)
        value = re.sub(r"\b(?:ev|src)[-_][A-Za-z0-9_.:-]+\b", "project evidence", value, flags=re.I)
        value = re.sub(
            r"\bevidence[_-](?!based\b|backed\b|driven\b|quality\b|maturity\b|source\b|used\b)[A-Za-z0-9_.:-]+\b",
            "project evidence",
            value,
            flags=re.I,
        )
        value = re.sub(r"\bupload:[^\s|,)>\]]+", "uploaded project document", value, flags=re.I)
        value = re.sub(r"\bstorage_ref\s*[:=]\s*[^\s|,)>\]]+", "", value, flags=re.I)
        value = re.sub(r"\bsource_ref\s*[:=]\s*[^\s|,)>\]]+", "", value, flags=re.I)
        value = re.sub(
            r"\bproject evidence\s+(?:provides evidence interpretation context|suggests|indicates)\b[^.?!]*(?:[.?!]|$)",
            "",
            value,
            flags=re.I,
        )
        value = re.sub(r"[ \t]{2,}", " ", value)
        value = re.sub(r"\s+([,.;:])", r"\1", value)
        lines.append(value.rstrip())
    return "\n".join(lines)


def _replace_client_threshold_placeholders(markdown: str, state: ProjectState) -> str:
    catalog = _client_threshold_catalog(state)
    lines: list[str] = []
    for line in str(markdown or "").splitlines():
        lines.append(_replace_client_threshold_placeholders_in_line(line, catalog))
    return "\n".join(lines)


def _replace_client_threshold_placeholders_in_line(line: str, catalog: list[dict[str, Any]]) -> str:
    value = str(line or "")
    if not re.search(
        r"(?:≥|>=)\s*threshold|\bthreshold\s+(?:activation(?:\s+rate)?|conversion|baseline)\b|"
        r"\bexpand\s+to\s+threshold\b|\bthreshold\s+of\s+[^.;,\n|]+|"
        r"\bfalls\s+below\s+threshold\b|\btarget\s+of\s+threshold\b|"
        r"\b(?:activation|baseline)\s+validation\s+gate\b|"
        r"\b(?:above|below)\s+the\s+threshold\b|\bmargin\s+to validate in Sprint 0\b",
        value,
        re.I,
    ):
        return value

    def candidate_for(metric: str | None = None, *, prefer_positive: bool = False) -> dict[str, Any] | None:
        preferred = tuple(filter(None, [metric]))
        return _select_client_threshold_candidate(catalog, value, preferred=preferred, prefer_positive=prefer_positive)

    def target_repl(_: re.Match[str]) -> str:
        return "target " + _client_threshold_at_least_phrase(candidate_for(prefer_positive=True))

    def comparator_repl(_: re.Match[str]) -> str:
        return _client_threshold_at_least_phrase(candidate_for(prefer_positive=True))

    def threshold_of_repl(match: re.Match[str]) -> str:
        metric = match.group(1).strip()
        return _client_threshold_of_phrase(metric, candidate_for(metric))

    def falls_below_repl(_: re.Match[str]) -> str:
        return "falls " + _client_threshold_below_phrase(candidate_for())

    def target_of_repl(_: re.Match[str]) -> str:
        return _client_threshold_target_phrase(candidate_for(prefer_positive=True) or candidate_for())

    def above_threshold_repl(_: re.Match[str]) -> str:
        return _client_threshold_above_phrase(candidate_for(prefer_positive=True) or candidate_for())

    def below_threshold_repl(_: re.Match[str]) -> str:
        return _client_threshold_below_phrase(candidate_for())

    value = re.sub(r"\btarget\s*(?:≥|>=)\s*threshold\b", target_repl, value, flags=re.I)
    value = re.sub(r"(?<![\w-])(?:≥|>=)\s*threshold\b", comparator_repl, value, flags=re.I)
    value = re.sub(r"\bthreshold\s+of\s+([^.;,\n|]+)", threshold_of_repl, value, flags=re.I)
    value = re.sub(r"\bfalls\s+below\s+threshold\b", falls_below_repl, value, flags=re.I)
    value = re.sub(r"\btarget\s+of\s+threshold\b", target_of_repl, value, flags=re.I)
    value = re.sub(
        r"\bthreshold\s+activation\s+rate\b",
        lambda _: _client_threshold_metric_phrase("activation rate", candidate_for("activation rate")),
        value,
        flags=re.I,
    )
    value = re.sub(
        r"\bthreshold\s+activation\b",
        lambda _: _client_threshold_metric_phrase("activation", candidate_for("activation")),
        value,
        flags=re.I,
    )
    value = re.sub(
        r"\bthreshold\s+conversion\b",
        lambda _: _client_threshold_metric_phrase("conversion", candidate_for("conversion")),
        value,
        flags=re.I,
    )
    value = re.sub(
        r"\bthreshold\s+baseline\b",
        lambda _: _client_threshold_metric_phrase("baseline", candidate_for("baseline")),
        value,
        flags=re.I,
    )
    value = re.sub(
        r"\bactivation\s+validation\s+gate\b",
        lambda _: _client_threshold_metric_phrase("activation", candidate_for("activation")),
        value,
        flags=re.I,
    )
    value = re.sub(
        r"\bbaseline\s+validation\s+gate\b",
        lambda _: _client_threshold_metric_phrase("baseline", candidate_for("baseline")),
        value,
        flags=re.I,
    )
    value = re.sub(
        r"\bexpand\s+to\s+threshold\b",
        lambda _: _client_threshold_expand_phrase(candidate_for(prefer_positive=True)),
        value,
        flags=re.I,
    )
    value = re.sub(r"\babove\s+the\s+threshold\b", above_threshold_repl, value, flags=re.I)
    value = re.sub(r"\bbelow\s+the\s+threshold\b", below_threshold_repl, value, flags=re.I)
    value = re.sub(
        r"\b(above|below)\s+(\d+(?:\.\d+)?\s*(?:%|pp|h|hrs?|hours?|d|days?|weeks?|months?))\s+to validate in Sprint 0\s+([^.;,\n|]+)",
        lambda match: f"{match.group(1)} {match.group(2)} {match.group(3).strip()}",
        value,
        flags=re.I,
    )
    value = re.sub(r"\bvalidated gate to validate in Sprint 0\b", "validated gate", value, flags=re.I)
    value = re.sub(r"\bmargin to validate in Sprint 0\b", "validation margin", value, flags=re.I)
    value = re.sub(r"[ \t]{2,}", " ", value)
    return re.sub(r"\s+([,.;:])", r"\1", value)


def _client_threshold_catalog(state: ProjectState) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    seen: set[str] = set()
    for text in _client_threshold_source_texts(state):
        candidate = _client_threshold_candidate(text)
        if not candidate:
            continue
        key = str(candidate["phrase"]).lower()
        if key in seen:
            continue
        seen.add(key)
        candidates.append(candidate)
    return candidates


def _client_threshold_source_texts(state: ProjectState) -> list[str]:
    parts: list[str] = []
    strategy = getattr(state, "strategy", None)
    if strategy:
        parts.extend(str(item) for item in list(getattr(strategy, "success_metrics", []) or []))
        parts.extend([
            getattr(strategy, "monitoring_plan", ""),
            getattr(strategy, "implementation_sequence", ""),
            getattr(strategy, "reentry_check", ""),
        ])
        for verdict in list(getattr(strategy, "preliminary_verdicts", []) or []):
            parts.extend([getattr(verdict, "evidence", ""), getattr(verdict, "monitoring_plan", "")])
        for action in list(getattr(strategy, "strategies", []) or []):
            parts.extend([
                getattr(action, "action", ""),
                getattr(action, "justification", ""),
                getattr(action, "expected_impact", ""),
                getattr(action, "risk_if_ignored", ""),
            ])
    monitor = getattr(state, "monitor", None)
    if monitor:
        for breaker in list(getattr(monitor, "circuit_breakers", []) or []):
            parts.extend([getattr(breaker, "strategy_ref", ""), getattr(breaker, "trip", ""), getattr(breaker, "reset", "")])
        for canary in list(getattr(monitor, "canaries", []) or []):
            parts.extend([getattr(canary, "signal", ""), getattr(canary, "direction", ""), getattr(canary, "window", ""), getattr(canary, "meaning", "")])
        schedule = getattr(monitor, "ooda_schedule", None)
        if schedule:
            for cadence in ("daily", "weekly", "monthly"):
                for item in list(getattr(schedule, cadence, []) or []):
                    parts.extend([getattr(item, "metric", ""), getattr(item, "source", "")])
        parts.extend(str(item) for item in list(getattr(monitor, "reentry_watch", []) or []))
    for hypothesis in list(getattr(state, "hypotheses", []) or []):
        parts.extend([
            getattr(hypothesis, "signal", ""),
            getattr(hypothesis, "confirm", ""),
            getattr(hypothesis, "reject", ""),
        ])
    parts.extend(str(line) for line in str(getattr(state, "report", "") or "").splitlines())
    return [part for part in parts if str(part or "").strip()]


def _client_threshold_candidate(text: str) -> dict[str, Any] | None:
    phrase = _redact_unsafe_string(_strip_inline_markdown(str(text or ""))).strip()
    if not phrase or not re.search(r"\d", phrase):
        return None
    if re.search(
        r"(?:≥|>=)\s*threshold|\bthreshold\s+(?:activation(?:\s+rate)?|conversion|baseline)\b|"
        r"\bexpand\s+to\s+threshold\b|\bthreshold\s+of\s+[^.;,\n|]+|"
        r"\bfalls\s+below\s+threshold\b|\btarget\s+of\s+threshold\b|"
        r"\b(?:activation|baseline)\s+validation\s+gate\b",
        phrase,
        re.I,
    ):
        return None
    value, direction = _extract_client_threshold_value(phrase)
    if not value:
        return None
    return {
        "phrase": phrase.rstrip(". "),
        "value": value,
        "direction": direction,
        "keywords": _client_threshold_keywords(phrase),
    }


def _extract_client_threshold_value(text: str) -> tuple[str, str]:
    unit = r"(?:%|pp|h|hrs?|hours?|d|days?|weeks?|months?)"
    patterns = [
        (r"\b(\d+(?:\.\d+)?\s*%\s+of\s+[^.;,\n|]+)", "positive", ""),
        (rf"(?:≥|>=|>|at\s+least|greater\s+than|more\s+than|above|reaches?|exceeds?)\s*(\d+(?:\.\d+)?\s*{unit})", "positive", "at least"),
        (rf"(?:≤|<=|<|below|under|less\s+than|falls?\s+below)\s*(\d+(?:\.\d+)?\s*{unit})", "negative", "below"),
        (rf"\bdown\s+(\d+(?:\.\d+)?\s*{unit})", "negative", "down"),
    ]
    for pattern, direction, label in patterns:
        match = re.search(pattern, text, re.I)
        if not match:
            continue
        value = _client_threshold_value_with_tail(match.group(1).strip(), text[match.end():])
        return f"{label} {value}".strip(), direction
    rolling = re.search(r"\b\d+(?:\.\d+)?[- ]day\s+rolling(?:\s+baseline)?\b", text, re.I)
    if rolling:
        return rolling.group(0), "neutral"
    day = re.search(r"\bby\s+Day\s+\d+\b", text, re.I)
    if day:
        return day.group(0), "neutral"
    return "", "neutral"


def _client_threshold_value_with_tail(value: str, tail: str) -> str:
    tail_match = re.match(
        r"\s*((?:by|within|for)\s+[^.;,\n|]{1,60})",
        str(tail or ""),
        re.I,
    )
    if tail_match:
        return f"{value} {tail_match.group(1).strip()}"
    return value


def _client_threshold_keywords(text: str) -> set[str]:
    normalized = str(text or "").lower()
    keys: set[str] = set()
    keyword_patterns = {
        "activation rate": r"\bactivation\s+rate\b",
        "activation": r"\bactivation\b",
        "conversion": r"\bconversion\b",
        "baseline": r"\bbaseline\b",
        "retention": r"\bretention\b",
        "churn": r"\bchurn\b",
        "cac": r"\bcac\b",
        "nrr": r"\bnrr\b",
        "pipeline": r"\bpipeline\b",
        "time-to-value": r"\btime[-\s]to[-\s]value\b",
        "new trials": r"\bnew\s+trials\b",
        "trials": r"\btrials?\b",
    }
    for key, pattern in keyword_patterns.items():
        if re.search(pattern, normalized, re.I):
            keys.add(key)
    return keys


def _select_client_threshold_candidate(
    catalog: list[dict[str, Any]],
    line: str,
    *,
    preferred: tuple[str, ...] = (),
    prefer_positive: bool = False,
) -> dict[str, Any] | None:
    def direction_ok(candidate: dict[str, Any]) -> bool:
        return not prefer_positive or candidate.get("direction") == "positive"

    if "baseline" in preferred:
        for candidate in catalog:
            keys = candidate.get("keywords", set())
            if "baseline" in keys and re.search(r"\bbaseline\b", str(candidate.get("value", "")), re.I) and direction_ok(candidate):
                return candidate
    for candidate in catalog:
        keys = candidate.get("keywords", set())
        if preferred and any(key in keys for key in preferred) and direction_ok(candidate):
            return candidate
    line_keys = _client_threshold_keywords(line)
    for candidate in catalog:
        keys = candidate.get("keywords", set())
        if line_keys.intersection(keys) and direction_ok(candidate):
            return candidate
    for candidate in catalog:
        if direction_ok(candidate):
            return candidate
    return catalog[0] if catalog else None


def _client_threshold_at_least_phrase(candidate: dict[str, Any] | None) -> str:
    if not candidate:
        return "at the validated gate"
    value = str(candidate.get("value", "")).strip()
    if not value:
        return "at the validated gate"
    if value.lower().startswith("at least "):
        return value
    return f"at the validation gate ({value})"


def _client_threshold_metric_phrase(metric: str, candidate: dict[str, Any] | None) -> str:
    label = metric.replace("-", " ")
    if candidate and candidate.get("value"):
        return f"{label} target ({candidate['value']})"
    return f"{label} target"


def _client_threshold_plain_value(candidate: dict[str, Any] | None) -> str:
    if not candidate:
        return ""
    value = str(candidate.get("value", "")).strip()
    value = re.sub(r"^(?:at\s+least|below|down)\s+", "", value, flags=re.I).strip()
    return value


def _client_threshold_of_phrase(metric: str, candidate: dict[str, Any] | None) -> str:
    plain = _client_threshold_plain_value(candidate)
    if plain:
        return plain
    return f"validated gate for {metric.strip()}"


def _client_threshold_above_phrase(candidate: dict[str, Any] | None) -> str:
    plain = _client_threshold_plain_value(candidate)
    if plain:
        return f"above {plain}"
    return "above the validated gate"


def _client_threshold_below_phrase(candidate: dict[str, Any] | None) -> str:
    plain = _client_threshold_plain_value(candidate)
    if plain:
        return f"below {plain}"
    return "below the validated gate"


def _client_threshold_target_phrase(candidate: dict[str, Any] | None) -> str:
    plain = _client_threshold_plain_value(candidate)
    if plain:
        return f"target of {plain}"
    return "target of the validated gate"


def _client_threshold_expand_phrase(candidate: dict[str, Any] | None) -> str:
    if candidate and candidate.get("phrase"):
        return f"expand after the validation gate is met ({candidate['phrase']})"
    return "expand after the validation gate is met"


def _state_contains_starter_price(state: ProjectState) -> bool:
    parts = [
        getattr(state, "brief", ""),
        getattr(state, "data", ""),
        getattr(state, "report", ""),
    ]
    for item in list(getattr(getattr(state, "knowledge_layer", None), "items", []) or []):
        parts.extend([getattr(item, "title", ""), getattr(item, "summary", ""), getattr(item, "source_ref", "")])
    strategy = getattr(state, "strategy", None)
    if strategy:
        parts.extend([
            getattr(strategy, "executive_strategy", ""),
            getattr(strategy, "implementation_sequence", ""),
            getattr(strategy, "monitoring_plan", ""),
        ])
        for action in list(getattr(strategy, "strategies", []) or []):
            parts.extend([
                getattr(action, "action", ""),
                getattr(action, "justification", ""),
                getattr(action, "expected_impact", ""),
            ])
    return bool(re.search(r"\$\s*499\s*/\s*month|\$499/month", "\n".join(str(part) for part in parts), re.I))


def _normalize_client_evidence_count_language(markdown: str, state: ProjectState, quality) -> str:
    replacements = {
        "partial": (
            "Direct project evidence: Partial — supplied evidence exists, but several "
            "decision-critical evidence channels remain incomplete or unavailable."
        ),
        "moderate": (
            "Direct project evidence: Moderate — supplied project documents provide "
            "planning-level evidence, but validation gaps remain."
        ),
    }
    number_token = (
        r"(?:\d+|n|one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|"
        r"thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen|twenty)"
    )
    count_pattern = re.compile(
        rf"\bDirect project evidence:\s*(Partial|Moderate)\s*[—-]\s*{number_token}\s+"
        r"(?:(?:parsed|supplied|project|direct|evidence|source)\s+)*"
        r"(?:documents?|files?)\s*(?:supplied)?\b(?:\s*\([^)\n|]*\))?(?:[^\n|]*?)(?=\s*(?:\||$))",
        re.I,
    )

    def normalize_line(line: str) -> str:
        return count_pattern.sub(lambda match: replacements[match.group(1).lower()], line)

    lines: list[str] = []
    in_code_block = False
    source_lines = str(markdown or "").splitlines()
    index = 0
    while index < len(source_lines):
        line = source_lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            lines.append(line)
            index += 1
            continue
        if (
            in_code_block
            or _looks_like_json_line(stripped)
            or stripped.startswith(">")
            or _line_has_quoted_confirmed(line)
        ):
            lines.append(line)
            index += 1
            continue
        if _looks_like_markdown_table_line(line):
            table_lines: list[str] = []
            while index < len(source_lines) and _looks_like_markdown_table_line(source_lines[index]):
                table_lines.append(source_lines[index])
                index += 1
            if _is_protected_client_cleanup_table_block(table_lines):
                if count_pattern.search("\n".join(table_lines)):
                    lines.extend(normalize_line(table_line) for table_line in table_lines)
                else:
                    lines.extend(table_lines)
            else:
                lines.extend(normalize_line(table_line) for table_line in table_lines)
            continue
        lines.append(normalize_line(line))
        index += 1
    return "\n".join(lines)


def _label_client_bf_trace_language(markdown: str) -> str:
    lines: list[str] = []
    in_code_block = False
    for line in str(markdown or "").splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            lines.append(line)
            continue
        if (
            in_code_block
            or _looks_like_json_line(stripped)
            or stripped.startswith(">")
            or stripped.startswith("|")
        ):
            lines.append(line)
            continue
        value = re.sub(
            r"\bBF\s*=\s*(\d+(?:\.\d+)?)\b(?:\s*[—-]\s*domain complexity confirmed\b)?",
            r"structural BF estimate=\1 (operator trace, not measured posterior)",
            line,
            flags=re.I,
        )
        lines.append(value)
    return "\n".join(lines)


def _soften_unvalidated_confirmed_language(markdown: str, state: ProjectState, quality) -> str:
    projection = evidence_maturity_projection(state, quality)
    if projection.maturity == "Validated":
        return markdown
    accounting = evidence_accounting_projection(state)
    support_phrase = (
        "supported by multiple supplied evidence files"
        if accounting.parsed_file_count >= 2 or accounting.uploaded_file_count >= 2
        else "directionally supported"
    )
    lines: list[str] = []
    in_code_block = False
    for line in str(markdown or "").splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            lines.append(line)
            continue
        if (
            in_code_block
            or _looks_like_json_line(stripped)
            or stripped.startswith(">")
            or _line_has_quoted_confirmed(line)
            or _line_looks_like_operator_label(line)
            or _line_looks_like_future_confirmation_gate(line)
        ):
            lines.append(line)
            continue
        value = line
        value = re.sub(
            r"\b(onboarding friction) is confirmed(?:\s+by\s+[^.;\n]+)?",
            lambda match: f"{match.group(1)} is {support_phrase}",
            value,
            flags=re.I,
        )
        value = re.sub(
            r"\bconfirmed causal hypothesis\b",
            "candidate causal hypothesis pending validation",
            value,
            flags=re.I,
        )
        value = re.sub(
            r"\bcausal hypothesis is confirmed\b",
            "causal hypothesis remains pending validation",
            value,
            flags=re.I,
        )
        value = re.sub(
            r"\bconfirmed (driver|cause|bottleneck|friction|finding|growth lever|lever)\b",
            lambda match: f"{support_phrase} {match.group(1)}",
            value,
            flags=re.I,
        )
        value = re.sub(
            r"\bmechanically explained ([^.|]+)",
            lambda match: f"best-supported working diagnosis for {match.group(1).strip()} pending validation",
            value,
            flags=re.I,
        )
        value = re.sub(
            r"\bmechanically explained\b",
            "is a best-supported working diagnosis pending validation",
            value,
            flags=re.I,
        )
        value = re.sub(
            r"\bmechanically explains\b",
            "is the best-supported working diagnosis for",
            value,
            flags=re.I,
        )
        value = re.sub(
            r"\bmechanical explanation\b",
            "working diagnosis pending validation",
            value,
            flags=re.I,
        )
        value = re.sub(
            r"\bconfirmed root cause\b",
            "working diagnosis pending validation",
            value,
            flags=re.I,
        )
        value = re.sub(r"\bconfirm(?:s|ed)?\s+that\b", "indicate that", value, flags=re.I)
        value = re.sub(r"\bconfirms\b", "supports", value, flags=re.I)
        value = re.sub(r"\bis confirmed\b", f"is {support_phrase}", value, flags=re.I)
        value = re.sub(r"\bwas confirmed\b", f"was {support_phrase}", value, flags=re.I)
        value = re.sub(r"\bdomain complexity confirmed\b", "domain complexity is directionally supported", value, flags=re.I)
        lines.append(value)
    return "\n".join(lines)


def _line_has_quoted_confirmed(line: str) -> bool:
    return bool(re.search(r"[\"'“][^\"'”]*\bconfirm(?:s|ed)?\b[^\"'”]*[\"'”]", str(line or ""), re.I))


def _line_looks_like_operator_label(line: str) -> bool:
    return bool(re.search(r"\b(operator trace|not measured posterior|diagnostic score|provisional risk estimate)\b", str(line or ""), re.I))


def _line_looks_like_future_confirmation_gate(line: str) -> bool:
    value = str(line or "")
    return bool(
        re.search(r"\b(if|when|once|before|after)\b[^\n.]{0,160}\bconfirm(?:s|ed)?\b", value, re.I)
        and re.search(r"\b(gate|threshold|proceed|stop|extend|kill|sprint\s*0|decision)\b", value, re.I)
    )


def _looks_like_markdown_table_line(value: str) -> bool:
    stripped = str(value or "").strip()
    return stripped.startswith("|") and "|" in stripped[1:]


def _is_protected_client_cleanup_table_block(lines: list[str]) -> bool:
    text = "\n".join(str(line or "") for line in lines)
    if not text.strip():
        return False
    protected_patterns = [
        r"\b(?:source\s+excerpt|source\s+text|raw\s+source|quoted\s+source|source\s+quote|what\s+it\s+says)\b",
        r"\b(?:citation\s+marker|evidence\s+marker|evidence\s+ids?|evidence\s+locator|"
        r"locator\s+availability|locators?|provenance|source\s+ref|source_id|source\s+phase|"
        r"storage_ref|file_id|knowledge_id)\b",
        r"\b(?:machine_archive|machine\s+archive|project_state\.json|phase_outputs\.json|"
        r"decision_objects\.json|clarifications\.json|evidence_locator_register\.json|"
        r"uploaded_file_manifest\.json|export_manifest\.json)\b",
        r"\[Evidence:\s*[A-Za-z0-9_.:-]+",
        r"\[#\d+\]",
        r"\b(?:ev|evidence|knowledge|source|file)[-_][A-Za-z0-9_.:-]+\b",
        r"https?://|file://|upload:[^\s|)>\]]+|storage_ref\s*[:=]|upload_store[\\/]",
        r"\b[A-Za-z]:[\\/]|\\\\",
        r"\{[^{}\n]*\"?[A-Za-z0-9_.-]+\"?\s*:[^{}\n]*\}|\"[A-Za-z0-9_.-]+\"\s*:",
        r"`[^`]+`",
    ]
    return any(re.search(pattern, text, re.I) for pattern in protected_patterns)


def _looks_like_json_line(value: str) -> bool:
    stripped = str(value or "").strip()
    return bool(stripped) and (
        (stripped.startswith("{") and stripped.endswith("}"))
        or (stripped.startswith("[") and stripped.endswith("]"))
    )


def client_section_from_report_or_fallback(
    state: ProjectState,
    sections: dict[str, str],
    report_headings: tuple[str, ...],
    fallback: str,
) -> str:
    for heading in report_headings:
        value = sections.get(_normalize_heading(heading), "").strip()
        if value:
            return value
    return fallback


def _evidence_maturity_badge_markdown(state: ProjectState, quality, *, client: bool) -> list[str]:
    projection = evidence_maturity_projection(state, quality)
    accounting = evidence_accounting_projection(state)
    lines = [
        "## Evidence maturity",
        f"Evidence maturity: {projection.maturity}",
        f"Client-use status: {projection.client_use_status}",
        f"Validation required: {projection.validation_required}",
    ]
    if client:
        return lines
    rows = [
        ["Field", "Status"],
        ["citation_marker_count", str(accounting.citation_marker_count)],
        ["citation_markers_resolved_count", str(accounting.citation_markers_resolved_count)],
        ["citation_markers_resolved", str(accounting.citation_markers_resolved)],
        ["concrete_source_locator_count", str(accounting.concrete_source_locator_count)],
        ["concrete_source_locators_available", str(accounting.concrete_source_locators_available)],
        ["uploaded_file_count", str(accounting.uploaded_file_count)],
        ["parsed_file_count", str(accounting.parsed_file_count)],
        ["rejected_or_unsupported_file_count", str(accounting.rejected_or_unsupported_file_count)],
        ["imported_evidence_count", str(accounting.imported_evidence_count)],
        ["imported_evidence_available", str(accounting.imported_evidence_available)],
        ["imported_signal_count", str(accounting.imported_signal_count)],
        ["imported_signals_available", str(accounting.imported_signals_available)],
    ]
    if accounting.parsed_file_display_names:
        rows.append(["parsed_file_display_names", ", ".join(accounting.parsed_file_display_names)])
    if accounting.rejected_file_display_names:
        rows.append(["rejected_file_display_names", ", ".join(accounting.rejected_file_display_names)])
    return [*lines, _markdown_table(rows)]


def _requires_sparse_growth_decision_package(quality) -> bool:
    return bool(getattr(quality, "sparse_evidence", False) and getattr(quality, "decision_domain", "") == "growth")


def _sparse_growth_decision_package_sections(*, client: bool) -> list[str]:
    monitoring_label = "Monitoring Details" if client else "Operator Controls"
    return [
        "## Decision Gates",
        SPARSE_GROWTH_DECISION_GATES,
        "## Sprint 0 Evidence Pack Required",
        SPARSE_GROWTH_SPRINT0_EVIDENCE_PACK,
        f"## {monitoring_label}",
        "The Decision Gates matrix is the source of truth for proceed, extend, stop, and escalation choices. Additional canaries or circuit breakers are implementation controls, not separate decision systems.",
        "## Governance fallback if leadership overrides the diagnostic hold",
        SPARSE_GROWTH_GOVERNANCE_FALLBACK,
        "## Spend Gate Owner and Enforcement",
        SPARSE_GROWTH_SPEND_GATE,
        "## What the team may do during Sprint 0",
        SPARSE_GROWTH_SPRINT0_ACTIONS,
        "## Minimum staffing assumption",
        SPARSE_GROWTH_CAPACITY_NOTE,
        "## Main limitation of this recommendation",
        SPARSE_GROWTH_MAIN_LIMITATION,
    ]


def _quality_warning_blocks(state: ProjectState, quality, *, client: bool) -> list[str]:
    blocks: list[str] = []
    projection = evidence_maturity_projection(state, quality)
    accounting = evidence_accounting_projection(state)
    if quality.sparse_evidence:
        warnings = [quality.sparse_evidence_caveat]
        if accounting.unsupported_or_missing_warning:
            warnings.append(accounting.unsupported_or_missing_warning)
        elif accounting.category_coverage_warning:
            warnings.append(accounting.category_coverage_warning)
        blocks.extend(["## Evidence maturity warning", "\n\n".join(warnings)])
    elif projection.maturity == "Partial evidence":
        warnings = [PARTIAL_EVIDENCE_CAVEAT]
        if accounting.parsed_file_count > 0 and accounting.imported_evidence_count == 0:
            warnings.append(UPLOADED_KNOWLEDGE_NO_IMPORTED_EVIDENCE_NOTE)
        if accounting.unsupported_or_missing_warning:
            warnings.append(accounting.unsupported_or_missing_warning)
        elif accounting.category_coverage_warning:
            warnings.append(accounting.category_coverage_warning)
        blocks.extend(["## Evidence maturity warning", "\n\n".join(warnings)])
    elif quality.evidence_warning and not client:
        blocks.extend(["## Evidence maturity warning", "Some evidence channels are missing: " + "; ".join(quality.sparse_reasons)])
    if quality.provisional_report:
        blocks.extend([
            "## Provisional report warning",
            "\n\n".join([
                quality.provisional_clarification_caveat,
                quality.provisional_clarification_next_action,
            ]),
        ])
    if client and not quality.has_concrete_locators:
        blocks.extend(["## Source locator note", NO_CONCRETE_LOCATORS_CLIENT_NOTE])
    if not client and (quality.telemetry_privacy_required or requires_telemetry_privacy_caveat(_quality_content_text(state))):
        blocks.extend(["## Telemetry privacy note", TELEMETRY_PRIVACY_CAVEAT])
    if not client:
        risk_gate = assess_risk_classification_gate(state)
        if risk_gate.warning_applies:
            blocks.extend(["## Risk classification note", _risk_classification_gate_markdown(risk_gate)])
        constraint_projection = constraint_adherence_projection(state)
        if constraint_projection.warning_applies:
            blocks.extend([
                "## Constraint adherence warning",
                _constraint_adherence_warning_markdown(constraint_projection),
            ])
    warnings = threshold_consistency_warnings(state, quality)
    if client and warnings:
        blocks.extend(["## Threshold note", "\n".join(f"- {warning}" for warning in warnings)])
    elif not client:
        if warnings:
            blocks.extend(["## Threshold warnings", "\n".join(f"- {warning}" for warning in warnings)])
    return blocks


def _constraint_adherence_warning_markdown(projection) -> str:
    lines = [projection.warning_text]
    if getattr(projection, "detected_constraints", None):
        lines.append("")
        lines.append("Detected operator constraints:")
        lines.extend(f"- {item}" for item in projection.detected_constraints)
    if getattr(projection, "contradiction_signals", None):
        lines.append("")
        lines.append("Generated contradiction signals:")
        lines.extend(f"- {item}" for item in projection.contradiction_signals)
    return "\n".join(lines)


def _quality_content_text(state: ProjectState) -> str:
    parts = [
        state.brief or "",
        state.data or "",
        state.report or "",
    ]
    if state.strategy:
        parts.extend([
            state.strategy.executive_strategy,
            state.strategy.implementation_sequence,
            state.strategy.monitoring_plan,
            state.strategy.confidence,
            *state.strategy.success_metrics,
        ])
        for action in state.strategy.strategies or []:
            parts.extend([action.action, action.justification, action.expected_impact, action.risk_if_ignored])
    if state.monitor:
        parts.append(state.monitor.commitment_rationale)
        parts.extend(monitor_success_metric_lines(state.monitor, limit=50))
    return "\n".join(str(part) for part in parts if part)


def _risk_classification_gate_markdown(risk_gate) -> str:
    source_counts = ", ".join(
        f"{source}: {count}" for source, count in sorted((risk_gate.source_counts or {}).items())
    ) or "none"
    rows = [
        ["Diagnostic", "Value"],
        ["Selected classification", risk_gate.selected_classification or "unavailable"],
        ["Normalized classification", risk_gate.normalized_classification or "unavailable"],
        ["Highest generated risk severity", risk_gate.highest_generated_risk_severity or "none"],
        ["High/critical structured risk count", str(risk_gate.high_or_critical_risk_count)],
        ["Source counts", source_counts],
    ]
    parts = [risk_gate.warning_text or RISK_CLASSIFICATION_WARNING, _markdown_table(rows)]
    risk_rows = [["Source", "Severity", "Title", "Summary"]]
    for row in list(risk_gate.high_or_critical_risks or [])[:12]:
        risk_rows.append([
            row.get("source", ""),
            row.get("severity", ""),
            row.get("title", ""),
            row.get("summary", ""),
        ])
    if len(risk_rows) > 1:
        parts.extend(["### High/critical structured risk diagnostics", _markdown_table(risk_rows)])
    return "\n\n".join(parts)


def _client_decision_reviewed(state: ProjectState) -> str:
    brief = _short_text(state.brief, 700)
    return brief or "No decision brief is available yet."


def _client_recommended_path(state: ProjectState) -> str:
    if state.strategy and state.strategy.executive_strategy:
        return _redact_unsafe_string(state.strategy.executive_strategy)
    return EMPTY_STRATEGY


def _client_why_recommended(state: ProjectState) -> str:
    if state.strategy and state.strategy.strategies:
        rows = [["Reason", "Expected impact"]]
        for action in state.strategy.strategies[:5]:
            rows.append([
                action.justification or action.action or "Recommended action",
                action.expected_impact,
            ])
        return _markdown_table(rows)
    if state.audit and state.audit.top_findings:
        return "\n".join(f"- {item}" for item in state.audit.top_findings[:6])
    return EMPTY_STRATEGY


def _client_evidence_used(state: ProjectState) -> str:
    quality = assess_report_quality_context(state)
    rows = [["Source", "What it contributes"]]
    for evidence in state.imported_evidence or []:
        rows.append([evidence.title or evidence.evidence_id, evidence.summary or evidence.category])
    for item in list(getattr(getattr(state, "knowledge_layer", None), "items", []) or [])[:12]:
        title = getattr(item, "title", "") or getattr(item, "item_id", "")
        summary = getattr(item, "summary", "") or getattr(item, "source_ref", "")
        rows.append([title, summary])
    if len(rows) > 1:
        citation_summary = (
            _citation_locator_summary_markdown(state, include_registry=True)
            if quality.has_concrete_locators
            else NO_CONCRETE_LOCATORS_CLIENT_NOTE
        )
        sprint_pack = "" if _requires_sparse_growth_decision_package(quality) else _client_sprint0_pack(quality)
        return "\n\n".join(part for part in [_markdown_table(rows), citation_summary, sprint_pack] if part)
    if quality.sparse_evidence:
        sprint_pack = "" if _requires_sparse_growth_decision_package(quality) else _client_sprint0_pack(quality)
        return "\n\n".join(part for part in [EMPTY_EVIDENCE, sprint_pack] if part)
    return EMPTY_EVIDENCE


def _client_next_steps(state: ProjectState) -> str:
    if state.strategy and state.strategy.strategies:
        return "\n".join(
            f"- {action.action} ({action.timeline or 'timing TBD'})"
            for action in state.strategy.strategies[:6]
            if action.action
        ) or EMPTY_STRATEGY
    if state.strategy and state.strategy.implementation_sequence:
        return _redact_unsafe_string(state.strategy.implementation_sequence)
    return EMPTY_STRATEGY


def _client_roadmap(state: ProjectState) -> str:
    if state.strategy and state.strategy.strategies:
        rows = [["Window", "Action"]]
        for label in ("7 days", "30 days", "60 days", "90 days"):
            action = next(
                (
                    item.action
                    for item in state.strategy.strategies
                    if label.split()[0] in (item.timeline or "")
                ),
                "",
            )
            rows.append([label, action or "To be confirmed by the operator."])
        return _markdown_table(rows)
    return EMPTY_STRATEGY


def _client_key_risks(state: ProjectState) -> str:
    if state.audit and (state.audit.top_findings or state.audit.observation_needs):
        items = list(state.audit.top_findings or []) + list(state.audit.observation_needs or [])
        return "\n".join(f"- {item}" for item in items[:8])
    risks = list(getattr(getattr(state, "decision_objects", None), "risks", []) or [])
    if risks:
        return "\n".join(f"- {risk.title or risk.summary}" for risk in risks[:8])
    return EMPTY_AUDIT


def _client_monitoring(state: ProjectState) -> str:
    if state.monitor:
        return summarize_monitoring(state)
    if state.strategy and state.strategy.monitoring_plan:
        return _redact_unsafe_string(state.strategy.monitoring_plan)
    return EMPTY_MONITORING


def _client_open_questions(state: ProjectState) -> str:
    parts = []
    clarification_summary = _client_clarifications_markdown(state)
    if clarification_summary:
        parts.append(clarification_summary)
    if state.audit and state.audit.observation_needs:
        parts.append("\n".join(f"- {item}" for item in state.audit.observation_needs[:8]))
    return "\n\n".join(parts) if parts else "No open assumptions or questions are recorded yet."


def _client_sprint0_pack(quality) -> str:
    if not (quality.sparse_evidence or quality.evidence_warning):
        return ""
    if _requires_sparse_growth_decision_package(quality):
        return "### Sprint 0 Evidence Pack Required\n\n" + SPARSE_GROWTH_SPRINT0_EVIDENCE_PACK
    rows = [["Evidence to collect in Sprint 0", "Why it matters"]]
    for category in quality.evidence_categories[:8]:
        rows.append([category, "Validates assumptions before implementation."])
    return "Sprint 0 evidence collection should validate the current recommendation before acting.\n\n" + _markdown_table(rows)


def operator_project_metadata(state: ProjectState) -> str:
    return _markdown_table([
        ["Field", "Value"],
        ["Project ID", state.project_id],
        ["Project name", state.project_name],
        ["Created", _fmt_datetime(state.created_at)],
        ["Current phase", state.current_phase],
        ["Risk classification", state.risk_classification],
        ["Risk rationale", state.risk_classification_rationale],
    ])


def operator_executive_summary(state: ProjectState) -> str:
    sections = _extract_report_sections(state.report or "")
    report_summary = sections.get(_normalize_heading("Executive Summary"), "").strip()
    if report_summary:
        return report_summary
    return _operator_overview_markdown(state)


def operator_current_recommendation(state: ProjectState) -> str:
    if state.strategy and state.strategy.executive_strategy:
        return _redact_unsafe_string(state.strategy.executive_strategy)
    sections = _extract_report_sections(state.report or "")
    report_recommendation = sections.get(_normalize_heading("Recommended Path"), "").strip()
    return report_recommendation or EMPTY_STRATEGY


def operator_decision_snapshot(state: ProjectState) -> str:
    accounting = evidence_accounting_projection(state)
    rows = [
        ["Field", "Value"],
        ["Project status", _project_status_value(state)],
        ["Current phase", state.current_phase],
        ["Report status", _enum_value(state.phase_status.get("report", ""))],
        ["Risk classification", state.risk_classification],
        ["uploaded_file_count", str(accounting.uploaded_file_count)],
        ["parsed_file_count", str(accounting.parsed_file_count)],
        ["rejected_or_unsupported_file_count", str(accounting.rejected_or_unsupported_file_count)],
        ["imported_evidence_count", str(accounting.imported_evidence_count)],
        ["imported_signal_count", str(accounting.imported_signal_count)],
        ["concrete_source_locator_count", str(accounting.concrete_source_locator_count)],
        ["Clarification answers", str(len(state.clarification_answers or []))],
    ]
    return _markdown_table(rows)


def operator_dashboard_overview(state: ProjectState) -> str:
    return _operator_overview_markdown(state)


def operator_original_input(state: ProjectState) -> str:
    rows = [["Input", "Value"], ["Brief", state.brief or "No brief provided."]]
    if state.data:
        rows.append(["Original data/context", _short_text(state.data, 900)])
    else:
        rows.append(["Original data/context", "No original data/context was provided."])
    return _markdown_table(rows)


def operator_classification_summary(state: ProjectState) -> str:
    if not state.classify:
        return "This section will populate after the classify phase runs."
    c = state.classify
    return _markdown_table([
        ["Field", "Value"],
        ["Domain", c.domain],
        ["Justification", c.justification],
        ["Structural BF estimate (operator trace, not measured posterior)", _fmt_value(c.bf)],
        ["Reference class", c.reference_class],
        ["DQ diagnostic score", ", ".join(_fmt_value(value) for value in c.dq)],
        ["Variety gaps", c.variety_gaps],
        ["RPD pattern", c.rpd_pattern],
        ["Sensemaking anchors", c.sensemaking_anchors],
        ["OODA cadence", c.ooda.freq],
    ])


def operator_hypotheses_table(state: ProjectState) -> str:
    rows = [["ID", "Hypothesis", "Status", "Confirm", "Reject", "EVOI", "Cluster", "Evidence IDs"]]
    for hypothesis in state.hypotheses or []:
        rows.append([
            hypothesis.id,
            hypothesis.text,
            hypothesis.status,
            hypothesis.confirm,
            hypothesis.reject,
            hypothesis.evoi,
            hypothesis.portfolio_cluster,
            ", ".join(hypothesis.evidence_ids or []),
        ])
    if len(rows) <= 1:
        return EMPTY_HYPOTHESES
    coverage = _operator_variable_coverage_summary(state)
    return "\n\n".join(part for part in (_markdown_table(rows), coverage) if part)


def _operator_variable_coverage_summary(state: ProjectState) -> str:
    coverage = assess_hypothesis_variable_coverage(state)
    if not coverage.has_hypotheses:
        return ""
    covered = ", ".join(coverage.covered_categories) or "None detected."
    missing = ", ".join(coverage.missing_critical_categories) or "No decision-critical gaps flagged."
    lines = [
        "### Variable coverage summary",
        f"Covered categories: {covered}",
        f"Missing decision-critical categories: {missing}",
    ]
    if coverage.evidence_needs:
        lines.append("Evidence needs:")
        lines.extend(
            f"- {need.category}: {need.evidence_need}"
            for need in coverage.evidence_needs[:8]
        )
    return "\n".join(lines)


def operator_gauntlet_summary(state: ProjectState) -> str:
    if not state.gauntlet:
        return "This section will populate after the gauntlet phase runs."
    return _summarize_gauntlet(state)


def operator_audit_summary(state: ProjectState) -> str:
    if not state.audit and not state.audit_raw:
        return EMPTY_AUDIT
    return _summarize_audit(state)


def operator_evidence_summary(state: ProjectState) -> str:
    parts = []
    accounting = evidence_accounting_projection(state)
    rows = [
        ["Field", "Value"],
        ["citation_marker_count", str(accounting.citation_marker_count)],
        ["citation_markers_resolved_count", str(accounting.citation_markers_resolved_count)],
        ["citation_markers_resolved", str(accounting.citation_markers_resolved)],
        ["concrete_source_locator_count", str(accounting.concrete_source_locator_count)],
        ["concrete_source_locators_available", str(accounting.concrete_source_locators_available)],
        ["uploaded_file_count", str(accounting.uploaded_file_count)],
        ["parsed_file_count", str(accounting.parsed_file_count)],
        ["rejected_or_unsupported_file_count", str(accounting.rejected_or_unsupported_file_count)],
        ["imported_evidence_count", str(accounting.imported_evidence_count)],
        ["imported_evidence_available", str(accounting.imported_evidence_available)],
        ["imported_signal_count", str(accounting.imported_signal_count)],
        ["imported_signals_available", str(accounting.imported_signals_available)],
    ]
    if accounting.parsed_file_display_names:
        rows.append(["parsed_file_display_names", ", ".join(accounting.parsed_file_display_names)])
    if accounting.rejected_file_display_names:
        rows.append(["rejected_file_display_names", ", ".join(accounting.rejected_file_display_names)])
    parts.extend(["### Evidence accounting", _markdown_table(rows)])
    if accounting.unsupported_or_missing_warning:
        parts.append(accounting.unsupported_or_missing_warning)
    elif accounting.category_coverage_warning:
        parts.append(accounting.category_coverage_warning)
    threshold_debug = _operator_threshold_section_classification_markdown(state)
    if threshold_debug:
        parts.extend(["### Threshold section classification", threshold_debug])
    evidence_rows = [["Evidence ID", "Title", "Summary", "Source phase"]]
    for evidence in state.imported_evidence or []:
        evidence_rows.append([evidence.evidence_id, evidence.title, evidence.summary, evidence.source_phase])
    if len(evidence_rows) > 1:
        parts.append(_markdown_table(evidence_rows))
    else:
        parts.append(EMPTY_EVIDENCE)

    uploads = _uploaded_file_manifest_payload(state)
    if uploads:
        rows = [["File ID", "Filename", "Type", "Size", "Parse status"]]
        for upload in uploads:
            parse_summary = upload.get("parse_summary") or {}
            rows.append([
                upload.get("file_id", ""),
                upload.get("original_filename", ""),
                upload.get("content_type", ""),
                upload.get("size_bytes", ""),
                parse_summary.get("status", ""),
            ])
        parts.extend(["### Uploaded files", _markdown_table(rows)])
    else:
        parts.extend(["### Uploaded files", EMPTY_UPLOADED_FILES])

    signals = state.imported_signals or []
    if signals:
        rows = [["Signal ID", "Name", "Kind", "Cadence"]]
        for signal in signals[:20]:
            rows.append([signal.signal_id, signal.name, signal.kind, signal.cadence])
        parts.extend(["### Imported signals", _markdown_table(rows)])
    else:
        parts.extend(["### Imported signals", "No imported signals are available yet."])

    locator_register = _evidence_locator_register_markdown(state)
    if locator_register and "No evidence locator registry entries" not in locator_register:
        parts.extend(["### Evidence locator register", locator_register])
    citation_summary = _citation_locator_summary_markdown(state, include_registry=False)
    if citation_summary:
        parts.extend(["### Citation locator review", citation_summary])
    return "\n\n".join(parts)


def _operator_threshold_section_classification_markdown(state: ProjectState) -> str:
    classifications = threshold_section_classification(state)
    if not classifications:
        return ""
    rows = [["Section", "Classification", "Reason"]]
    for item in classifications:
        if item.classification == "ignored" and not item.threshold_like:
            continue
        rows.append([
            _safe_threshold_debug_text(item.section_name),
            _safe_threshold_debug_text(item.classification),
            _safe_threshold_debug_text(item.reason),
        ])
    return _markdown_table(rows) if len(rows) > 1 else ""


def _safe_threshold_debug_text(value: str) -> str:
    text = _redact_unsafe_string(str(value or ""))
    text = re.sub(r"\bupload:[^\s|]+", REDACTED, text, flags=re.I)
    text = re.sub(r"\bstorage:[^\s|]+", REDACTED, text, flags=re.I)
    text = re.sub(r"\bprovider[_-]?payload\b.*", "provider payload redacted", text, flags=re.I)
    return _short_text(text, 140)


def operator_strategy_summary(state: ProjectState) -> str:
    if not state.strategy and not state.strategy_raw:
        return EMPTY_STRATEGY
    return _summarize_strategy(state)


def operator_sqi_summary(state: ProjectState) -> str:
    if not state.sqi:
        return "This section will populate after the SQI phase runs."
    rows = [["Field", "Value"], ["SQI overall", _fmt_value(state.sqi.sqi_overall)], ["Weakest link", state.sqi.weakest_link]]
    for dimension in state.sqi.dimensions or []:
        rows.append([f"Dimension: {dimension.name}", f"{_fmt_value(dimension.score)} ({dimension.grade}) - {dimension.finding}"])
    if state.sqi.improvement_actions:
        rows.append(["Improvement actions", "; ".join(state.sqi.improvement_actions)])
    coverage = assess_hypothesis_variable_coverage(state)
    if coverage.missing_critical_categories:
        rows.append([
            "Variable coverage limitation",
            "Hypothesis set may under-cover: " + ", ".join(coverage.missing_critical_categories[:6]),
        ])
    return _markdown_table(rows)


def operator_monitoring_summary(state: ProjectState) -> str:
    summary = summarize_monitoring(state) if state.monitor else EMPTY_MONITORING
    if MONITORING_TEMPLATE_OPERATOR_NOTE in summary:
        return summary
    return "\n\n".join([summary, f"Operator note: {MONITORING_TEMPLATE_OPERATOR_NOTE}"])


def operator_workspace_summary(state: ProjectState) -> str:
    return _workspace_summary_markdown(state)


def operator_risks_and_questions(state: ProjectState) -> str:
    parts = []
    risks = list(getattr(getattr(state, "decision_objects", None), "risks", []) or [])
    if risks:
        rows = [["Risk ID", "Title", "Severity", "Status", "Summary"]]
        for risk in risks[:20]:
            rows.append([risk.risk_id, risk.title, risk.severity, risk.status, risk.summary])
        parts.append(_markdown_table(rows))
    elif state.audit and state.audit.top_findings:
        parts.append("\n".join(f"- {item}" for item in state.audit.top_findings[:10]))
    else:
        parts.append(EMPTY_AUDIT)
    clarification_summary = _client_clarifications_markdown(state)
    if clarification_summary:
        parts.extend(["### Open clarification questions", clarification_summary])
    return "\n\n".join(parts)


def operator_trace_summary(state: ProjectState) -> str:
    if not any([
        state.classify,
        state.hypotheses,
        state.gauntlet,
        state.audit,
        state.strategy,
        state.sqi,
        state.monitor,
        state.report,
    ]):
        return EMPTY_TRACE
    summary = summarize_trace(state)
    return summary if summary and "unavailable" not in summary.lower() else EMPTY_TRACE


def operator_clarifications_summary(state: ProjectState) -> str:
    summary = _operator_clarifications_markdown(state)
    if summary == "No clarification questions or answers saved.":
        return EMPTY_CLARIFICATIONS
    if not state.clarification_answers:
        return "\n\n".join([summary, EMPTY_CLARIFICATIONS])
    return summary


def operator_report_appendix(state: ProjectState) -> str:
    return _redact_unsafe_string(state.report or "") if state.report else "No report is available yet."


def operator_technical_appendix(state: ProjectState) -> str:
    parts = [
        "### Risk and policy summary",
        summarize_policy(state),
        "### Budget summary",
        _budget_summary_markdown(state),
        "### Approvals",
        _approvals_summary_markdown(state),
        "### Calibration and prediction summary",
        _calibration_prediction_summary_markdown(state),
    ]
    return "\n\n".join(parts)


def _project_status_value(state: ProjectState) -> str:
    statuses = {_enum_value(value) for value in (state.phase_status or {}).values()}
    if "running" in statuses:
        return "running"
    if "failed" in statuses:
        return "failed"
    if statuses and statuses <= {"completed"}:
        return "completed"
    return "in progress"


def build_machine_archive_payload(state: ProjectState) -> dict[str, Any]:
    decision_objects = _decision_objects_payload(state)
    files: dict[str, Any] = {
        "project_state.json": sanitize_for_export(state, "machine_archive", mode="redact"),
        "report.md": _redact_unsafe_string(state.report or "No report available."),
        "phase_outputs.json": sanitize_for_export(_phase_outputs_payload(state), "machine_archive", mode="redact"),
        "decision_objects.json": sanitize_for_export(decision_objects, "machine_archive", mode="redact"),
        "clarifications.json": sanitize_for_export(_clarifications_payload(state), "machine_archive", mode="redact"),
        "evidence_locator_register.json": sanitize_for_export(_evidence_locator_payload(state), "machine_archive", mode="redact"),
        "uploaded_file_manifest.json": sanitize_for_export(_uploaded_file_manifest_payload(state), "machine_archive", mode="redact"),
        "policy_summary.json": sanitize_for_export(_policy_summary_payload(state), "machine_archive", mode="redact"),
    }

    calibration = _calibration_predictions_payload(state)
    if calibration:
        files["calibration_predictions.json"] = sanitize_for_export(calibration, "machine_archive", mode="redact")
    approvals = _approvals_payload(state)
    if approvals:
        files["approvals_summary.json"] = sanitize_for_export(approvals, "machine_archive", mode="redact")
    risks = _risk_summary_payload(state)
    if risks:
        files["risk_summary.json"] = sanitize_for_export(risks, "machine_archive", mode="redact")

    included_files = ["export_manifest.json", *sorted(files.keys())]
    files["export_manifest.json"] = build_export_manifest(
        state,
        "machine_archive",
        "zip",
        included_files=included_files,
    )
    return files


def build_export_manifest(
    state: ProjectState,
    profile: str,
    format: str,
    *,
    included_files: list[str] | None = None,
) -> dict[str, Any]:
    current_version = report_freshness.current_code_version()
    return sanitize_for_export(
        {
            "export_schema_version": "1.0",
            "project_id": state.project_id,
            "project_name": state.project_name,
            "export_profile": profile,
            "export_format": format,
            "generated_at": _utc_now(),
            "code_version": current_version,
            "report_freshness": report_freshness.report_freshness_manifest(
                state,
                current_version=current_version,
            ),
            "redaction_policy": "Recursive sensitive-key, unsafe-path, prompt/debug, provider-payload, and chain-of-thought redaction.",
            "included_files": included_files or [],
        },
        profile,
        mode="redact",
    )


def sanitize_for_export(value, profile, mode: str = "redact"):
    sanitized = _sanitize_for_export(value, profile=profile, mode=mode)
    if sanitized is _DROP:
        return None if mode == "redact" else {}
    return sanitized


def summarize_phase_outputs(state: ProjectState) -> str:
    rows = [["Phase", "Status", "Confidence", "Completed at", "Summary"]]
    for phase in ("classify", "hypotheses", "gauntlet", "audit", "strategy", "sqi", "monitor", "report"):
        status = state.phase_status.get(phase, "")
        rows.append([
            phase,
            _enum_value(status),
            _fmt_value(state.phase_confidence.get(phase)),
            state.phase_run_completed_at.get(phase, ""),
            state.phase_summaries.get(phase, _phase_output_digest(state, phase)),
        ])
    return _markdown_table(rows)


def _phase_output_digest(state: ProjectState, phase: str) -> str:
    if phase == "classify" and state.classify:
        return _short_text(f"{state.classify.domain}: {state.classify.justification}", 240)
    if phase == "hypotheses" and state.hypotheses:
        return f"{len(state.hypotheses)} hypotheses generated."
    if phase == "gauntlet" and state.gauntlet:
        return f"{len(state.gauntlet.results or [])} stress-test result(s); MECE gaps: {_short_text(state.gauntlet.mece_gaps, 120)}"
    if phase == "audit" and (state.audit or state.audit_raw):
        if state.audit:
            return f"{len(state.audit.top_findings or [])} finding(s), {len(state.audit.observation_needs or [])} observation need(s)."
        return "Audit output is saved; structured summary is unavailable."
    if phase == "strategy" and (state.strategy or state.strategy_raw):
        if state.strategy:
            return _short_text(state.strategy.executive_strategy or f"{len(state.strategy.strategies or [])} strategy action(s).", 240)
        return "Strategy output is saved; structured summary is unavailable."
    if phase == "sqi" and state.sqi:
        return f"SQI overall: {_fmt_value(state.sqi.sqi_overall)}; weakest link: {_short_text(state.sqi.weakest_link, 120)}"
    if phase == "monitor" and state.monitor:
        return f"{len(state.monitor.canaries or [])} canary signal(s), {len(state.monitor.circuit_breakers or [])} circuit breaker(s)."
    if phase == "report" and state.report:
        return _short_text(state.report, 240)
    return ""


def summarize_hypotheses(state: ProjectState) -> str:
    rows = [["ID", "Hypothesis", "Alpha/Beta structural priors", "Status", "Confirm", "Reject", "EVOI", "Cluster"]]
    for hypothesis in state.hypotheses or []:
        rows.append([
            hypothesis.id,
            hypothesis.text,
            f"{hypothesis.alpha}/{hypothesis.beta}",
            hypothesis.status,
            hypothesis.confirm,
            hypothesis.reject,
            hypothesis.evoi,
            hypothesis.portfolio_cluster,
        ])
    return _markdown_table(rows) if len(rows) > 1 else "No hypotheses saved."


def summarize_monitoring(state: ProjectState) -> str:
    if not state.monitor:
        return "No monitor output saved."
    lines = [
        commitment_score_text(state.monitor.commitment_score, state.monitor.commitment_rationale),
        f"Commitment rationale: {state.monitor.commitment_rationale or 'TBD — requires operator confirmation.'}",
        "### OODA Schedule",
        _markdown_table(
            [["Cadence", "Metric", "Owner", "Source"]]
            + [["Daily", item.metric, item.owner, item.source] for item in state.monitor.ooda_schedule.daily]
            + [["Weekly", item.metric, item.owner, item.source] for item in state.monitor.ooda_schedule.weekly]
            + [["Monthly", item.metric, item.owner, item.source] for item in state.monitor.ooda_schedule.monthly]
        ),
        "### Circuit Breakers",
        _markdown_table(
            [["Strategy", "Trip", "Reset"]]
            + [[item.strategy_ref, item.trip, item.reset] for item in state.monitor.circuit_breakers]
        ),
        "### Canaries",
        _markdown_table(
            [["Signal", "Direction", "Window", "Meaning"]]
            + [[item.signal, item.direction, item.window, item.meaning] for item in state.monitor.canaries]
        ),
    ]
    return "\n\n".join(lines)


def summarize_policy(state: ProjectState) -> str:
    event_counts: dict[str, int] = {}
    for event in state.policy_audit_log or []:
        event_type = str((event or {}).get("event_type", "unknown"))
        event_counts[event_type] = event_counts.get(event_type, 0) + 1
    rows = [
        ["Field", "Value"],
        ["Risk classification", state.risk_classification],
        ["Risk rationale", state.risk_classification_rationale],
        ["Set by", state.risk_classification_set_by],
        ["Kill switch active", str(bool(state.kill_switch_active))],
        ["Kill switch reason", state.kill_switch_reason],
        ["Policy event count", str(len(state.policy_audit_log or []))],
        ["Policy event types", ", ".join(f"{key}={value}" for key, value in sorted(event_counts.items()))],
    ]
    return _markdown_table(rows)


def summarize_trace(state: ProjectState) -> str:
    try:
        from explainability import build_project_trace

        trace = build_project_trace(state)
        rows = [["Phase", "Status", "Purpose", "Output summary"]]
        for phase in trace.phases:
            rows.append([
                phase.phase,
                phase.status,
                phase.purpose,
                _short_text(phase.output_summary, 260),
            ])
        return _markdown_table(rows)
    except Exception as exc:
        return f"Decision trace summary unavailable: {_short_text(str(exc), 180)}"


def _as_pdf_text(text: str) -> str:
    return escape((text or "").replace("\n", "<br/>"))


def _build_dossier_blocks(state: ProjectState) -> list[dict]:
    return _markdown_to_blocks(build_operator_dossier_markdown(state, include_freshness=False))

    blocks: list[dict] = []

    def heading(text: str, level: int = 1):
        blocks.append({"type": "heading", "text": text, "level": level})

    def paragraph(text: str):
        clean = (text or "").strip()
        if clean:
            blocks.append({"type": "paragraph", "text": clean})

    def bullets(items: list[str]):
        clean_items = [str(item).strip() for item in items if str(item).strip()]
        if clean_items:
            blocks.append({"type": "bullets", "items": clean_items})

    heading(state.project_name or "Decision Dossier", 1)
    paragraph(
        " | ".join(filter(None, [
            f"Project ID: {state.project_id}",
            f"Created: {_fmt_datetime(state.created_at)}",
            f"Risk: {state.risk_classification}",
            f"Current phase: {state.current_phase}",
        ]))
    )

    heading("Input", 2)
    paragraph("Brief")
    paragraph(state.brief or "No brief provided.")
    if state.data:
        paragraph("Supporting data")
        paragraph(state.data)
    if state.observations:
        bullets([f"Observation - {k}: {v}" for k, v in state.observations.items()])
    if state.timer_logs:
        bullets([f"Timer - {entry.get('time', '')}: {entry.get('label', '')}" for entry in state.timer_logs])

    heading("Classify", 2)
    if state.classify:
        c = state.classify
        paragraph(f"Domain: {c.domain}")
        paragraph(f"Justification: {c.justification}")
        paragraph(f"Bayes factor: {c.bf}")
        bullets([
            f"Variety environment: {c.variety_env}",
            f"Variety system: {c.variety_sys}",
            f"Variety gaps: {c.variety_gaps}",
            f"Variety decision: {c.variety_decision}",
            f"RPD pattern: {c.rpd_pattern}",
            f"Sensemaking anchors: {c.sensemaking_anchors}",
            f"Expectancy violations: {c.expectancy_violations}",
            f"Reference class: {c.reference_class}",
            f"DQ: {', '.join(str(v) for v in c.dq)}",
            f"Maturity assessment: {c.maturity_assessment}",
            f"Spiral depth: {c.spiral_depth}",
        ])
        heading("OODA", 3)
        bullets([
            f"Observe: {c.ooda.observe}",
            f"Orient: {c.ooda.orient}",
            f"Decide: {c.ooda.decide}",
            f"Act: {c.ooda.act}",
            f"Frequency: {c.ooda.freq}",
        ])
    else:
        paragraph("No classify output saved.")

    heading("Hypotheses", 2)
    if state.hypotheses:
        for hypothesis in state.hypotheses:
            heading(f"{hypothesis.id}: {hypothesis.text}", 3)
            paragraph(f"Justification: {hypothesis.justification}")
            bullets([
                f"Signal: {hypothesis.signal}",
                f"Alpha/Beta: {hypothesis.alpha}/{hypothesis.beta}",
                f"Confirm: {hypothesis.confirm}",
                f"Reject: {hypothesis.reject}",
                f"EVOI: {hypothesis.evoi}",
                f"Portfolio cluster: {hypothesis.portfolio_cluster}",
                f"Status: {hypothesis.status}",
            ])
    else:
        paragraph("No hypotheses saved.")

    heading("Gauntlet", 2)
    if state.gauntlet:
        paragraph(f"Portfolio correlation: {state.gauntlet.portfolio_correlation}")
        paragraph(f"MECE gaps: {state.gauntlet.mece_gaps}")
        paragraph(f"Thompson priority: {state.gauntlet.thompson_priority}")
        paragraph(f"EVOI ranking: {state.gauntlet.evoi_ranking}")
        for result in state.gauntlet.results:
            heading(f"{result.id} - risk rank {result.risk_rank}", 3)
            bullets([
                f"Crux: {result.crux}",
                f"Top FMEA: {result.top_fmea}",
                f"FTA cut set: {result.fta_cut_set}",
            ])
            if result.frameworks:
                bullets([
                    f"{fw.get('fw', 'FW')}: {fw.get('finding', '')} | action={fw.get('action', '')}"
                    for fw in result.frameworks
                ])
    else:
        paragraph("No gauntlet output saved.")

    heading("Audit", 2)
    if state.audit:
        paragraph(f"Data based: {state.audit.data_based}")
        bullets([f"Top finding: {item}" for item in state.audit.top_findings])
        bullets([f"Observation need: {item}" for item in state.audit.observation_needs])
        bullets([
            f"FMEA {item.component}: {item.failure_mode} | RPN {item.rpn} | action {item.action}"
            for item in state.audit.fmea
        ])
    elif state.audit_raw:
        paragraph(state.audit_raw)
    else:
        paragraph("No audit output saved.")

    heading("Strategy", 2)
    if state.strategy:
        paragraph(state.strategy.executive_strategy)
        heading("Preliminary verdicts", 3)
        bullets([
            f"{item.id}: {item.verdict} | evidence: {item.evidence} | monitoring: {item.monitoring_plan}"
            for item in state.strategy.preliminary_verdicts
        ])
        heading("Actions", 3)
        for action in state.strategy.strategies:
            heading(f"{action.priority}: {action.action}", 4)
            bullets([
                f"Justification: {action.justification}",
                f"Evidence chain: {action.evidence_chain}",
                f"Expected impact: {action.expected_impact}",
                f"Effort: {action.effort}",
                f"Timeline: {action.timeline}",
                f"Risk if ignored: {action.risk_if_ignored}",
                f"Framework source: {action.framework_source}",
            ])
        if state.strategy.success_metrics:
            bullets([f"Success metric: {metric}" for metric in state.strategy.success_metrics])
        elif monitor_has_signals(state.monitor):
            paragraph("Success metrics are captured in the monitoring plan below.")
            bullets(monitor_success_metric_lines(state.monitor))
        paragraph(f"Implementation sequence: {state.strategy.implementation_sequence}")
        paragraph(f"Monitoring plan: {state.strategy.monitoring_plan}")
        paragraph(f"Review date: {state.strategy.review_date}")
        paragraph(f"Confidence: {state.strategy.confidence}")
        paragraph(f"Re-entry check: {state.strategy.reentry_check}")
    elif state.strategy_raw:
        paragraph(state.strategy_raw)
    else:
        paragraph("No strategy output saved.")

    heading("Monitor", 2)
    if state.monitor:
        heading("OODA schedule", 3)
        bullets(_schedule_lines("Daily", state.monitor.ooda_schedule.daily))
        bullets(_schedule_lines("Weekly", state.monitor.ooda_schedule.weekly))
        bullets(_schedule_lines("Monthly", state.monitor.ooda_schedule.monthly))
        bullets([
            f"Circuit breaker - {item.strategy_ref}: trip {item.trip} | reset {item.reset}"
            for item in state.monitor.circuit_breakers
        ])
        bullets([
            f"Canary - {item.signal}: {item.direction} over {item.window} | {item.meaning}"
            for item in state.monitor.canaries
        ])
        bullets([
            f"Chaos drill - {item.what}: when {item.when} | measure {item.measure}"
            for item in state.monitor.chaos_drills
        ])
        bullets([f"HRO principle: {item}" for item in state.monitor.hro_principles_active])
        bullets([f"Re-entry watch: {item}" for item in state.monitor.reentry_watch])
        paragraph(commitment_score_text(state.monitor.commitment_score, state.monitor.commitment_rationale))
        paragraph(f"Commitment rationale: {state.monitor.commitment_rationale}")
    else:
        paragraph("No monitor output saved.")

    heading("Report", 2)
    if state.report:
        _append_markdownish_report(blocks, state.report)
    else:
        paragraph("No report saved.")

    return blocks


def _schedule_lines(label: str, entries: list) -> list[str]:
    return [
        f"{label} - {entry.metric} | owner: {entry.owner} | source: {entry.source}"
        for entry in entries
    ]


def _append_markdownish_report(blocks: list[dict], report: str) -> None:
    current_bullets: list[str] = []

    def flush_bullets():
        nonlocal current_bullets
        if current_bullets:
            blocks.append({"type": "bullets", "items": current_bullets})
            current_bullets = []

    for raw_line in (report or "").splitlines():
        line = raw_line.strip()
        if not line:
            flush_bullets()
            continue
        if line.startswith("### "):
            flush_bullets()
            blocks.append({"type": "heading", "text": line[4:].strip(), "level": 3})
        elif line.startswith("## "):
            flush_bullets()
            blocks.append({"type": "heading", "text": line[3:].strip(), "level": 2})
        elif line.startswith("# "):
            flush_bullets()
            blocks.append({"type": "heading", "text": line[2:].strip(), "level": 2})
        elif line.startswith("- "):
            current_bullets.append(line[2:].strip())
        else:
            flush_bullets()
            blocks.append({"type": "paragraph", "text": line})
    flush_bullets()


def _fmt_datetime(value) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat()
    return str(value)


def _profile_filename(state: ProjectState, profile: str, fmt: str) -> str:
    project_id = re.sub(r"[^A-Za-z0-9-]+", "-", state.project_id or "project").strip("-") or "project"
    return f"{project_id}-{profile}-{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}.{fmt}"


def _safe_report_markdown(state: ProjectState) -> str:
    markdown = _redact_unsafe_string(state.report or "No report available.")
    quality = assess_report_quality_context(state)
    if quality.sparse_evidence:
        markdown = _simplify_sparse_report_markdown(markdown, quality)
    markdown = _prepend_evidence_maturity_badge(markdown, state, quality)
    if quality.provisional_report and quality.provisional_clarification_caveat not in markdown:
        provisional_warning = "\n\n".join([
            "## Provisional report warning",
            quality.provisional_clarification_caveat,
            quality.provisional_clarification_next_action,
        ])
        markdown = "\n\n".join([provisional_warning, markdown.strip()])
    if _requires_sparse_growth_decision_package(quality):
        markdown = _with_sparse_growth_decision_package(markdown)
    if requires_productization_wave_matrix(state, quality) and "Wave 2 Graduation Matrix" not in markdown:
        markdown = "\n\n".join([markdown.strip(), WAVE2_GRADUATION_MATRIX])
    if not quality.has_concrete_locators:
        markdown = suppress_client_raw_evidence_ids(markdown)
    markdown = normalize_export_text(markdown, audience="client")
    if _requires_sparse_growth_decision_package(quality):
        markdown = _dedupe_client_sparse_growth_sprint0(markdown)
        markdown = guard_client_bf_confidence(markdown, state)
    return _finalize_export_markdown(markdown, state, audience="client", quality=quality)


def _prepend_report_freshness_warning(
    markdown: str,
    state: ProjectState,
    *,
    current_code_version: str | None = None,
) -> str:
    warning = _report_freshness_warning(state, current_code_version=current_code_version)
    if not warning:
        return markdown
    return "\n\n".join([warning, markdown])


def _prepend_evidence_maturity_badge(markdown: str, state: ProjectState, quality) -> str:
    if "Evidence maturity:" in str(markdown or ""):
        return markdown
    badge = "\n".join(_evidence_maturity_badge_markdown(state, quality, client=True))
    return "\n\n".join([badge, markdown.strip()])


def _with_sparse_growth_decision_package(markdown: str) -> str:
    source = _remove_markdown_section(markdown, "Decision Gates")
    source = _remove_markdown_section(source, "Sprint 0 Evidence Pack Required")
    source = _remove_markdown_section(source, "Governance fallback if leadership overrides the diagnostic hold")
    source = _remove_markdown_section(source, "Spend Gate Owner and Enforcement")
    source = _remove_markdown_section(source, "What the team may do during Sprint 0")
    source = _remove_markdown_section(source, "Minimum staffing assumption")
    source = _remove_markdown_section(source, "Main limitation of this recommendation")
    package = "\n\n".join(_sparse_growth_decision_package_sections(client=True))
    road_map = re.search(r"(?im)^#{1,6}\s+(?:Roadmap|Timeline / 7-30-60-90 roadmap)\s*$", source)
    if road_map:
        return "\n\n".join([source[: road_map.start()].strip(), package, source[road_map.start():].strip()])
    return "\n\n".join([source.strip(), package])


def _dedupe_client_sparse_growth_sprint0(markdown: str) -> str:
    source = str(markdown or "")
    split_match = re.search(
        r"(?im)^#{1,6}\s+(?:Appendix|Technical appendix|SQI / quality review|Detailed evidence appendix)\b.*$",
        source,
    )
    if split_match:
        main = source[: split_match.start()]
        appendix = source[split_match.start():]
    else:
        main = source
        appendix = ""

    seen = False

    def sprint_section_repl(match: re.Match[str]) -> str:
        nonlocal seen
        if seen:
            return ""
        seen = True
        return match.group(0)

    main = re.sub(
        r"(?ims)^#{1,6}\s+Sprint 0 Evidence Pack Required\s*$.*?(?=^#{1,6}\s+|\Z)",
        sprint_section_repl,
        main,
    )
    if seen:
        main = re.sub(
            r"(?ims)^\|\s*Evidence Item\s*\|\s*Why It Is Needed\s*\|\s*Decision It Validates\s*\|.*?(?=^\s*$|^#{1,6}\s+|\Z)",
            "",
            main,
        )
    return "\n\n".join(part.strip() for part in (main, appendix) if part.strip())


def _remove_markdown_section(markdown: str, heading: str) -> str:
    source = str(markdown or "")
    pattern = re.compile(
        rf"(?ims)^#{{1,6}}\s+{re.escape(heading)}\s*$.*?(?=^#{{1,6}}\s+|\Z)"
    )
    return pattern.sub("", source).strip()


def _report_freshness_warning(
    state: ProjectState,
    *,
    current_code_version: str | None = None,
) -> str:
    freshness = report_freshness.assess_report_freshness(
        state,
        current_version=current_code_version,
    )
    return freshness.warning


def _operator_report_freshness_warning(
    state: ProjectState,
    *,
    current_code_version: str | None = None,
) -> str:
    freshness = report_freshness.assess_report_freshness(
        state,
        current_version=current_code_version,
    )
    metadata_rows = [
        ["Field", "Value"],
        ["Freshness status", freshness.status],
        ["Generated code version", freshness.generated_code_version or "unavailable"],
        ["Current code version", freshness.current_code_version or "unavailable"],
        ["Generated at", freshness.generated_at or "unavailable"],
        ["Matching report hash", str(bool(freshness.matching_report_hash))],
    ]
    parts = []
    if freshness.warning:
        parts.append(freshness.warning)
    parts.append("Freshness metadata:\n\n" + _markdown_table(metadata_rows))
    return "\n\n".join(parts)


def _simplify_sparse_report_markdown(markdown: str, quality) -> str:
    lines: list[str] = []
    in_technical_appendix = False
    for raw_line in str(markdown or "").splitlines():
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", raw_line.strip())
        if heading:
            in_technical_appendix = _is_technical_appendix_heading(heading.group(2))
            lines.append(raw_line)
            continue
        if in_technical_appendix:
            lines.append(raw_line)
        else:
            lines.append(_client_safe_text(raw_line, quality))
    value = _collapse_markdown_blank_lines("\n".join(lines))
    if SPARSE_CONFIDENCE_RULE not in value:
        value = "\n\n".join([f"**Sparse confidence note:** {SPARSE_CONFIDENCE_RULE}", value])
    return value


def _collapse_markdown_blank_lines(value: str) -> str:
    lines = [line.rstrip() for line in str(value or "").splitlines()]
    collapsed: list[str] = []
    blank = False
    for line in lines:
        if not line.strip():
            if not blank:
                collapsed.append("")
            blank = True
        else:
            collapsed.append(line)
            blank = False
    return "\n".join(collapsed).strip()


def _ensure_single_client_source_locator_note(markdown: str, quality) -> str:
    source = str(markdown or "")
    note_pattern = re.compile(
        r"(?:\n{0,2}#{1,6}\s+(?:Citation|Source)\s+locator\s+note\s*\n+)?"
        r"No concrete (?:citation|source) locators were available for this project; "
        r"evidence should be validated in Sprint 0\.",
        re.I,
    )
    cleaned = note_pattern.sub("\n", source)
    cleaned = _collapse_markdown_blank_lines(cleaned)
    if getattr(quality, "has_concrete_locators", False):
        return cleaned

    note_block = f"## Source locator note\n\n{NO_CONCRETE_LOCATORS_CLIENT_NOTE}"
    lines = cleaned.splitlines()
    insert_at = 0
    if lines and lines[0].startswith("# "):
        insert_at = 1
        if len(lines) > 1 and lines[1].strip() and not lines[1].startswith("#"):
            insert_at = 2
    lines[insert_at:insert_at] = ["", note_block, ""]
    return _collapse_markdown_blank_lines("\n".join(lines))


def _client_safe_text(text: str, quality) -> str:
    protected_text, metric_fragments = _protect_client_concrete_metric_values(str(text or ""))
    value = client_simplify_text(protected_text, sparse_evidence=quality.sparse_evidence or quality.evidence_warning)
    if quality.decision_domain == "growth":
        replacements = {
            "Search Console": "growth analytics",
            "GA4": "product/revenue analytics",
            "crawl/technical evidence": "operating evidence",
            "crawl": "diagnostic review",
            "editorial evidence": "customer/revenue evidence",
            "editorial workflow evidence": "customer/revenue workflow evidence",
            "CMS/schema capability": "operating capability",
            "SEO Lead": "Growth Lead",
            "Editorial Lead": "Growth Lead",
            "Web/CMS Owner": "Product Analytics Lead",
        }
        for unsafe, safe in replacements.items():
            value = re.sub(re.escape(unsafe), safe, value, flags=re.I)
    if quality.decision_domain == "productization" and not any("CMS" in item for item in quality.evidence_categories):
        replacements = {
            "CMS/schema capability": "template schema / field registry validation",
            "CMS or schema capability": "reusable template schema, field registry, or product instrumentation",
            "CMS": "template system",
        }
        for unsafe, safe in replacements.items():
            value = re.sub(re.escape(unsafe), safe, value, flags=re.I)
    return _polish_client_report_citation_rendering(_restore_client_concrete_metric_values(value, metric_fragments))


def _remove_empty_client_citation_marker_columns(markdown: str) -> str:
    source = str(markdown or "")
    lines = source.splitlines()
    output: list[str] = []
    index = 0
    while index < len(lines):
        if lines[index].lstrip().startswith("|"):
            block: list[str] = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                block.append(lines[index])
                index += 1
            output.extend(_clean_client_citation_table_block(block))
            continue
        output.append(lines[index])
        index += 1
    return "\n".join(output)


def _clean_client_citation_table_block(block: list[str]) -> list[str]:
    if len(block) < 2:
        return block
    rows = [_split_markdown_table_row(line) for line in block]
    if not rows or not rows[0] or not _is_client_markdown_separator_cells(rows[1]):
        return block
    width = len(rows[0])
    candidate_indexes = [
        index
        for index, header in enumerate(rows[0])
        if _is_client_citation_column_header(header)
    ]
    if not candidate_indexes:
        return block
    removable = {
        index
        for index in candidate_indexes
        if all(
            _is_empty_client_citation_cell((row + [""] * width)[index])
            for row in rows[2:]
        )
    }
    if not removable:
        return block

    cleaned: list[str] = []
    for row_index, row in enumerate(rows):
        padded = (row + [""] * width)[:width]
        kept = [cell for index, cell in enumerate(padded) if index not in removable]
        if not kept:
            continue
        if row_index == 1:
            kept = ["---"] * len(kept)
        if row_index > 1 and not any(cell.strip() for cell in kept):
            continue
        cleaned.append(_format_markdown_table_row(kept))
    return cleaned


def _split_markdown_table_row(line: str) -> list[str]:
    text = line.strip()
    if not text.startswith("|"):
        return []
    if text.endswith("|"):
        text = text[1:-1]
    else:
        text = text[1:]
    return [cell.strip() for cell in text.split("|")]


def _format_markdown_table_row(cells: list[str]) -> str:
    return "| " + " | ".join(cells) + " |"


def _is_client_markdown_separator_cells(cells: list[str]) -> bool:
    return bool(cells) and all(
        re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells
    )


def _is_client_citation_column_header(value: str) -> bool:
    normalized = _normalize_heading(value)
    return normalized in {
        "citation",
        "citation marker",
        "citation markers",
        "citation status",
        "citation statuses",
        "citation locator",
        "citation locators",
        "locator",
        "locators",
    } or ("citation" in normalized and ("marker" in normalized or "locator" in normalized))


def _is_empty_client_citation_cell(value: str) -> bool:
    text = re.sub(r"<br\s*/?>", " ", str(value or ""), flags=re.I).strip()
    text = text.strip("`*_ ")
    if not text:
        return True
    normalized = _normalize_heading(text)
    return normalized in {
        "n a",
        "na",
        "none",
        "unknown",
        "unavailable",
        "not available",
        "citation unavailable",
        "no citation available",
        "no concrete locator registered",
        "no concrete citation locator registered",
        "no concrete citation locators available",
        "evidence source unavailable",
        "no locator",
        "no locators",
    } or bool(re.fullmatch(r"[-]+", text))


def _is_technical_appendix_heading(value: str) -> bool:
    normalized = _normalize_heading(value)
    return normalized in {
        "appendix technical analysis",
        "technical appendix",
        "appendix technical",
    } or ("technical" in normalized and "appendix" in normalized)


def _project_metadata_line(state: ProjectState) -> str:
    return " | ".join(
        part
        for part in [
            f"Project ID: {state.project_id}",
            f"Project name: {_redact_unsafe_string(state.project_name or '')}",
            f"Generated: {_utc_now()}",
            f"Risk: {_redact_unsafe_string(state.risk_classification or '')}",
        ]
        if part.strip()
    )


def _client_project_metadata_line(state: ProjectState) -> str:
    return " | ".join(
        part
        for part in [
            f"Project name: {_redact_unsafe_string(state.project_name or '')}",
            f"Generated: {_utc_now()}",
        ]
        if part.strip()
    )


def _extract_report_sections(report: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current_key = ""
    current_level = 0
    for raw_line in (report or "").splitlines():
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", raw_line.strip())
        if match:
            heading = match.group(2).strip()
            key = _normalize_heading(heading)
            if key in {_normalize_heading(item) for item in REPORT_CLARITY_HEADINGS}:
                current_key = key
                current_level = len(match.group(1))
                sections.setdefault(current_key, [])
                continue
            if current_key and len(match.group(1)) <= current_level:
                current_key = ""
                current_level = 0
        if current_key:
            sections[current_key].append(raw_line)
    return {key: _redact_unsafe_string("\n".join(lines).strip()) for key, lines in sections.items()}


def _section_or_fallback(sections: dict[str, str], heading: str) -> str:
    value = sections.get(_normalize_heading(heading), "").strip()
    return value or "Not available in current report output."


def _normalize_heading(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", (value or "").lower()).strip()


def _citation_locator_summary_markdown(state: ProjectState, *, include_registry: bool) -> str:
    try:
        from cdp.citation_resolvability import build_defense_pass_result

        result = build_defense_pass_result(state)
    except Exception:
        return ""

    has_summary = bool(result.registry_entries or result.markers or result.summary_counts)
    if not has_summary:
        return ""

    lines = [
        "Citation locator review summary — confirms citation marker resolution only, not semantic support or concrete source locator availability.",
    ]
    if result.summary_counts:
        rows = [["Metric", "Count"]]
        for key in (
            "canonical_marker_count",
            "resolved_exact",
            "resolved_id_only",
            "unknown_evidence_id",
            "locator_mismatch",
            "malformed",
            "load_bearing_review_count",
        ):
            if key in result.summary_counts:
                rows.append([key, str(result.summary_counts.get(key, 0))])
        if len(rows) > 1:
            lines.append(_markdown_table(rows))

    if include_registry and result.registry_entries:
        rows = [["Evidence ID", "Locator availability", "Title"]]
        for entry in result.registry_entries[:12]:
            locators = ", ".join(entry.locators or [])
            rows.append([
                entry.evidence_id,
                _redact_unsafe_string(locators or "No concrete locator registered"),
                _redact_unsafe_string(", ".join(entry.titles or [])),
            ])
        lines.append(_markdown_table(rows))
    return "\n\n".join(lines)


def _evidence_locator_register_markdown(state: ProjectState) -> str:
    entries = _evidence_locator_payload(state)
    if not entries:
        return "No evidence locator registry entries available."
    rows = [["Evidence ID", "Locators", "Source", "Title"]]
    for entry in entries[:30]:
        rows.append([
            entry.get("evidence_id", ""),
            ", ".join(entry.get("locators", []) or []) or "No concrete locator registered",
            entry.get("source_id", ""),
            entry.get("title", ""),
        ])
    return _markdown_table(rows)


def _client_clarifications_markdown(state: ProjectState) -> str:
    questions = []
    for cycle in state.clarification_cycles or []:
        for question in cycle.questions or []:
            status = _enum_value(question.status)
            if status in {"open", "unavailable"}:
                questions.append([
                    question.text,
                    question.why_it_matters,
                    question.affected_phase,
                    status,
                ])
    if not questions:
        return ""
    return _markdown_table([["Open question", "Why it matters", "Affected phase", "Status"], *questions])


def _operator_clarifications_markdown(state: ProjectState) -> str:
    rows = [["Question ID", "Question", "Affected phase", "Status", "Answer"]]
    answers = {answer.question_id: answer for answer in state.clarification_answers or []}
    for cycle in state.clarification_cycles or []:
        for question in cycle.questions or []:
            answer = answers.get(question.question_id)
            rows.append([
                question.question_id,
                question.text,
                question.affected_phase,
                _enum_value(question.status),
                answer.answer_text if answer else "",
            ])
    return _markdown_table(rows) if len(rows) > 1 else "No clarification questions or answers saved."


def _operator_overview_markdown(state: ProjectState) -> str:
    try:
        from overview import build_operator_overview

        overview = build_operator_overview(state)
        rows = [
            ["Field", "Value"],
            ["Project status", overview.project_status],
            ["Current recommendation", _short_text(overview.current_recommendation, 260)],
            ["Decision summary", _short_text(overview.decision_summary, 260)],
            ["Sources/files", _short_text(overview.sources_and_files_message, 260)],
            ["Next operator action", _short_text(overview.next_operator_action, 260)],
        ]
        return _markdown_table(rows)
    except Exception as exc:
        return f"Operator overview unavailable: {_short_text(str(exc), 180)}"


def _workspace_summary_markdown(state: ProjectState) -> str:
    try:
        from workspace import build_workspace_summary

        workspace = build_workspace_summary(state)
        rows = [
            ["Field", "Value"],
            ["Project status", workspace.project_status],
            ["Current phase", workspace.current_phase],
            ["Active risks", str(workspace.active_risk_count)],
            ["Decision object status", workspace.decision_object_health.status],
            ["Knowledge status", workspace.knowledge_health.status],
            ["SQI", _fmt_value(workspace.score_summary.sqi_overall)],
            ["Deterministic score", _fmt_value(workspace.score_summary.det_score_overall)],
            ["Brier score", _fmt_value(workspace.score_summary.brier_score)],
        ]
        if workspace.blocking_reasons:
            rows.append(["Blocking reasons", "; ".join(workspace.blocking_reasons)])
        return _markdown_table(rows)
    except Exception as exc:
        return f"Workspace summary unavailable: {_short_text(str(exc), 180)}"


def _summarize_gauntlet(state: ProjectState) -> str:
    if not state.gauntlet:
        return "No gauntlet output saved."
    rows = [["ID", "Risk rank", "Crux", "Top FMEA", "FTA cut set"]]
    for result in state.gauntlet.results or []:
        rows.append([
            result.id,
            _fmt_value(result.risk_rank),
            result.crux,
            str(result.top_fmea),
            str(result.fta_cut_set),
        ])
    return "\n\n".join(
        [
            f"Portfolio correlation: {_fmt_value(state.gauntlet.portfolio_correlation)}",
            f"MECE gaps: {state.gauntlet.mece_gaps or 'TBD — requires operator confirmation.'}",
            _markdown_table(rows),
        ]
    )


def _summarize_audit(state: ProjectState) -> str:
    if not state.audit:
        return _redact_unsafe_string(state.audit_raw or "No audit output saved.")
    rows = [["Type", "Summary"]]
    rows.extend([["Top finding", item] for item in state.audit.top_findings or []])
    rows.extend([["Observation need", item] for item in state.audit.observation_needs or []])
    for item in state.audit.fmea or []:
        rows.append(["FMEA", f"{item.component}: {item.failure_mode} | RPN {item.rpn} | action {item.action}"])
    return _markdown_table(rows)


def _summarize_strategy(state: ProjectState) -> str:
    if not state.strategy:
        return _redact_unsafe_string(state.strategy_raw or "No strategy output saved.")
    rows = [["Priority", "Action", "Expected impact", "Timeline", "Risk if ignored"]]
    for action in state.strategy.strategies or []:
        rows.append([
            _enum_value(action.priority),
            action.action,
            action.expected_impact,
            action.timeline,
            action.risk_if_ignored,
        ])
    metrics = "\n".join(f"- {metric}" for metric in state.strategy.success_metrics or [])
    if not metrics and monitor_has_signals(state.monitor):
        monitor_lines = "\n".join(f"- {line}" for line in monitor_success_metric_lines(state.monitor))
        metrics = "Success metrics are captured in the monitoring plan below."
        if monitor_lines:
            metrics = f"{metrics}\n{monitor_lines}"
    return "\n\n".join([
        state.strategy.executive_strategy or "No executive strategy saved.",
        _markdown_table(rows),
        "### Success Metrics",
        metrics or "No success metrics saved.",
    ])


def _budget_summary_markdown(state: ProjectState) -> str:
    rows = [["Budget field", "Value"]]
    for key, value in (state.budget_caps or {}).items():
        rows.append([f"cap.{key}", _fmt_value(value)])
    for key, value in (state.budget_consumed or {}).items():
        rows.append([f"consumed.{key}", _fmt_value(value)])
    return _markdown_table(rows)


def _approvals_summary_markdown(state: ProjectState) -> str:
    approvals = _approvals_payload(state)
    if not approvals:
        return "No approvals saved."
    rows = [["Scope", "Status", "Requested by", "Resolved by", "Reason"]]
    for approval in approvals:
        rows.append([
            approval.get("scope", ""),
            approval.get("status", ""),
            approval.get("requested_by", ""),
            approval.get("resolved_by", ""),
            approval.get("reason", ""),
        ])
    return _markdown_table(rows)


def _calibration_prediction_summary_markdown(state: ProjectState) -> str:
    payload = _calibration_predictions_payload(state)
    if not payload:
        return "No calibration or prediction records saved."
    rows = [["Type", "ID/Phase", "Probability/Score", "Outcome/Notes"]]
    for prediction in payload.get("predictions", []):
        rows.append([
            "Prediction",
            f"{prediction.get('hypothesis_id', '')}/{prediction.get('phase', '')}",
            _fmt_value(prediction.get("predicted_probability")),
            _fmt_value(prediction.get("actual_outcome")),
        ])
    for snapshot in payload.get("calibration_snapshots", []):
        rows.append([
            "Calibration",
            snapshot.get("snapshot_id", ""),
            _fmt_value(snapshot.get("brier_score")),
            snapshot.get("notes", ""),
        ])
    return _markdown_table(rows)


def _phase_outputs_payload(state: ProjectState) -> dict[str, Any]:
    return {
        "classify": _model_dump(state.classify),
        "hypotheses": _model_dump(state.hypotheses),
        "gauntlet": _model_dump(state.gauntlet),
        "audit": _model_dump(state.audit),
        "audit_raw": state.audit_raw,
        "strategy": _model_dump(state.strategy),
        "strategy_raw": state.strategy_raw,
        "sqi": _model_dump(state.sqi),
        "monitor": _model_dump(state.monitor),
        "phase_summaries": dict(state.phase_summaries or {}),
        "phase_status": _model_dump(state.phase_status),
        "phase_confidence": dict(state.phase_confidence or {}),
        "phase_run_completed_at": dict(state.phase_run_completed_at or {}),
    }


def _decision_objects_payload(state: ProjectState) -> dict[str, Any]:
    if state.decision_objects is None:
        try:
            from decision_objects import ensure_decision_objects

            return _model_dump(ensure_decision_objects(state, trigger="export_profile"))
        except Exception:
            return {}
    return _model_dump(state.decision_objects)


def _clarifications_payload(state: ProjectState) -> dict[str, Any]:
    return {
        "cycles": _model_dump(state.clarification_cycles),
        "answers": _model_dump(state.clarification_answers),
    }


def _evidence_locator_payload(state: ProjectState) -> list[dict[str, Any]]:
    try:
        from cdp.citation_resolvability import build_evidence_locator_registry

        return [_model_dump(entry) for entry in build_evidence_locator_registry(state)]
    except Exception:
        return []


def _uploaded_file_manifest_payload(state: ProjectState) -> list[dict[str, Any]]:
    manifests = list(getattr(getattr(state, "knowledge_layer", None), "uploaded_files", []) or [])
    rows = []
    for manifest in manifests:
        storage_ref = getattr(manifest, "storage_ref", "")
        rows.append(
            {
                "file_id": getattr(manifest, "file_id", ""),
                "source_id": getattr(manifest, "source_id", ""),
                "original_filename": getattr(manifest, "filename", ""),
                "content_type": getattr(manifest, "media_type", ""),
                "size_bytes": getattr(manifest, "size_bytes", 0),
                "evidence_id": getattr(manifest, "evidence_id", ""),
                "uploaded_at": getattr(manifest, "uploaded_at", ""),
                "uploaded_by": getattr(manifest, "uploaded_by", ""),
                "parser_kind": getattr(manifest, "parser_kind", ""),
                "import_mode": getattr(manifest, "import_mode", ""),
                "storage_ref": REDACTED if _looks_like_unsafe_path(str(storage_ref)) else storage_ref,
                "parse_summary": _model_dump(getattr(manifest, "parse_summary", None)),
            }
        )
    return rows


def _policy_summary_payload(state: ProjectState) -> dict[str, Any]:
    event_counts: dict[str, int] = {}
    for event in state.policy_audit_log or []:
        event_type = str((event or {}).get("event_type", "unknown"))
        event_counts[event_type] = event_counts.get(event_type, 0) + 1
    return {
        "risk_classification": state.risk_classification,
        "risk_classification_rationale": state.risk_classification_rationale,
        "risk_classification_set_by": state.risk_classification_set_by,
        "kill_switch_active": state.kill_switch_active,
        "kill_switch_reason": state.kill_switch_reason,
        "kill_switch_triggered_by": state.kill_switch_triggered_by,
        "kill_switch_triggered_at": state.kill_switch_triggered_at,
        "phase_breaker_count": len(state.phase_breakers or {}),
        "policy_audit_event_count": len(state.policy_audit_log or []),
        "policy_audit_event_types": event_counts,
        "intake_sanitization_summary": _sanitization_summary(state.intake_sanitization_findings),
    }


def _calibration_predictions_payload(state: ProjectState) -> dict[str, Any]:
    payload = {
        "brier_score": state.brier_score,
        "predictions": _model_dump(state.predictions),
        "calibration_snapshots": _model_dump(
            getattr(getattr(state, "decision_objects", None), "calibration_snapshots", [])
            or []
        ),
    }
    return {key: value for key, value in payload.items() if value not in (None, [], {})}


def _approvals_payload(state: ProjectState) -> list[dict[str, Any]]:
    approvals = []
    decision_objects = state.decision_objects
    for approval in list(getattr(decision_objects, "approvals", []) or []):
        approvals.append(_model_dump(approval))
    for action, approval in (state.approvals_granted or {}).items():
        row = _model_dump(approval)
        if isinstance(row, dict):
            row.setdefault("scope", action)
            approvals.append(row)
    return approvals


def _risk_summary_payload(state: ProjectState) -> dict[str, Any]:
    risks = list(getattr(getattr(state, "decision_objects", None), "risks", []) or [])
    return {
        "risk_classification": state.risk_classification,
        "risk_classification_rationale": state.risk_classification_rationale,
        "risks": _model_dump(risks),
    }


def _sanitization_summary(findings: dict | None) -> dict[str, Any]:
    if not findings:
        return {}
    return {
        "highest_severity": findings.get("highest_severity"),
        "recommendation": findings.get("recommendation"),
        "brief_length": findings.get("brief_length"),
        "truncated": findings.get("truncated"),
        "finding_count_by_severity": findings.get("finding_count_by_severity", {}),
    }


def _zip_archive_bytes(files: dict[str, Any]) -> bytes:
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for filename in sorted(files.keys()):
            if filename == "raw_project_state.json":
                continue
            payload = files[filename]
            if filename.endswith(".json"):
                content = json.dumps(
                    sanitize_for_export(payload, "machine_archive", mode="redact"),
                    ensure_ascii=False,
                    indent=2,
                    sort_keys=True,
                    default=str,
                )
            else:
                content = _redact_unsafe_string(str(payload or ""))
            archive.writestr(filename, content.encode("utf-8"))
    return buf.getvalue()


def _export_markdown_docx_bytes(markdown: str, *, title: str) -> bytes:
    document = Document()
    document.core_properties.title = title
    document.core_properties.subject = "Decision Engine profile export"
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    for block in _markdown_to_blocks(markdown):
        if block["type"] == "heading":
            document.add_heading(block["text"], level=min(max(block["level"], 1), 4))
        elif block["type"] == "paragraph":
            document.add_paragraph(_strip_inline_markdown(block["text"]))
        elif block["type"] == "bullets":
            for item in block["items"]:
                document.add_paragraph(_strip_inline_markdown(item), style="List Bullet")
        elif block["type"] == "numbered":
            for item in block["items"]:
                document.add_paragraph(_strip_inline_markdown(item), style="List Number")
        elif block["type"] == "table":
            _add_docx_table(document, block["rows"])
        elif block["type"] == "divider":
            _add_docx_divider(document)

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def _export_markdown_pdf_bytes(markdown: str, *, title: str, client_visible: bool = False) -> bytes:
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=title,
    )
    styles = getSampleStyleSheet()
    heading1 = ParagraphStyle("ProfileHeading1", parent=styles["Heading1"], fontSize=17, leading=21, spaceAfter=8)
    heading2 = ParagraphStyle("ProfileHeading2", parent=styles["Heading2"], fontSize=13, leading=17, spaceBefore=8, spaceAfter=6)
    heading3 = ParagraphStyle("ProfileHeading3", parent=styles["Heading3"], fontSize=11, leading=14, spaceBefore=6, spaceAfter=4)
    body = ParagraphStyle("ProfileBody", parent=styles["BodyText"], fontSize=9.5, leading=13, spaceAfter=4)

    story = []
    render_markdown = _polish_client_report_citation_rendering(markdown) if client_visible else markdown
    for block in _markdown_to_blocks(render_markdown):
        if block["type"] == "heading":
            style = heading1 if block["level"] == 1 else heading2 if block["level"] == 2 else heading3
            story.append(Paragraph(_as_pdf_text(_strip_inline_markdown(block["text"])), style))
        elif block["type"] == "paragraph":
            story.append(Paragraph(_as_pdf_text(_strip_inline_markdown(block["text"])), body))
        elif block["type"] == "bullets":
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(_as_pdf_text(_strip_inline_markdown(item)), body), leftIndent=10) for item in block["items"]],
                    bulletType="bullet",
                    leftIndent=12,
                )
            )
        elif block["type"] == "numbered":
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(_as_pdf_text(_strip_inline_markdown(item)), body), leftIndent=10) for item in block["items"]],
                    bulletType="1",
                    leftIndent=12,
                )
            )
        elif block["type"] == "table":
            story.extend(
                _pdf_table_flowables(
                    block["rows"],
                    body,
                    doc.width,
                    client_visible=client_visible,
                )
            )
        elif block["type"] == "divider":
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CBD5E1"), spaceBefore=4, spaceAfter=4))
        story.append(Spacer(1, 2))
    doc.build(story)
    return buf.getvalue()


def _markdown_to_blocks(markdown: str) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    lines = (markdown or "").splitlines()
    idx = 0
    bullets: list[str] = []
    numbered: list[str] = []

    def flush_lists():
        nonlocal bullets, numbered
        if bullets:
            blocks.append({"type": "bullets", "items": bullets})
            bullets = []
        if numbered:
            blocks.append({"type": "numbered", "items": numbered})
            numbered = []

    while idx < len(lines):
        line = _normalize_markdown_line(lines[idx]).strip()
        if not line:
            flush_lists()
            idx += 1
            continue
        table = _consume_table(lines, idx)
        if table:
            flush_lists()
            rows, next_idx = table
            blocks.append({"type": "table", "rows": rows})
            idx = next_idx
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush_lists()
            blocks.append({"type": "heading", "level": len(heading.group(1)), "text": heading.group(2).strip()})
        elif _is_markdown_divider(line):
            flush_lists()
            blocks.append({"type": "divider"})
        elif re.match(r"^(-|\*|•)\s+", line):
            if numbered:
                flush_lists()
            bullets.append(re.sub(r"^(-|\*|•)\s+", "", line).strip())
        elif re.match(r"^\d+\.\s+", line):
            if bullets:
                flush_lists()
            numbered.append(re.sub(r"^\d+\.\s+", "", line).strip())
        else:
            flush_lists()
            blocks.append({"type": "paragraph", "text": line})
        idx += 1
    flush_lists()
    return blocks


def _consume_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines):
        return None
    first = _normalize_markdown_line(lines[start]).strip()
    second = _normalize_markdown_line(lines[start + 1]).strip()
    if "|" not in first or not _is_markdown_separator_row(second):
        return None
    rows = [_split_table_row(first)]
    idx = start + 2
    while idx < len(lines):
        line = _normalize_markdown_line(lines[idx]).strip()
        if "|" not in line:
            break
        rows.append(_split_table_row(line))
        idx += 1
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    return rows, idx


def _normalize_markdown_line(line: str) -> str:
    value = str(line or "").strip()
    if value == ">":
        return ""
    if value.startswith(">"):
        return value[1:].lstrip()
    return value


def _is_markdown_divider(line: str) -> bool:
    return bool(re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", (line or "").strip()))


def _is_markdown_separator_row(line: str) -> bool:
    cells = _split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [_strip_inline_markdown(cell.strip()) for cell in stripped.split("|")]


def _add_docx_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            table.cell(row_idx, col_idx).text = _strip_inline_markdown(str(value))


def _add_docx_divider(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)


def _pdf_table_flowables(
    rows: list[list[str]],
    body_style: ParagraphStyle,
    available_width: float,
    *,
    client_visible: bool = False,
) -> list[Any]:
    column_count = max(len(row) for row in rows) if rows else 1
    if column_count > 4:
        return _pdf_wide_table_cards(rows, body_style, available_width, client_visible=client_visible)
    return [_pdf_table(rows, body_style, available_width)]


def _pdf_table(rows: list[list[str]], body_style: ParagraphStyle, available_width: float) -> Table:
    column_count = max(len(row) for row in rows) if rows else 1
    col_widths = [available_width / column_count] * column_count
    table = Table(
        [[Paragraph(_as_pdf_text(str(cell)), body_style) for cell in row] for row in rows],
        colWidths=col_widths,
        repeatRows=1,
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E5E7EB")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#111827")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def _pdf_wide_table_cards(
    rows: list[list[str]],
    body_style: ParagraphStyle,
    available_width: float,
    *,
    client_visible: bool = False,
) -> list[Any]:
    if not rows:
        return []
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    headers = normalized[0]
    cards: list[Any] = []
    label_style = ParagraphStyle(
        "WideTableCardLabel",
        parent=body_style,
        fontName="Helvetica-Bold",
        fontSize=max(body_style.fontSize - 0.5, 7),
        leading=body_style.leading,
    )
    for row_index, row in enumerate(normalized[1:], start=1):
        card_rows = []
        for header, value in zip(headers, row):
            if not str(header).strip() and not str(value).strip():
                continue
            if client_visible and _is_client_citation_column_header(str(header)) and not str(value).strip():
                continue
            card_rows.append([
                Paragraph(_as_pdf_text(_strip_inline_markdown(str(header or f"Field {row_index}"))), label_style),
                Paragraph(_as_pdf_text(_strip_inline_markdown(str(value))), body_style),
            ])
        if not card_rows:
            continue
        card = Table(card_rows, colWidths=[available_width * 0.28, available_width * 0.72])
        card.setStyle(
            TableStyle(
                [
                    ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#E5E7EB")),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F8FAFC")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        cards.extend([card, Spacer(1, 5)])
    return cards


def _markdown_table(rows: list[list[Any]]) -> str:
    if not rows:
        return ""
    clean_rows = [[_table_cell(value) for value in row] for row in rows]
    width = max(len(row) for row in clean_rows)
    clean_rows = [row + [""] * (width - len(row)) for row in clean_rows]
    header = clean_rows[0]
    separator = ["---"] * width
    body = clean_rows[1:]
    return "\n".join(
        ["| " + " | ".join(header) + " |", "| " + " | ".join(separator) + " |"]
        + ["| " + " | ".join(row) + " |" for row in body]
    )


def _table_cell(value: Any) -> str:
    return _redact_unsafe_string(str(value if value is not None else "")).replace("\n", " ").replace("|", "/").strip()


def _strip_inline_markdown(text: str) -> str:
    value = str(text or "")
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = value.replace("**", "").replace("__", "")
    return value.strip()


def _spell_visible_comparators(text: str) -> str:
    value = str(text or "")
    value = re.sub(r"(?<![-=])\s*>=\s*", " at least ", value)
    value = re.sub(r"(?<![-])\s*>\s*", " greater than ", value)
    return re.sub(r"\s{2,}", " ", value).strip()


def _sanitize_for_export(value, *, profile: str, mode: str):
    if hasattr(value, "model_dump"):
        value = value.model_dump(mode="json")
    elif isinstance(value, datetime):
        return value.isoformat()
    elif hasattr(value, "value") and not isinstance(value, (str, bytes)):
        return getattr(value, "value")

    if isinstance(value, dict):
        result = {}
        for key, item in value.items():
            key_text = str(key)
            if _is_sensitive_key(key_text):
                if mode == "drop":
                    continue
                result[key_text] = REDACTED
                continue
            sanitized = _sanitize_for_export(item, profile=profile, mode=mode)
            if sanitized is not _DROP:
                result[key_text] = sanitized
        return result
    if isinstance(value, (list, tuple, set)):
        items = []
        for item in value:
            sanitized = _sanitize_for_export(item, profile=profile, mode=mode)
            if sanitized is not _DROP:
                items.append(sanitized)
        return items
    if isinstance(value, str):
        return _redact_unsafe_string(value)
    return value


def _is_sensitive_key(key: str) -> bool:
    normalized = re.sub(r"[^a-z0-9]+", "_", key.lower())
    return any(fragment in normalized for fragment in SENSITIVE_KEY_SUBSTRINGS)


def _redact_unsafe_string(value: str) -> str:
    text = str(value or "")
    text = re.sub(r"(?i)(api[_-]?key|token|password|secret|credential)\s*[:=]\s*[^,\s|;]+", r"\1=" + REDACTED, text)
    text = re.sub(r"(?i)\bBearer\s+[A-Za-z0-9._~+/=-]+", "Bearer " + REDACTED, text)
    text = re.sub(r"\bsk-[A-Za-z0-9_\-]{8,}\b", REDACTED, text)
    text = re.sub(r"(?i)\b[A-Z]:[\\/][^\s|]+", REDACTED, text)
    text = re.sub(r"\\\\[A-Za-z0-9_.-]+\\[^\s|]+", REDACTED, text)
    text = re.sub(r"file://[^\s|]+", REDACTED, text)
    text = re.sub(r"(?<!https:)(?<!http:)\b/(Users|home|mnt|var|tmp|root|workspace|Volumes|opt)/[^\s|]+", REDACTED, text)
    return text


def _looks_like_unsafe_path(value: str) -> bool:
    return _redact_unsafe_string(value) != str(value or "")


def _model_dump(value):
    if hasattr(value, "model_dump"):
        return value.model_dump(mode="json")
    if isinstance(value, list):
        return [_model_dump(item) for item in value]
    if isinstance(value, dict):
        return {str(key): _model_dump(item) for key, item in value.items()}
    if hasattr(value, "value"):
        return getattr(value, "value")
    return value


def _phase_output_text(state: ProjectState, phase: str) -> str:
    if phase == "audit" and state.audit_raw and not state.audit:
        return state.audit_raw
    if phase == "strategy" and state.strategy_raw and not state.strategy:
        return state.strategy_raw
    if phase == "report":
        return state.report or ""
    return json.dumps(_model_dump(getattr(state, phase, None)), ensure_ascii=False, default=str)


def _short_text(value: str, limit: int = 180) -> str:
    text = _redact_unsafe_string(str(value or "")).replace("\n", " ").strip()
    return text if len(text) <= limit else text[: limit - 1].rstrip() + "…"


def _fmt_value(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        return f"{value:.4g}"
    return _redact_unsafe_string(str(_enum_value(value)))


def _enum_value(value):
    return value.value if hasattr(value, "value") else value


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def _code_version() -> str:
    return report_freshness.current_code_version()
