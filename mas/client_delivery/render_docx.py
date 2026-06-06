"""DOCX renderer for the Client Delivery Generator v0.5."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from .models import DeliveryPackage
from .utils import display_join, display_text


HUMAN_REVIEW_OPERATOR_NOTE = (
    "This artifact is operator-generated and requires human review before client delivery."
)
HUMAN_REVIEW_MEMO_NOTE = (
    "This memo is a decision-support artifact. A human operator should review evidence, "
    "assumptions, risks, and recommendations before client delivery."
)


def render_board_memo_docx(package: DeliveryPackage, path) -> Path:
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.core_properties.title = "Strategic Decision Board Memo"
    document.core_properties.subject = "Client Delivery Generator v0.5"
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    document.add_heading("Board Memo", 0)

    document.add_heading("Decision", level=1)
    document.add_paragraph(display_text(package.decision_statement, "Decision statement unavailable."))
    document.add_paragraph(f"Project ID: {display_text(package.project_id)}")

    document.add_heading("Recommendation", level=1)
    recommendation_rows = [
        ["Selected option", package.recommendation.selected_option],
        ["Rationale", package.recommendation.rationale],
        ["Confidence", package.recommendation.confidence],
        ["Evidence strength", package.recommendation.evidence_strength],
        ["Evidence", display_join(package.recommendation.evidence)],
    ]
    _add_key_value_table(document, recommendation_rows)

    document.add_heading("Critical Assumptions", level=1)
    _add_table(
        document,
        ["Assumption", "Falsification trigger", "Confidence", "Evidence"],
        [
            [
                assumption.assumption,
                assumption.falsification_trigger,
                assumption.confidence,
                display_join(assumption.evidence),
            ]
            for assumption in package.critical_assumptions
        ],
        empty_text="No critical assumptions were extracted.",
    )

    document.add_heading("30/60/90 Execution Plan", level=1)
    _add_table(
        document,
        ["Phase", "Action", "Owner", "Dependencies", "Evidence", "Notes"],
        [
            [
                action.phase,
                action.action,
                action.owner,
                display_join(action.dependencies),
                display_join(action.evidence),
                action.notes,
            ]
            for action in package.execution_plan
        ],
        empty_text="No execution actions were extracted.",
    )

    document.add_heading("KPIs & Review Triggers", level=1)
    _add_table(
        document,
        ["KPI", "Type", "Red threshold", "Amber threshold", "Actual", "Status", "Owner"],
        [
            [
                kpi.name,
                kpi.indicator_type,
                kpi.threshold_red,
                kpi.threshold_amber,
                kpi.actual_value,
                kpi.status,
                kpi.owner,
            ]
            for kpi in package.kpis
        ],
        empty_text="No KPIs were extracted.",
    )
    document.add_paragraph("Review triggers:")
    for trigger in package.review.reentry_triggers:
        document.add_paragraph(display_text(trigger), style="List Bullet")

    document.add_heading("Human Review Required", level=1)
    document.add_paragraph(HUMAN_REVIEW_OPERATOR_NOTE)
    document.add_paragraph(HUMAN_REVIEW_MEMO_NOTE)

    document.save(output_path)
    return output_path


def _add_key_value_table(document: Document, rows: list[list[object]]) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        table.cell(row_idx, 0).text = display_text(row[0])
        table.cell(row_idx, 1).text = display_text(row[1])


def _add_table(
    document: Document,
    headers: list[str],
    rows: list[list[object]],
    *,
    empty_text: str,
) -> None:
    if not rows:
        document.add_paragraph(empty_text)
        return
    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = display_text(header)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            table.cell(row_idx, col_idx).text = display_text(value)
