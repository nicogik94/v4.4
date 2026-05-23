from __future__ import annotations

import asyncio
import io
import json
import re
import sys
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from config import MODEL_ROUTING  # noqa: E402
from exporters import export_project_docx_bytes  # noqa: E402
from extensions.runtime import GatewayRequest, RoutingContext  # noqa: E402
from knowledge.registry import ensure_knowledge_layer, upsert_source_entry  # noqa: E402
from knowledge.sync import sync_offline_source  # noqa: E402
from llm_client import _get_provider_gateway  # noqa: E402
from orchestrator import build_report_prompt, build_system_prompt  # noqa: E402
from state import (  # noqa: E402
    KnowledgeItem,
    KnowledgeLayerState,
    PhaseStatus,
    ProjectState,
    Provenance,
    SourceRegistryEntry,
)
from tests.test_decision_objects import make_state  # noqa: E402
from tests.test_workflow_runner import make_completed_state  # noqa: E402


CANONICAL_MARKER_PATTERN = r"\[Evidence:\s+[^\s|]+\s+\|\s+[^\]]+\]"
CANONICAL_MARKER_RE = re.compile(CANONICAL_MARKER_PATTERN)
MARKER_DETAIL_RE = re.compile(
    r"\[Evidence:\s+(?P<evidence_id>[^\s|]+)\s+\|\s+(?P<locator>[^\]]+)\]"
)
EVIDENCE_CANDIDATE_RE = re.compile(r"\[Evidence:[^\]]*\]")
MARKDOWN_HEADING_RE = re.compile(r"^\s{0,3}#{1,6}\s+(.+?)\s*$")
LOAD_BEARING_SECTIONS = {
    "executive summary",
    "decision logic",
    "evidence strength",
    "final verdicts",
    "strategy results",
    "monitoring and kill criteria",
}
PRIMARY_PROVIDER = "anthropic"
PRIMARY_REPORT_MODEL = "claude-sonnet-4-6"


@dataclass(frozen=True)
class ValidationCase:
    case_id: str
    state: ProjectState
    min_load_bearing_sections_with_markers: int


def _mark_completed(state: ProjectState) -> ProjectState:
    for phase in (
        "classify",
        "hypotheses",
        "gauntlet",
        "audit",
        "strategy",
        "sqi",
        "monitor",
        "report",
    ):
        state.phase_status[phase] = PhaseStatus.COMPLETED
        state.phase_confidence[phase] = 1.0
    state.current_phase = "report"
    state.report = None
    return state


def _manual_source(source_id: str, name: str, **overrides: Any) -> SourceRegistryEntry:
    payload = {
        "source_id": source_id,
        "name": name,
        "source_kind": "offline_fixture",
        "connector_type": "offline_fixture",
        "access_mode": "manual",
        "sensitivity": "internal",
        "freshness_policy_id": "project_evidence",
        "trust_tier": "operator_curated",
    }
    payload.update(overrides)
    return SourceRegistryEntry(**payload)


def _build_case_a() -> ValidationCase:
    state = _mark_completed(make_state("t1a-gate2-case-a"))
    ensure_knowledge_layer(state)
    now = datetime.now()
    observed = now - timedelta(hours=2)

    sources = [
        _manual_source("src-audit", "Audit fixture"),
        _manual_source("src-strategy", "Strategy fixture"),
        _manual_source("src-provenance", "Official note", trust_tier="official"),
        _manual_source("src-manual", "Operator evidence", trust_tier="analyst_verified"),
        _manual_source("src-official", "Official source", trust_tier="official"),
    ]
    for source in sources:
        upsert_source_entry(state, source)

    # Fixture values are reused from existing repo tests; this script only
    # assembles them into a report-phase validation state.
    sync_offline_source(
        state,
        "src-audit",
        [
            {
                "source_ref": "fixture://audit/eligible",
                "title": "Fresh audit note",
                "summary": "Recent operational signal suggests page-speed regressions in archive traffic.",
                "observed_at": observed.isoformat(),
                "structured_payload": {"region": "mx", "score": 0.77, "trend": "down"},
                "claim_targets": ["phase:audit"],
            }
        ],
        actor="operator",
        requested_at=now,
    )
    sync_offline_source(
        state,
        "src-strategy",
        [
            {
                "source_ref": "fixture://strategy/eligible",
                "title": "Fresh strategy note",
                "summary": "Recent demand signals favor archive refresh before net-new content.",
                "observed_at": observed.isoformat(),
                "structured_payload": {"region": "mx", "score": 0.82},
                "claim_targets": ["phase:strategy"],
            }
        ],
        actor="operator",
        requested_at=now,
    )
    sync_offline_source(
        state,
        "src-provenance",
        [
            {
                "source_ref": "fixture://provenance/1",
                "title": "Official market note",
                "summary": "Demand has shifted toward the recommended plan.",
                "published_at": "2026-04-19T00:00:00",
                "claim_targets": ["phase:strategy", "hypothesis:H1"],
            }
        ],
        actor="operator",
        requested_at=now,
    )
    sync_offline_source(
        state,
        "src-manual",
        [
            {
                "source_ref": "fixture://manual/1",
                "title": "Pricing sheet",
                "summary": "Current pricing sheet supports phased rollout and margin discipline.",
                "published_at": "2026-04-20T00:00:00",
                "claim_targets": ["phase:strategy"],
            }
        ],
        actor="operator",
        requested_at=now,
    )
    sync_offline_source(
        state,
        "src-official",
        [
            {
                "source_ref": "fixture://official/rollout-note",
                "title": "Official rollout note",
                "summary": "Official product note confirms rollout remains supported.",
                "published_at": "2026-04-21T09:55:00",
                "claim_targets": ["phase:strategy"],
            }
        ],
        actor="operator",
        requested_at=now,
    )
    return ValidationCase(
        case_id="case_a_existing_curated_evidence",
        state=state,
        min_load_bearing_sections_with_markers=3,
    )


def _build_case_b() -> ValidationCase:
    state = _mark_completed(make_completed_state("t1a-gate2-case-b"))
    state.knowledge_layer = KnowledgeLayerState(
        items=[
            KnowledgeItem(
                evidence_id="ev-market-note",
                source_id="src-market",
                source_ref="fixture://market-note",
                locator="upload:file-1:market-note.pdf#chunk=2",
                title="Market note",
                provenance=Provenance(external_uri="upload:file-1"),
            ),
            KnowledgeItem(
                evidence_id="ev-no-mutation",
                source_ref="fixture://no-mutation",
                locator="fixture://no-mutation#row=1",
                title="No mutation note",
            ),
            KnowledgeItem(
                evidence_id="ev-no-summary",
                locator="fixture://summary#chunk=1",
                title="Safe title",
            ),
        ]
    )
    if state.hypotheses:
        state.hypotheses[0].evidence_ids = ["ev-no-locator"]
    return ValidationCase(
        case_id="case_b_existing_t1a_marker_fixtures",
        state=state,
        min_load_bearing_sections_with_markers=2,
    )


def _registry_from_prompt(prompt: str) -> dict[str, set[str]]:
    start = prompt.find("PROJECT EVIDENCE LOCATORS:")
    end = prompt.find("MANDATORY REPORT CITATION DISCIPLINE:")
    if start < 0 or end < 0 or end <= start:
        return {}
    section = prompt[start:end]
    registry: dict[str, set[str]] = {}
    for line in section.splitlines():
        stripped = line.strip()
        if not stripped.startswith("- [Evidence:"):
            continue
        match = MARKER_DETAIL_RE.search(stripped)
        if not match:
            continue
        registry.setdefault(match.group("evidence_id"), set()).add(match.group("locator").strip())
    return registry


def _marker_details(report_text: str) -> list[dict[str, Any]]:
    details: list[dict[str, Any]] = []
    for match in MARKER_DETAIL_RE.finditer(report_text or ""):
        details.append(
            {
                "marker": match.group(0),
                "evidence_id": match.group("evidence_id"),
                "locator": match.group("locator").strip(),
                "start": match.start(),
                "end": match.end(),
            }
        )
    return details


def _malformed_candidates(report_text: str) -> list[str]:
    malformed: list[str] = []
    for match in EVIDENCE_CANDIDATE_RE.finditer(report_text or ""):
        candidate = match.group(0)
        detail = MARKER_DETAIL_RE.fullmatch(candidate)
        content = candidate[len("[Evidence:"):-1].strip()
        placeholder = (
            content == "..."
            or "<" in candidate
            or ">" in candidate
            or (
                detail is not None
                and (
                    detail.group("evidence_id").strip().lower() in {"evidence_id", "..."}
                    or detail.group("locator").strip().lower() in {"locator", "..."}
                )
            )
        )
        if not CANONICAL_MARKER_RE.fullmatch(candidate) or placeholder:
            malformed.append(candidate)
    return malformed


def _invented_findings(markers: list[dict[str, Any]], registry: dict[str, set[str]]) -> dict[str, list[dict[str, str]]]:
    unknown_ids: list[dict[str, str]] = []
    locator_mismatches: list[dict[str, str]] = []
    for marker in markers:
        evidence_id = str(marker["evidence_id"])
        locator = str(marker["locator"])
        if evidence_id not in registry:
            unknown_ids.append({"marker": marker["marker"], "evidence_id": evidence_id, "locator": locator})
            continue
        if locator not in registry[evidence_id]:
            locator_mismatches.append(
                {
                    "marker": marker["marker"],
                    "evidence_id": evidence_id,
                    "locator": locator,
                    "registered_locators": sorted(registry[evidence_id]),
                }
            )
    return {"unknown_evidence_ids": unknown_ids, "locator_mismatches": locator_mismatches}


def _load_bearing_counts(report_text: str) -> dict[str, Any]:
    counts = {section: 0 for section in sorted(LOAD_BEARING_SECTIONS)}
    current_section = ""
    for raw_line in (report_text or "").splitlines():
        line = raw_line.strip()
        heading = MARKDOWN_HEADING_RE.match(line)
        if heading:
            normalized = heading.group(1).strip().lower()
            normalized = re.sub(r"\s+\[#\d+\]\s*$", "", normalized).strip()
            current_section = normalized
        if current_section in counts:
            counts[current_section] += len(CANONICAL_MARKER_RE.findall(line))
    return {
        "counts_by_section": counts,
        "total": sum(counts.values()),
        "sections_with_markers": sum(1 for value in counts.values() if value > 0),
    }


def _docx_text(payload: bytes) -> str:
    document = Document(io.BytesIO(payload))
    lines: list[str] = []
    lines.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        lines.append(text)
    return "\n".join(lines)


def _validate_docx_preservation(state: ProjectState, markers: list[str]) -> dict[str, Any]:
    docx_payload = export_project_docx_bytes(state)
    rendered = _docx_text(docx_payload)
    missing = [marker for marker in markers if marker not in rendered]
    return {
        "raw_marker_count": len(markers),
        "preserved_count": len(markers) - len(missing),
        "missing_markers": missing,
        "all_preserved": not missing,
    }


async def _generate_report(case: ValidationCase) -> dict[str, Any]:
    prompt = build_report_prompt(case.state)
    system = build_system_prompt("report", json_mode=False)
    request = GatewayRequest(
        phase="report",
        system_prompt=system,
        user_prompt=prompt,
        project_id="",
        agent_name="report",
        routing_context=RoutingContext(phase="report"),
        allow_cache=False,
    )
    gateway = _get_provider_gateway()
    response = await gateway.call(request, config_override=MODEL_ROUTING["report"])
    return {
        "prompt": prompt,
        "response": response,
    }


async def _run_case(case: ValidationCase) -> dict[str, Any]:
    generated = await _generate_report(case)
    response = generated["response"]
    provider = str(getattr(response, "provider_used", "") or "")
    model = str(getattr(response, "model_used", "") or "")
    fallback_used = bool(getattr(response, "fallback_used", False))
    error = str(getattr(response, "error", "") or "")
    error_type = str(getattr(response, "error_type", "") or "")
    route_known = bool(provider and model and isinstance(fallback_used, bool))
    primary_route = provider == PRIMARY_PROVIDER and model == PRIMARY_REPORT_MODEL and fallback_used is False

    route_status = "primary_anthropic" if primary_route else "not_primary_or_unknown"
    report_text = str(getattr(response, "text", "") or "")
    case.state.report = report_text

    registry = _registry_from_prompt(generated["prompt"])
    raw_markers = CANONICAL_MARKER_RE.findall(report_text)
    marker_details = _marker_details(report_text)
    malformed = _malformed_candidates(report_text)
    invented = _invented_findings(marker_details, registry)
    load_bearing = _load_bearing_counts(report_text)
    docx = _validate_docx_preservation(case.state, raw_markers) if report_text else {
        "raw_marker_count": 0,
        "preserved_count": 0,
        "missing_markers": [],
        "all_preserved": False,
    }

    validation_failures: list[str] = []
    if not report_text:
        validation_failures.append("report_generation_returned_empty_text")
    if not raw_markers:
        validation_failures.append("no_canonical_markers_in_raw_report")
    if malformed:
        validation_failures.append("malformed_evidence_markers_present")
    if invented["unknown_evidence_ids"]:
        validation_failures.append("invented_or_unknown_evidence_ids_present")
    if invented["locator_mismatches"]:
        validation_failures.append("invented_or_mismatched_locators_present")
    if load_bearing["sections_with_markers"] < case.min_load_bearing_sections_with_markers:
        validation_failures.append("insufficient_load_bearing_sections_with_markers")
    if not docx["all_preserved"]:
        validation_failures.append("docx_did_not_preserve_all_raw_markers")

    if not route_known or error or not primary_route:
        gate_status = "INCONCLUSIVE"
    elif validation_failures:
        gate_status = "FAIL"
    else:
        gate_status = "PASS"

    return {
        "case_id": case.case_id,
        "provider_status": {
            "provider": provider,
            "model": model,
            "fallback_used": fallback_used,
            "route_status": route_status,
            "route_known": route_known,
            "error": error,
            "error_type": error_type,
            "input_tokens": int(getattr(response, "input_tokens", 0) or 0),
            "output_tokens": int(getattr(response, "output_tokens", 0) or 0),
            "cache_read_tokens": int(getattr(response, "cache_read_tokens", 0) or 0),
            "cost_usd": float(getattr(response, "cost_usd", 0.0) or 0.0),
            "latency_ms": int(getattr(response, "latency_ms", 0) or 0),
        },
        "registry": {
            "evidence_id_count": len(registry),
            "evidence_ids": sorted(registry),
        },
        "raw_report_validation": {
            "canonical_marker_regex": CANONICAL_MARKER_PATTERN,
            "canonical_marker_count": len(raw_markers),
            "canonical_markers": raw_markers,
            "load_bearing_marker_counts": load_bearing,
            "malformed_marker_count": len(malformed),
            "malformed_markers": malformed,
            "invented_findings": invented,
        },
        "docx_validation": docx,
        "validation_failures": validation_failures,
        "gate_status": gate_status,
    }


def _overall_status(case_results: list[dict[str, Any]]) -> str:
    statuses = [case["gate_status"] for case in case_results]
    if any(status == "INCONCLUSIVE" for status in statuses):
        return "INCONCLUSIVE"
    if any(status == "FAIL" for status in statuses):
        return "FAIL"
    return "PASS"


async def main() -> int:
    cases = [_build_case_a(), _build_case_b()]
    case_results = []
    for case in cases:
        case_results.append(await _run_case(case))
    gate_status = _overall_status(case_results)
    payload = {
        "validation": "t1a_gate2_primary_anthropic",
        "generated_at": datetime.now().isoformat(),
        "scope": {
            "t1b_implemented": False,
            "cdp_implemented": False,
            "product_code_modified_by_script": False,
            "persistent_project_state_written": False,
            "docx_checked_for_marker_preservation_only": True,
        },
        "primary_route_required": {
            "provider": PRIMARY_PROVIDER,
            "model": PRIMARY_REPORT_MODEL,
            "fallback_used": False,
        },
        "case_results": case_results,
        "gate_2_status": gate_status,
        "t1a_behaviorally_validated_under_primary_anthropic": gate_status == "PASS",
        "t1b_safe_to_plan_only": gate_status == "PASS",
    }
    print(json.dumps(payload, indent=2, sort_keys=True))
    return 0 if gate_status == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(asyncio.run(main()))
